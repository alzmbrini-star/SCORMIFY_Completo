"""Quiz questions CRUD, generation and submission routes"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Depends
from typing import List, Optional, Any
import uuid
import os
import json
import logging
import re
import aiofiles

from routes.deps import db, now_utc, serialize_doc, UPLOADS_DIR
from routes.auth import require_auth
from routes.projects_common import load_authorized_project
from models import (
    QuizQuestion, QuizQuestionCreate, QuizQuestionUpdate, QuizAlternative,
    QuizConfig, QuizAttempt, QuizGenerateRequest, QuizSubmitRequest
)

logger = logging.getLogger("server")

router = APIRouter(tags=["Questions"])


def _quiz_generation_credentials() -> tuple[str, str]:
    """Use the OpenAI secret already configured for the Editor and Tutor."""
    api_key = os.environ.get("OPENAI_API_KEY", "").strip()
    model = (
        os.environ.get("OPENAI_QUIZ_MODEL", "").strip()
        or os.environ.get("OPENAI_TEXT_MODEL", "").strip()
        or "gpt-4o"
    )
    return api_key, model


def _extract_quiz_json(value: object) -> dict:
    """Extract one JSON object without trusting surrounding model prose."""
    text = str(value or "").strip()
    fenced = re.search(r"```(?:json)?\s*([\s\S]*?)```", text, re.IGNORECASE)
    if fenced:
        text = fenced.group(1).strip()
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        start, end = text.find("{"), text.rfind("}")
        if start < 0 or end <= start:
            raise ValueError("A OpenAI não retornou um objeto JSON")
        parsed = json.loads(text[start:end + 1])
    if not isinstance(parsed, dict):
        raise ValueError("A resposta da OpenAI não contém um objeto JSON")
    return parsed


def _clean_quiz_text(value: Any, max_length: int) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()[:max_length]


def _normalize_generated_questions(
    payload: dict,
    requested_type: str,
    requested_count: int,
) -> list[dict]:
    """Validate AI output before any generated question is persisted."""
    raw_questions = payload.get("questions")
    if not isinstance(raw_questions, list):
        raise ValueError("A resposta não contém uma lista de questões")

    normalized: list[dict] = []
    for raw in raw_questions[:requested_count]:
        if not isinstance(raw, dict):
            continue
        text = _clean_quiz_text(raw.get("text") or raw.get("questionText"), 1000)
        explanation = _clean_quiz_text(raw.get("explanation"), 3000)
        if not text:
            continue

        raw_type = str(raw.get("type") or "").strip()
        q_type = requested_type if requested_type != "mixed" else raw_type
        if q_type not in ("multiple_choice", "true_false"):
            continue

        raw_alternatives = raw.get("alternatives")
        if not isinstance(raw_alternatives, list):
            continue
        alternatives = []
        for raw_alt in raw_alternatives:
            if not isinstance(raw_alt, dict):
                continue
            alt_text = _clean_quiz_text(raw_alt.get("text"), 500)
            if alt_text:
                alternatives.append(
                    {"text": alt_text, "isCorrect": raw_alt.get("isCorrect") is True}
                )

        expected_alternatives = 2 if q_type == "true_false" else 4
        if len(alternatives) != expected_alternatives:
            continue
        if sum(1 for alt in alternatives if alt["isCorrect"]) != 1:
            continue

        if q_type == "true_false":
            true_alt = next(
                (
                    alt
                    for alt in alternatives
                    if alt["text"].casefold() in ("verdadeiro", "true")
                ),
                None,
            )
            false_alt = next(
                (
                    alt
                    for alt in alternatives
                    if alt["text"].casefold() in ("falso", "false")
                ),
                None,
            )
            if not true_alt or not false_alt:
                continue
            correct_is_true = true_alt["isCorrect"]
            alternatives = [
                {"text": "Verdadeiro", "isCorrect": correct_is_true},
                {"text": "Falso", "isCorrect": not correct_is_true},
            ]

        normalized.append(
            {
                "type": q_type,
                "text": text,
                "alternatives": alternatives,
                "explanation": explanation,
            }
        )

    if len(normalized) != requested_count:
        raise ValueError(
            f"A OpenAI retornou {len(normalized)} de {requested_count} questões válidas"
        )
    return normalized


@router.get("/questions")
async def list_questions(project_id: Optional[str] = None):
    query = {"projectId": project_id} if project_id else {}
    docs = await db.questions.find(query, {"_id": 0}).to_list(1000)
    return docs


@router.get("/questions/{question_id}")
async def get_question(question_id: str):
    doc = await db.questions.find_one({"id": question_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Question not found")
    return doc


@router.post("/questions")
async def create_question(data: QuizQuestionCreate):
    from models import generate_id
    q = {
        "id": generate_id(),
        "projectId": data.projectId,
        "slideId": data.slideId,
        "type": data.type,
        "questionText": data.questionText,
        "alternatives": [a.model_dump() for a in data.alternatives] if data.alternatives else [],
        "correctAnswer": data.correctAnswer,
        "explanation": data.explanation or "",
        "points": data.points,
        "order": data.order,
        "createdAt": now_utc().isoformat(),
    }
    await db.questions.insert_one(q)
    q.pop('_id', None)
    return q


@router.put("/questions/{question_id}")
async def update_question(question_id: str, data: QuizQuestionUpdate):
    existing = await db.questions.find_one({"id": question_id})
    if not existing:
        raise HTTPException(404, "Question not found")
    update = data.model_dump(exclude_unset=True)
    if "alternatives" in update and update["alternatives"]:
        update["alternatives"] = [a if isinstance(a, dict) else a.model_dump() for a in update["alternatives"]]
    update["updatedAt"] = now_utc().isoformat()
    await db.questions.update_one({"id": question_id}, {"$set": update})
    doc = await db.questions.find_one({"id": question_id}, {"_id": 0})
    return doc


@router.delete("/questions/{question_id}")
async def delete_question(question_id: str):
    result = await db.questions.delete_one({"id": question_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Question not found")
    return {"status": "deleted"}


@router.post("/questions/generate")
async def generate_questions_with_ai(
    request: QuizGenerateRequest,
    user: dict = Depends(require_auth),
):
    project = None
    if request.projectId:
        project = await load_authorized_project(request.projectId, user)

    api_key, model = _quiz_generation_credentials()
    if not api_key:
        raise HTTPException(
            status_code=503,
            detail=(
                "A geração de questões ainda não possui uma chave OpenAI "
                "configurada. Cadastre OPENAI_API_KEY no backend do Render."
            ),
        )

    prompt_text = (request.prompt or "").strip()
    document_text = (request.documentContent or "").strip()
    if request.source == "document" and not document_text:
        raise HTTPException(status_code=400, detail="Envie um documento com conteúdo.")
    if request.source == "prompt" and not prompt_text:
        raise HTTPException(status_code=400, detail="Informe o tema das questões.")

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        if request.questionType == "true_false":
            type_instruction = """Gere apenas questões de Verdadeiro ou Falso. Cada questão deve ter type "true_false" e exatamente 2 alternativas: "Verdadeiro" e "Falso"."""
        elif request.questionType == "multiple_choice":
            type_instruction = """Gere apenas questões de múltipla escolha. Cada questão deve ter exatamente 4 alternativas, sendo apenas 1 correta."""
        else:
            type_instruction = """Gere uma mistura de questões de múltipla escolha e Verdadeiro/Falso."""
        system_message = f"""Você é um especialista em avaliação educacional.
{type_instruction}
Crie exatamente {request.count} questões.
REGRAS:
1. Responda em português brasileiro.
2. Crie questões claras, objetivas e baseadas somente no tema/contexto fornecido.
3. Evite ambiguidades, pegadinhas e alternativas obviamente incorretas.
4. Inclua uma explicação didática para cada resposta correta.
5. Cada questão deve ter exatamente uma alternativa correta.
6. Retorne somente JSON válido, sem Markdown ou comentários.

FORMATO:
{{"questions": [
  {{"type": "multiple_choice", "text": "Pergunta", "alternatives": [{{"text": "Alternativa A", "isCorrect": false}}, {{"text": "Alternativa B", "isCorrect": true}}, {{"text": "Alternativa C", "isCorrect": false}}, {{"text": "Alternativa D", "isCorrect": false}}], "explanation": "Explicação"}},
  {{"type": "true_false", "text": "Afirmação", "alternatives": [{{"text": "Verdadeiro", "isCorrect": true}}, {{"text": "Falso", "isCorrect": false}}], "explanation": "Explicação"}}
]}}
"""
        chat = (
            LlmChat(
                api_key=api_key,
                session_id=f"quiz-gen-{uuid.uuid4()}",
                system_message=system_message,
            )
            .with_model("openai", model)
            .with_params(temperature=0.2)
        )
        if request.source == "document" and request.documentContent:
            full_prompt = (
                f"Gere {request.count} questões com base exclusivamente neste documento:\n\n"
                f"{document_text}"
            )
        else:
            full_prompt = f"Gere {request.count} questões sobre: {prompt_text}"
            if request.context:
                full_prompt += f"\n\nContexto adicional: {request.context.strip()}"
        if project:
            project_title = _clean_quiz_text(
                project.get("name") or project.get("title"), 300
            )
            if project_title:
                full_prompt += f"\n\nCurso: {project_title}"

        response = await chat.send_message(UserMessage(text=full_prompt))
        parsed = _extract_quiz_json(response)
        questions_data = _normalize_generated_questions(
            parsed,
            request.questionType,
            request.count,
        )

        prepared_questions = []
        company_id = project.get("companyId") if project else user.get("companyId")
        for q_data in questions_data:
            alternatives = [
                QuizAlternative(**alt).model_dump()
                for alt in q_data["alternatives"]
            ]
            question = QuizQuestion(
                projectId=request.projectId,
                type=q_data["type"],
                text=q_data["text"],
                alternatives=alternatives,
                explanation=q_data["explanation"],
                tags=["ai-generated"],
            )
            question_dict = question.model_dump()
            question_dict["createdAt"] = question.createdAt.isoformat()
            question_dict["updatedAt"] = question.updatedAt.isoformat()
            if company_id:
                question_dict["companyId"] = company_id
            prepared_questions.append(question_dict)

        saved_questions = []
        for question_dict in prepared_questions:
            await db.questions.insert_one(question_dict)
            saved_questions.append(serialize_doc(question_dict))
        return {"success": True, "questions": saved_questions, "count": len(saved_questions)}
    except (json.JSONDecodeError, ValueError) as exc:
        logger.warning("Quiz generation returned invalid data: %s", exc)
        raise HTTPException(
            status_code=422,
            detail=(
                "A OpenAI retornou questões em formato incompleto. "
                "Tente novamente ou descreva o tema com mais detalhes."
            ),
        )
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="A integração de IA não está disponível no servidor.",
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Quiz generation error: {e}")
        from routes.ai_gen import _friendly_text_generation_error

        status_code, detail = _friendly_text_generation_error(e)
        raise HTTPException(status_code=status_code, detail=detail)


@router.post("/questions/parse-doc")
async def parse_doc_file(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(('.doc', '.docx')):
        raise HTTPException(status_code=400, detail="Only .doc and .docx files are accepted")
    try:
        from docx import Document
        content = await file.read()
        temp_path = UPLOADS_DIR / f"temp_{uuid.uuid4()}_{file.filename}"
        async with aiofiles.open(temp_path, 'wb') as f:
            await f.write(content)
        doc = Document(str(temp_path))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        os.remove(temp_path)
        extracted_text = "\n\n".join(full_text)
        return {"success": True, "filename": file.filename, "text": extracted_text, "wordCount": len(extracted_text.split())}
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library not installed")
    except Exception as e:
        logger.error(f"Doc parsing error: {e}")
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")


@router.post("/quiz/submit")
async def submit_quiz(request: QuizSubmitRequest):
    question_ids = [a["questionId"] for a in request.answers]
    questions = await db.questions.find({"id": {"$in": question_ids}}, {"_id": 0}).to_list(100)
    questions_map = {q["id"]: q for q in questions}
    total_points = 0
    earned_points = 0
    results = []
    for answer in request.answers:
        question = questions_map.get(answer["questionId"])
        if not question:
            continue
        total_points += question.get("points", 1)
        correct_alt = None
        selected_alt = None
        for alt in question.get("alternatives", []):
            if alt.get("isCorrect"):
                correct_alt = alt
            if alt.get("id") == answer.get("selectedAlternativeId"):
                selected_alt = alt
        is_correct = selected_alt and selected_alt.get("isCorrect", False)
        if is_correct:
            earned_points += question.get("points", 1)
        results.append({
            "questionId": answer["questionId"], "questionText": question.get("text"),
            "selectedAlternativeId": answer.get("selectedAlternativeId"),
            "selectedText": selected_alt.get("text") if selected_alt else None,
            "correctAlternativeId": correct_alt.get("id") if correct_alt else None,
            "correctText": correct_alt.get("text") if correct_alt else None,
            "isCorrect": is_correct, "explanation": question.get("explanation")
        })
    percentage = (earned_points / total_points * 100) if total_points > 0 else 0
    final_score = round(percentage / 10, 1)
    passed = percentage >= 60
    attempt = QuizAttempt(quizId=request.quizId, projectId="", answers=results, score=final_score, percentage=round(percentage, 1), passed=passed, completedAt=now_utc())
    attempt_dict = attempt.model_dump()
    attempt_dict['createdAt'] = attempt.createdAt.isoformat()
    attempt_dict['completedAt'] = attempt.completedAt.isoformat() if attempt.completedAt else None
    await db.quiz_attempts.insert_one(attempt_dict)
    return {"success": True, "attemptId": attempt.id, "score": final_score, "percentage": round(percentage, 1), "passed": passed, "totalQuestions": len(results), "correctAnswers": sum(1 for r in results if r["isCorrect"]), "results": results}


@router.get("/quiz/attempts/{project_id}")
async def get_quiz_attempts(project_id: str, quiz_id: Optional[str] = None):
    query = {"projectId": project_id}
    if quiz_id:
        query["quizId"] = quiz_id
    attempts = await db.quiz_attempts.find(query, {"_id": 0}).sort("createdAt", -1).to_list(100)
    return attempts
