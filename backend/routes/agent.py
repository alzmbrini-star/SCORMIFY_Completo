"""AI Instructional Design Agent routes"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Request, Depends, BackgroundTasks
from pydantic import BaseModel, ConfigDict as PydanticConfigDict
from typing import Optional, List, Dict, Any
import uuid
import os
import asyncio
import json
import logging
import base64
import re
import httpx
from pathlib import Path
from datetime import datetime, timezone

from routes.deps import (
    db, now_utc, serialize_doc, get_project_by_id, update_project,
    PROJECTS_DIR, STORAGE_DIR, mongo_url, ELEVENLABS_API_KEY, HEYGEN_API_KEY,
    HEYGEN_BASE_URL, HEYGEN_HEADERS
)
from routes.auth import require_agent_access, require_auth, get_current_user

logger = logging.getLogger("server")

router = APIRouter(tags=["Agent"])


# =============================================================================
# AI INSTRUCTIONAL DESIGN AGENT
# =============================================================================

class AgentSessionCreate(BaseModel):
    """Create agent session - optionally with initial text content"""
    contentText: Optional[str] = None
    fileName: Optional[str] = None

class GenerateHtmlRequest(BaseModel):
    prompt: str
    courseContext: Optional[str] = None

class AgentConfigUpdate(BaseModel):
    model_config = PydanticConfigDict(extra="allow")
    title: Optional[str] = None
    depth: Optional[str] = "intermediario"
    duration: Optional[int] = 30
    modules: Optional[int] = 3
    interactivity: Optional[str] = "media"
    visualStyle: Optional[str] = "moderno e profissional"
    format: Optional[str] = "curso_completo"
    description: Optional[str] = None

class AgentChatMessage(BaseModel):
    message: str


def _is_light_color(hex_color: str) -> bool:
    """Check if a hex color is light (luminance > 0.5)."""
    try:
        c = hex_color.lstrip('#')
        if len(c) < 6:
            return True
        r, g, b = int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
        luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
        return luminance > 0.5
    except Exception:
        return False


def _apply_design_token_to_slide(slide: dict, design_token: dict, sb_slide: dict = None):
    """Apply a design template's visual styling to a single slide."""
    import re as _re
    palette = design_token["palette"]
    font_heading = design_token["fonts"]["heading"]
    font_body = design_token["fonts"]["body"]
    header_style = design_token.get("headerStyle", "solid")
    corner_radius = design_token.get("cornerRadius", "12px")
    primary = palette["primary"]
    accent = palette["accent"]
    content_bg = palette.get("contentBg", "#f0fdf4")
    text_color = palette.get("text", "#1e293b")

    slide_type = slide.get("type", "")

    # Better slide type detection for slides without explicit type
    if not slide_type:
        title = (slide.get("title") or "").lower()
        if any(k in title for k in ("capa", "cover", "título", "title", "intro")):
            slide_type = "title"
        elif any(k in title for k in ("quiz", "prova", "teste", "avaliação")):
            slide_type = "quiz"
        elif any(k in title for k in ("resumo", "summary", "conclus")):
            slide_type = "summary"
        else:
            slide_type = "content"

    # Set background based on slide type
    is_dark_slide = slide_type in ("title", "cover", "quiz", "summary")
    if is_dark_slide:
        slide["background"] = primary
        slide_text_color = "#ffffff"
        slide_text_muted = "rgba(255,255,255,0.75)"
    else:
        slide["background"] = content_bg
        # Determine text color based on background lightness
        if _is_light_color(content_bg):
            slide_text_color = text_color
            slide_text_muted = text_color + "cc"
        else:
            slide_text_color = "#ffffff"
            slide_text_muted = "rgba(255,255,255,0.75)"

    # Update elements
    for el in slide.get("elements", []):
        # Update header bars (html element at top, reasonable height, wide)
        if el.get("type") == "html" and el.get("y", -1) <= 5 and el.get("height", 0) <= 55 and el.get("height", 0) >= 40 and el.get("width", 0) >= 1700:
            from services.ai_agent import _build_header_bar
            pal = {**palette, "fontHeading": font_heading, "fontBody": font_body, "headerStyle": header_style}
            module_name = ""
            slide_title = ""
            hc = el.get("htmlContent", "")
            import re
            spans = re.findall(r'>([^<]+)<', hc)
            if spans:
                module_name = spans[0].strip()
                if len(spans) > 1:
                    slide_title = spans[1].strip()
            el["htmlContent"] = _build_header_bar(pal, module_name, slide_title)
            continue

        # Update fonts AND text colors in html content
        if el.get("type") in ("html", "text") and el.get("htmlContent"):
            html = el["htmlContent"]
            # Replace font-family
            html = _re.sub(r"font-family:[^;\"']+", f"font-family:{font_body}", html)
            html = _re.sub(r'(<h[1-3][^>]*style="[^"]*?)font-family:[^;"]+', lambda m: m.group(1) + f"font-family:{font_heading}", html)

            # Replace text colors in inline styles
            def _replace_text_color(match):
                prefix = match.group(1)
                old_color = match.group(2).strip()
                # Skip accent-colored elements
                if accent.lower().replace('#', '') in old_color.lower().replace('#', '').replace(' ', ''):
                    return match.group(0)
                # Simple rule: dark slide = white text, light slide = dark text
                if is_dark_slide:
                    return prefix + "#ffffff"
                else:
                    return prefix + slide_text_color

            html = _re.sub(r'((?:^|;|\s|")color\s*:\s*)(#[0-9a-fA-F]{3,8}|rgba?\([^)]+\))', _replace_text_color, html)
            el["htmlContent"] = html

            if el.get("style"):
                el["style"]["fontFamily"] = font_body

        # Update image corner radius
        if el.get("type") == "image":
            if not el.get("style"):
                el["style"] = {}
            el["style"]["borderRadius"] = corner_radius

        # Update accent colors in small decorative elements
        if el.get("type") == "html" and el.get("htmlContent") and el.get("height", 0) <= 12:
            hc = el["htmlContent"]
            if "background:" in hc:
                hc = _re.sub(r'background:#[0-9a-fA-F]{3,8}', f'background:{accent}', hc)
                el["htmlContent"] = hc



@router.post("/generate-html")
async def generate_html_with_ai(body: GenerateHtmlRequest, request: Request):
    """Generate interactive HTML+JS content from a text prompt using Gemini."""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Não autenticado")
    
    emergent_key = os.environ.get("EMERGENT_LLM_KEY", "")
    if not emergent_key:
        raise HTTPException(status_code=500, detail="Chave de IA não configurada")
    
    system_msg = """Você é um especialista em desenvolvimento web front-end.
Sua tarefa é gerar código HTML completo, autocontido, com CSS e JavaScript embutidos.

REGRAS OBRIGATÓRIAS:
- Retorne APENAS o código HTML, sem explicações, sem markdown, sem blocos ```html
- O HTML deve ser COMPLETO e autocontido (todo CSS e JS inline/embutido)
- Use design moderno, bonito e responsivo
- Use cores vibrantes e profissionais
- O JavaScript DEVE ser funcional e interativo
- Inclua animações CSS quando apropriado
- Use fontes do Google Fonts quando necessário (via link CDN)
- Todo texto deve ser em português brasileiro
- O conteúdo deve caber em um container de ~800x500px
- NÃO use bibliotecas externas pesadas, apenas vanilla JS, CSS e CDNs leves
- Inclua tratamento de erros no JavaScript
- Use flexbox/grid para layouts modernos"""

    context_hint = ""
    if body.courseContext:
        context_hint = f"\n\nContexto do curso: {body.courseContext}"
    
    prompt = f"""Crie o seguinte conteúdo HTML interativo:{context_hint}

Solicitação do usuário: {body.prompt}

Retorne APENAS o código HTML completo, começando com <div> ou <!DOCTYPE html>. Sem explicações."""

    try:
        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"html-gen-{uuid.uuid4().hex[:8]}",
            system_message=system_msg,
        ).with_model("gemini", "gemini-3-flash-preview")
        
        resp = await chat.send_message(UserMessage(text=prompt))
        html_code = resp.strip() if isinstance(resp, str) else str(resp).strip()
        
        # Clean up markdown wrappers if the model wraps in ```html blocks
        if html_code.startswith("```"):
            lines = html_code.split("\n")
            if lines[0].strip().startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            html_code = "\n".join(lines)
        
        return {"html": html_code}
    except Exception as e:
        logger.error(f"HTML generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao gerar HTML: {str(e)}")


@router.get("/agent/check-access")
async def check_agent_access(request: Request):
    """Check if the current user has access to the AI Agent feature."""
    user = await get_current_user(request)
    if not user:
        return {"hasAccess": False, "reason": "not_authenticated"}
    
    # Super admins always have access
    if user.get("role") == "super_admin":
        return {"hasAccess": True, "reason": "super_admin"}
    
    # Check company permissions
    company = user.get("company")
    if not company:
        return {"hasAccess": False, "reason": "no_company"}
    
    if company.get("permissions", {}).get("agentAccess", False):
        return {"hasAccess": True, "reason": "company_permission"}
    
    return {"hasAccess": False, "reason": "not_authorized"}

@router.post("/agent/sessions")
async def create_agent_session(data: AgentSessionCreate, request: Request, user: dict = Depends(require_agent_access)):
    """Create a new AI agent session."""
    session_id = str(uuid.uuid4())
    session = {
        "id": session_id,
        "step": "created",  # created, analyzed, configured, structured, storyboarded, generated
        "contentText": data.contentText or "",
        "fileName": data.fileName or "",
        "analysis": None,
        "config": {},
        "structure": None,
        "storyboard": None,
        "projectId": None,
        "chatHistory": [],
        "userId": user.get("user_id"),
        "companyId": user.get("companyId"),
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    await db.agent_sessions.insert_one({**session, "_id": session_id})
    return session

@router.get("/agent/sessions/{session_id}")
async def get_agent_session(session_id: str, request: Request, user: dict = Depends(require_agent_access)):
    """Get agent session state."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    return s

@router.post("/agent/sessions/{session_id}/upload")
async def agent_upload_content(session_id: str, request: Request, file: UploadFile = File(None), text: str = Form(None), url: str = Form(None), user: dict = Depends(require_agent_access)):
    """Upload content to agent session (file, text, or URL)."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    content_text = text or ""
    file_name = ""

    # Handle URL scraping
    if url and url.strip():
        file_name = url.strip()
        try:
            async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
                resp = await client.get(url.strip(), headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "Accept": "text/html,application/xhtml+xml",
                    "Accept-Language": "pt-BR,pt;q=0.9,en;q=0.8",
                })
                if resp.status_code == 200:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(resp.text, "html.parser")
                    # Remove scripts, styles, nav, footer
                    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe"]):
                        tag.decompose()
                    # Try article/main content first
                    main = soup.find("article") or soup.find("main") or soup.find("body")
                    if main:
                        content_text = main.get_text(separator="\n", strip=True)
                    else:
                        content_text = soup.get_text(separator="\n", strip=True)
                    # Clean up excessive whitespace
                    import re as _re
                    content_text = _re.sub(r'\n{3,}', '\n\n', content_text).strip()
                    content_text = content_text[:50000]  # Limit
                else:
                    content_text = f"[Erro ao acessar URL: HTTP {resp.status_code}]"
        except Exception as e:
            logger.warning(f"URL scraping failed: {e}")
            content_text = f"[Erro ao acessar URL: {str(e)[:200]}]"

    # Handle file upload
    elif file:
        file_name = file.filename or ""
        file_bytes = await file.read()
        ext = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else ""

        if ext == "txt":
            content_text = file_bytes.decode("utf-8", errors="ignore")

        elif ext == "pdf":
            # Native PDF extraction
            try:
                from PyPDF2 import PdfReader
                import io as _io
                reader = PdfReader(_io.BytesIO(file_bytes))
                pages = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
                content_text = "\n\n".join(pages)
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}")
                # Fallback to ConvertAPI
                content_text = await _convert_api_extract(file_name, file_bytes, ext)

        elif ext in ("docx",):
            # Native DOCX extraction
            try:
                from docx import Document
                import io as _io
                doc = Document(_io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            paragraphs.append(" | ".join(cells))
                content_text = "\n\n".join(paragraphs)
            except Exception as e:
                logger.warning(f"python-docx extraction failed: {e}")
                content_text = await _convert_api_extract(file_name, file_bytes, ext)

        elif ext in ("pptx", "ppt", "doc"):
            # Use ConvertAPI for PPT/legacy doc
            content_text = await _convert_api_extract(file_name, file_bytes, ext)

        else:
            content_text = file_bytes.decode("utf-8", errors="ignore")

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"contentText": content_text, "fileName": file_name, "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "ok", "contentLength": len(content_text), "fileName": file_name}


async def _convert_api_extract(file_name: str, file_bytes: bytes, ext: str) -> str:
    """Fallback: use ConvertAPI to extract text from a file."""
    convert_secret = os.environ.get("CONVERTAPI_SECRET", "")
    if not convert_secret:
        return f"[ConvertAPI não configurada - arquivo {file_name} recebido]"
    try:
        async with httpx.AsyncClient(timeout=120) as client:
            files_payload = {"File": (file_name, file_bytes)}
            resp = await client.post(
                f"https://v2.convertapi.com/convert/{ext}/to/txt?Secret={convert_secret}",
                files=files_payload,
            )
            if resp.status_code == 200:
                result = resp.json()
                text = ""
                for f_item in result.get("Files", []):
                    b64 = f_item.get("FileData", "")
                    if b64:
                        text += base64.b64decode(b64).decode("utf-8", errors="ignore")
                return text
    except Exception as e:
        logger.warning(f"ConvertAPI text extraction failed: {e}")
    return f"[Erro ao extrair texto de {file_name}]"

@router.post("/agent/sessions/{session_id}/analyze")
async def agent_analyze(session_id: str):
    """Step 1: Analyze content with AI."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    if not s.get("contentText"):
        raise HTTPException(400, "No content uploaded")

    from services.ai_agent import analyze_content
    analysis = await analyze_content(session_id, s["contentText"], s.get("fileName", ""))

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"analysis": analysis, "step": "analyzed", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return analysis

@router.post("/agent/sessions/{session_id}/configure")
async def agent_configure(session_id: str, data: AgentConfigUpdate):
    """Set course configuration parameters."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    config = data.model_dump(exclude_unset=True)
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"config": config, "step": "configured", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "ok", "config": config}

@router.post("/agent/sessions/{session_id}/generate-structure")
async def agent_generate_structure(session_id: str, request: Request):
    """Step 2: Generate course structure (optionally from a template)."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    # Check if body contains a templateId
    template_id = None
    try:
        body = await request.json()
        template_id = body.get("templateId")
    except Exception:
        pass

    if template_id:
        from services.ai_agent import generate_structure_from_template
        structure = await generate_structure_from_template(session_id, s["contentText"], s.get("config", {}), template_id)
        if not structure:
            from services.ai_agent import generate_structure
            structure = await generate_structure(session_id, s["contentText"], s.get("config", {}))
    else:
        from services.ai_agent import generate_structure
        structure = await generate_structure(session_id, s["contentText"], s.get("config", {}))

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"structure": structure, "step": "structured", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return structure

@router.post("/agent/sessions/{session_id}/generate-storyboard")
async def agent_generate_storyboard(session_id: str, background_tasks: BackgroundTasks):
    """Step 3: Generate detailed storyboard (runs in thread pool to avoid blocking event loop)."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    if not s.get("structure"):
        raise HTTPException(400, "Structure not generated yet")

    # If storyboard already exists, return it
    if s.get("step") == "storyboarded" and s.get("storyboard"):
        return {"status": "already_done", "message": "Storyboard already generated"}

    # If already generating, don't start another thread
    if s.get("step") == "storyboarding":
        return {"status": "processing", "message": "Storyboard generation already in progress"}

    # Mark as processing
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"step": "storyboarding", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )

    def _sync_generate():
        """Run storyboard generation in its own event loop inside a thread."""
        import asyncio as _asyncio
        loop = _asyncio.new_event_loop()
        _asyncio.set_event_loop(loop)
        try:
            from services.ai_agent import generate_storyboard
            from motor.motor_asyncio import AsyncIOMotorClient
            _client = AsyncIOMotorClient(os.environ.get("MONGO_URL"), serverSelectionTimeoutMS=30000, connectTimeoutMS=30000)
            _db = _client[os.environ.get("DB_NAME")]

            async def _progress(batch_num, total, message):
                await _db.agent_sessions.update_one(
                    {"id": session_id},
                    {"$set": {
                        "storyboardProgress": {"batch": batch_num, "total": total, "message": message},
                        "updatedAt": datetime.now(timezone.utc).isoformat()
                    }}
                )

            storyboard = loop.run_until_complete(
                generate_storyboard(session_id, s["contentText"], s["structure"], s.get("config", {}), progress_callback=_progress)
            )
            loop.run_until_complete(
                _db.agent_sessions.update_one(
                    {"id": session_id},
                    {"$set": {"storyboard": storyboard, "step": "storyboarded", "storyboardProgress": None, "updatedAt": datetime.now(timezone.utc).isoformat()}}
                )
            )
            _client.close()
        except Exception as e:
            err_msg = str(e)
            logger.error(f"Storyboard generation error: {err_msg}")
            error_detail = "Erro ao gerar storyboard."
            if "Budget" in err_msg or "budget" in err_msg:
                error_detail = "Orçamento da chave LLM excedido. Acesse Profile > Universal Key > Add Balance para adicionar saldo."
            elif "502" in err_msg or "BadGateway" in err_msg:
                error_detail = "Serviço de IA temporariamente indisponível (502). Tente novamente em alguns minutos."
            try:
                from motor.motor_asyncio import AsyncIOMotorClient
                _client = AsyncIOMotorClient(os.environ.get("MONGO_URL"), serverSelectionTimeoutMS=30000, connectTimeoutMS=30000)
                _db = _client[os.environ.get("DB_NAME")]
                loop.run_until_complete(
                    _db.agent_sessions.update_one(
                        {"id": session_id},
                        {"$set": {"step": "structured", "error": error_detail, "storyboardProgress": None, "updatedAt": datetime.now(timezone.utc).isoformat()}}
                    )
                )
                _client.close()
            except Exception:
                pass
        finally:
            loop.close()

    # Run in thread pool - this keeps the main event loop free for other requests
    import threading
    thread = threading.Thread(target=_sync_generate, daemon=True)
    thread.start()
    return {"status": "processing", "message": "Storyboard being generated..."}


@router.post("/agent/generate-bg-image")
async def agent_generate_bg_image(data: dict):
    """Generate a background image using Gemini Nano Banana based on a text prompt."""
    prompt = data.get("prompt", "")
    if not prompt:
        raise HTTPException(400, "Prompt required")
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        import uuid as _uuid
        emergent_key = os.environ.get("EMERGENT_LLM_KEY", "")
        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"bg_{_uuid.uuid4().hex[:8]}",
            system_message="You are an image generator.",
        ).with_model("gemini", "gemini-3-pro-image-preview").with_params(modalities=["image", "text"])

        full_prompt = f"Abstract artistic background for presentation slide: {prompt}. No text, no watermarks, smooth gradients, professional, suitable as slide background."
        text_resp, images = await chat.send_message_multimodal_response(UserMessage(text=full_prompt))
        if images and len(images) > 0:
            img_bytes = base64.b64decode(images[0]['data'])
            import hashlib
            seed = hashlib.md5(prompt.encode()).hexdigest()[:10]
            fname = f"bg_ai_{seed}.png"
            fpath = os.path.join(str(PROJECTS_DIR), "bg_temp", fname)
            os.makedirs(os.path.dirname(fpath), exist_ok=True)
            with open(fpath, "wb") as f:
                f.write(img_bytes)
            
            # Persist in MongoDB for production environments with ephemeral storage (non-blocking)
            import threading
            try:
                from services.asset_store import store_asset_sync
                threading.Thread(
                    target=store_asset_sync,
                    args=(mongo_url, os.environ['DB_NAME'], "bg_temp", fname, fpath),
                    daemon=True
                ).start()
                logger.info(f"BG image generated and persisting to MongoDB: {fname}")
            except Exception as e:
                logger.warning(f"Failed to persist BG image in MongoDB (non-fatal): {e}")
            
            return {"imageUrl": f"/api/projects/bg_temp/assets/{fname}", "imageBase64": f"data:image/png;base64,{images[0]['data']}"}
        raise HTTPException(500, "No image generated")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"BG image generation error: {e}")
        raise HTTPException(500, f"Failed to generate background image: {str(e)}")


@router.get("/agent/sessions/by-project/{project_id}")
async def agent_get_session_by_project(project_id: str):
    """Find the agent session associated with a project (for editing media of existing course)."""
    s = await db.agent_sessions.find_one(
        {"projectId": project_id},
        {"_id": 0, "contentText": 0}
    )
    if not s:
        raise HTTPException(404, "No agent session found for this project")
    return s



@router.post("/agent/sessions/{session_id}/media-config")
async def agent_set_media_config(session_id: str, data: dict):
    """Save media and background configuration for each slide."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    media_config = data.get("mediaConfig", {})
    bg_config = data.get("bgConfig", {})
    global_text_color = data.get("globalTextColor", "")
    global_font_size = data.get("globalFontSize", "")
    global_animation = data.get("globalAnimation", "")
    design_template_id = data.get("designTemplateId", "")
    update_data = {
        "mediaConfig": media_config, "bgConfig": bg_config,
        "globalTextColor": global_text_color, "globalFontSize": global_font_size,
        "globalAnimation": global_animation,
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    if design_template_id:
        update_data["config.designTemplateId"] = design_template_id
    await db.agent_sessions.update_one({"id": session_id}, {"$set": update_data})
    return {"status": "ok", "configured": len(media_config), "backgrounds": len(bg_config)}


@router.post("/agent/sessions/{session_id}/apply-media-changes")
async def apply_media_changes(session_id: str, data: dict):
    """Apply media config changes (backgrounds, animations, text color) to an existing project.
    Supports changedSlides parameter to only update specific slides."""
    project_id = data.get("projectId")
    if not project_id:
        raise HTTPException(400, "projectId is required")

    # changedSlides: list of slide indices that were modified, or None for all
    changed_slides = data.get("changedSlides")  # None = apply to all, [] = none, [0,2,5] = specific

    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(404, "Project not found")

    bg_config = s.get("bgConfig") or {}
    global_text_color = s.get("globalTextColor") or ""
    global_font_size = s.get("globalFontSize") or ""
    global_animation = s.get("globalAnimation") or ""
    media_config = s.get("mediaConfig") or {}
    slides = project.get("course", {}).get("slides", [])
    storyboard = s.get("storyboard") or {}
    storyboard_slides = storyboard.get("slides", []) if isinstance(storyboard, dict) else []
    updated_count = 0

    # Design template support - apply visual theme to all slides
    design_template_id = s.get("config", {}).get("designTemplateId", "") or data.get("designTemplateId", "")
    design_token = None
    if design_template_id:
        from services.ai_agent import get_design_template_by_id
        design_token = get_design_template_by_id(design_template_id)

    for i, slide in enumerate(slides):
        # Skip slides that were not changed (if changedSlides is specified)
        if changed_slides is not None and i not in changed_slides:
            continue

        changed = False
        idx_str = str(i)

        # Apply background from bgConfig
        bg = bg_config.get(idx_str, {})
        if bg:
            bg_type = bg.get("type", "")
            if bg_type == "solid":
                slide["background"] = bg.get("color", "#FFFFFF")
                slide.pop("backgroundImage", None)
                changed = True
            elif bg_type == "gradient":
                c1 = bg.get("color1", "#1e293b")
                c2 = bg.get("color2", "#10b981")
                direction = bg.get("direction", "to right")
                slide["background"] = f"linear-gradient({direction}, {c1}, {c2})"
                slide.pop("backgroundImage", None)
                changed = True
            elif bg_type in ("image", "ai_image"):
                bg_url = bg.get("imageUrl", "")
                bg_data = bg.get("imageData", "")
                # If we have base64 imageData but no imageUrl, persist it first
                if bg_data and bg_data.startswith("data:image") and not bg_url:
                    try:
                        import hashlib
                        header, b64 = bg_data.split(",", 1)
                        ext = "png"
                        if "jpeg" in header or "jpg" in header:
                            ext = "jpg"
                        seed = hashlib.md5(b64[:200].encode()).hexdigest()[:12]
                        fname = f"bg_upload_{seed}.{ext}"
                        fpath = os.path.join(str(PROJECTS_DIR), "bg_temp", fname)
                        os.makedirs(os.path.dirname(fpath), exist_ok=True)
                        img_bytes = base64.b64decode(b64)
                        with open(fpath, "wb") as f:
                            f.write(img_bytes)
                        bg_url = f"/api/projects/bg_temp/assets/{fname}"
                        # Persist to MongoDB for production
                        try:
                            from services.asset_store import store_asset_sync
                            import threading
                            threading.Thread(
                                target=store_asset_sync,
                                args=(mongo_url, os.environ['DB_NAME'], "bg_temp", fname, fpath),
                                daemon=True
                            ).start()
                        except Exception as pe:
                            logger.warning(f"Failed to persist uploaded BG to MongoDB: {pe}")
                        logger.info(f"Persisted uploaded BG image: {fname}")
                    except Exception as e:
                        logger.error(f"Failed to save uploaded BG image: {e}")
                if bg_url:
                    slide["backgroundImage"] = bg_url
                    slide["backgroundOpacity"] = bg.get("opacity", 0.3)
                    changed = True

        # Apply global text color
        if global_text_color:
            for el in slide.get("elements", []):
                if el.get("type") in ("html", "text"):
                    if not el.get("style"):
                        el["style"] = {}
                    el["style"]["color"] = global_text_color
                    changed = True

        # Apply global font size scale
        if global_font_size:
            scale = int(global_font_size) / 100.0
            import re as _re
            for el in slide.get("elements", []):
                if el.get("type") in ("html", "text") and el.get("htmlContent"):
                    def _scale_font(m):
                        orig = float(m.group(1))
                        return f"font-size:{round(orig * scale)}px"
                    el["htmlContent"] = _re.sub(r'font-size:\s*(\d+(?:\.\d+)?)px', _scale_font, el["htmlContent"])
                    changed = True

        # Apply global animation
        if global_animation:
            stagger_idx = 0
            for el in slide.get("elements", []):
                if el.get("type") in ("html", "text"):
                    el["animation"] = {
                        "type": "entrance",
                        "effect": global_animation,
                        "duration": 0.5,
                        "delay": stagger_idx * 0.2,
                    }
                    el["animations"] = [{
                        "type": "entrance",
                        "effect": global_animation,
                        "duration": 0.5,
                        "startTime": (el.get("startTime", 0) or 0) + stagger_idx * 0.2,
                    }]
                    stagger_idx += 1
                    changed = True

        # Apply design template styling (headers, fonts, backgrounds, corner radius)
        if design_token:
            _apply_design_token_to_slide(slide, design_token, storyboard_slides[i] if i < len(storyboard_slides) else {})
            changed = True

        # Apply media config (AI image generation, video embeds, etc.)
        mc = media_config.get(idx_str, {})
        media_type = mc.get("type", "")

        if media_type == "ai_image":
            # Generate AI image for this slide
            sb_slide = storyboard_slides[i] if i < len(storyboard_slides) else {}
            kw = sb_slide.get("imageKeywords", slide.get("title", "education"))
            try:
                from services.ai_agent import _fetch_stock_image
                img_url = await _fetch_stock_image(kw, str(PROJECTS_DIR), project_id)
                if img_url:
                    # Find existing image element and update, or add one
                    img_found = False
                    for el in slide.get("elements", []):
                        if el.get("type") == "image":
                            el["src"] = img_url
                            el["content"] = img_url
                            img_found = True
                            break
                    if not img_found:
                        from models import generate_id
                        # Add image element and resize text to two-column layout
                        slide.setdefault("elements", [])
                        slide["elements"].append({
                            "id": generate_id(), "type": "image",
                            "x": 1160, "y": 90, "width": 700, "height": 440,
                            "src": img_url, "content": img_url,
                            "style": {"borderRadius": "12px"}, "startTime": 0,
                            "animations": [{"id": generate_id(), "type": "entrance", "effect": "fade", "trigger": "withPrevious", "duration": 0.5, "delay": 0.3}],
                        })
                        # Resize text elements to left column
                        for el in slide.get("elements", []):
                            if el.get("type") in ("html", "text") and el.get("width", 0) > 1200:
                                el["width"] = 1050
                                el["x"] = 60
                    changed = True
                    logger.info(f"AI image generated for slide {i}: {img_url}")
            except Exception as e:
                logger.error(f"Failed to generate AI image for slide {i}: {e}")

        elif media_type == "gallery_image":
            # Use a pre-existing image from the gallery
            gallery_url = mc.get("galleryImageUrl", "")
            if gallery_url:
                img_found = False
                for el in slide.get("elements", []):
                    if el.get("type") == "image":
                        el["src"] = gallery_url
                        el["content"] = gallery_url
                        img_found = True
                        break
                if not img_found:
                    from models import generate_id
                    slide.setdefault("elements", [])
                    slide["elements"].append({
                        "id": generate_id(), "type": "image",
                        "x": 1160, "y": 90, "width": 700, "height": 440,
                        "src": gallery_url, "content": gallery_url,
                        "style": {"borderRadius": "12px"}, "startTime": 0,
                        "animations": [{"id": generate_id(), "type": "entrance", "effect": "fade", "trigger": "withPrevious", "duration": 0.5, "delay": 0.3}],
                    })
                    for el in slide.get("elements", []):
                        if el.get("type") in ("html", "text") and el.get("width", 0) > 1200:
                            el["width"] = 1050
                            el["x"] = 60
                changed = True
                logger.info(f"Gallery image applied to slide {i}: {gallery_url}")

        elif media_type in ("youtube", "vimeo"):
            video_url = mc.get("url", "")
            if video_url:
                from services.ai_agent import _parse_video_url
                video_info = _parse_video_url(video_url)
                if video_info:
                    from models import generate_id
                    embed_url = video_info["embedUrl"]
                    platform = "YouTube" if video_info["type"] == "youtube" else "Vimeo"
                    video_html = f'''<div style="width:100%;height:100%;border-radius:12px;overflow:hidden;background:#000;position:relative;">
<iframe src="{embed_url}" style="width:100%;height:100%;border:none;" allowfullscreen allow="autoplay; encrypted-media"></iframe>
<div style="position:absolute;bottom:0;left:0;right:0;padding:6px 12px;background:linear-gradient(transparent,rgba(0,0,0,0.7));">
<span style="color:rgba(255,255,255,0.6);font-size:11px;">{platform}</span>
</div>
</div>'''
                    # Replace existing image/video or add new
                    replaced = False
                    for el in slide.get("elements", []):
                        if el.get("type") == "image":
                            el["type"] = "html"
                            el["htmlContent"] = video_html
                            el.pop("src", None)
                            el.pop("content", None)
                            replaced = True
                            break
                    if not replaced:
                        slide.setdefault("elements", [])
                        slide["elements"].append({
                            "id": generate_id(), "type": "html",
                            "x": 1160, "y": 90, "width": 700, "height": 440,
                            "htmlContent": video_html,
                            "style": {}, "startTime": 0,
                            "animations": [{"id": generate_id(), "type": "entrance", "effect": "fade", "trigger": "withPrevious", "duration": 0.5, "delay": 0.3}],
                        })
                        for el in slide.get("elements", []):
                            if el.get("type") in ("html", "text") and el.get("width", 0) > 1200:
                                el["width"] = 1050
                                el["x"] = 60
                    changed = True

        elif media_type == "none":
            # Remove image/video elements and expand text to full width
            slide["elements"] = [el for el in slide.get("elements", []) if el.get("type") != "image"]
            for el in slide.get("elements", []):
                if el.get("type") in ("html", "text") and el.get("x", 0) <= 80:
                    el["width"] = 1760
                    el["x"] = 80
            changed = True

        # Apply narration config from media config
        narr = mc.get("narration", {})
        if narr.get("enabled") and narr.get("selectedScript") and narr.get("voiceId"):
            # Mark for narration generation
            slide.setdefault("_narrationPending", {
                "script": narr["selectedScript"],
                "voiceId": narr["voiceId"],
            })
            changed = True

        if changed:
            updated_count += 1

    # Save updated slides
    await db.projects.update_one(
        {"id": project_id},
        {"$set": {"course.slides": slides, "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )

    # Handle narration generation for new narration configs
    narration_tasks = []
    for i, slide in enumerate(slides):
        pending = slide.pop("_narrationPending", None)
        if pending:
            # Check if this slide already has narration audio
            has_narration = any(a.get("type") == "narration" for a in slide.get("audio", []))
            if not has_narration:
                narration_tasks.append({
                    "slideIndex": i,
                    "slideId": slide.get("id", ""),
                    "slideTitle": slide.get("title", f"Slide {i+1}"),
                    "text": pending["script"],
                    "status": "pending",
                })

    if narration_tasks and ELEVENLABS_API_KEY:
        voice_id = narration_tasks[0].get("voiceId", "") if narration_tasks else ""
        # Get voiceId from first narration config
        for idx_str, mc in media_config.items():
            narr = mc.get("narration", {})
            if narr.get("voiceId"):
                voice_id = narr["voiceId"]
                break
        if voice_id:
            await db.projects.update_one(
                {"id": project_id},
                {"$set": {"narrationPending": narration_tasks, "narrationVoiceId": voice_id}}
            )
            # Generate narrations in background
            import asyncio
            asyncio.create_task(_generate_narrations(project_id, narration_tasks, voice_id))

    return {"status": "ok", "updatedSlides": updated_count, "projectId": project_id}



@router.post("/agent/sessions/{session_id}/generate-slide-narration")
async def agent_generate_slide_narration(session_id: str, data: dict):
    """Generate 3 narration script options for an agent storyboard slide."""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    emergent_key = os.environ.get('EMERGENT_LLM_KEY', '')
    if not emergent_key:
        raise HTTPException(500, "AI key not configured")

    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    storyboard = s.get("storyboard", {})
    slides = storyboard.get("slides", [])
    slide_index = data.get("slideIndex", 0)
    if slide_index >= len(slides):
        raise HTTPException(400, "Invalid slide index")

    slide = slides[slide_index]
    style = data.get("style", "educational")
    language = data.get("language", "português brasileiro")

    # Extract text content from storyboard slide
    text_parts = []
    slide_title = slide.get("title", "")
    if slide_title:
        text_parts.append(f"Título: {slide_title}")
    for el in slide.get("elements", []):
        content = el.get("content", "")
        if content:
            import re as _re
            clean = _re.sub(r'<[^>]+>', '', content).strip()
            if clean:
                text_parts.append(clean)
    narration_script = slide.get("narrationScript", "")
    if narration_script:
        text_parts.append(f"Roteiro sugerido: {narration_script}")

    slide_text = "\n".join(text_parts)

    style_guide = {
        "educational": "educativo e didático, explicando conceitos de forma clara e objetiva",
        "conversational": "conversacional e descontraído, como se estivesse falando com um amigo",
        "formal": "formal e profissional, adequado para ambientes corporativos",
        "friendly": "amigável e acolhedor, criando conexão com o espectador"
    }

    try:
        system_msg = f"""Você é um especialista em criar textos de narração para slides de cursos e apresentações.

Suas diretrizes:
1. Escreva em {language}
2. Use tom {style_guide.get(style, style_guide['educational'])}
3. O texto deve ser adequado para narração em voz alta (TTS)
4. Escreva de forma natural, fluida e envolvente
5. Use pausas naturais (vírgulas, pontos) para dar ritmo
6. O texto deve complementar e explicar o conteúdo do slide
7. Cada opção deve ter entre 2 e 5 frases (ideal para 20-60 segundos de narração)
8. Cada opção deve ter uma abordagem ligeiramente diferente

FORMATO DE RESPOSTA OBRIGATÓRIO:
Retorne exatamente 3 opções, separadas por "---". Cada opção deve conter APENAS o texto de narração, sem numeração, títulos ou marcadores."""

        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"agent-narration-{uuid.uuid4()}",
            system_message=system_msg
        ).with_model("gemini", "gemini-3-flash-preview")

        prompt = f"Crie 3 opções de texto de narração para o seguinte slide:\n\n{slide_text}"
        response = await chat.send_message(UserMessage(text=prompt))

        options = [opt.strip() for opt in response.split("---") if opt.strip()]
        while len(options) < 3:
            options.append(options[-1] if options else "Narração não disponível.")
        options = options[:3]

        return {"options": options, "slideIndex": slide_index, "style": style}
    except Exception as e:
        logger.error(f"Agent narration generation error: {e}")
        raise HTTPException(500, f"Falha ao gerar narração: {str(e)}")



@router.post("/agent/sessions/{session_id}/save-narration-config")
async def agent_save_narration_config(session_id: str, data: dict):
    """Save per-slide narration configuration from the Storyboard screen."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    narration_slides = data.get("narrationSlides", {})
    narration_voice_id = data.get("narrationVoiceId", "")
    narration_enabled = data.get("narrationEnabled", False)

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {
            "config.narrationEnabled": narration_enabled,
            "config.narrationVoiceId": narration_voice_id,
            "config.narrationSlides": narration_slides,
            "updatedAt": datetime.now(timezone.utc).isoformat(),
        }}
    )

    enabled_count = sum(1 for v in narration_slides.values() if v)
    return {"status": "ok", "enabledSlides": enabled_count}


@router.post("/agent/sessions/{session_id}/save-type-overrides")
async def agent_save_type_overrides(session_id: str, data: dict):
    """Save slide type overrides from the Storyboard screen (e.g., avatar_scene -> content)."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    type_overrides = data.get("typeOverrides", {})

    # Apply type overrides to the storyboard slides
    storyboard = s.get("storyboard") or {}
    slides = storyboard.get("slides", [])
    for idx_str, new_type in type_overrides.items():
        idx = int(idx_str)
        if 0 <= idx < len(slides):
            slides[idx]["originalType"] = slides[idx].get("type", "content")
            slides[idx]["type"] = new_type

    # Build update dict - only set storyboard.slides if storyboard exists
    update_dict = {
        "config.slideTypeOverrides": type_overrides,
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    
    # Only update storyboard.slides if there's an existing storyboard
    if s.get("storyboard") is not None:
        update_dict["storyboard.slides"] = slides
    else:
        # Initialize storyboard with empty slides if it doesn't exist
        update_dict["storyboard"] = {"slides": slides}

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": update_dict}
    )

    return {"status": "ok", "overrides": len(type_overrides)}


@router.post("/agent/sessions/{session_id}/cost-estimate")
async def agent_cost_estimate(session_id: str):
    """Estimate the cost of generating the course before committing."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    structure = s.get("structure", {})
    config = s.get("config", {})
    media_config = s.get("mediaConfig", {})

    # Count slides by type
    total_slides = 0
    content_slides = 0
    for mod in structure.get("modules", []):
        for sl in mod.get("slides", []):
            total_slides += 1
            if sl.get("type") == "content":
                content_slides += 1

    # Count AI images needed
    ai_images = sum(1 for v in media_config.values() if v.get("type") == "ai_image")
    if ai_images == 0:
        ai_images = content_slides  # Default: all content slides get AI images

    # Storyboard batches (4 slides per batch)
    storyboard_batches = (total_slides + 3) // 4

    # Cost estimates (approximate, in USD)
    # Gemini 3 Flash: ~$0.001 per 1K input tokens, ~$0.004 per 1K output tokens
    # ~2K input + ~1K output per batch = ~$0.006 per batch
    text_cost_per_batch = 0.006
    # Gemini Nano Banana: ~$0.02 per image
    image_cost_per_image = 0.02

    text_cost = (storyboard_batches + 2) * text_cost_per_batch  # +2 for analysis & structure
    image_cost = ai_images * image_cost_per_image

    # ElevenLabs narration cost based on per-slide config
    narration_cost = 0.0
    narration_slides = config.get("narrationSlides", {})
    if config.get("narrationEnabled"):
        storyboard = s.get("storyboard", {})
        sb_slides = storyboard.get("slides", [])
        # ElevenLabs Starter: $5/30,000 chars = $0.000167/char
        cost_per_char = 5.0 / 30000.0
        for i, sb_slide in enumerate(sb_slides):
            # If narrationSlides configured, use it; otherwise default to all
            if narration_slides:
                if not narration_slides.get(str(i), False):
                    continue
            script = sb_slide.get("narrationScript", "")
            if script:
                narration_cost += len(script) * cost_per_char

    total_cost = text_cost + image_cost + narration_cost

    # Compare with old GPT-5.2 + GPT Image 1 costs
    old_text_cost = (storyboard_batches + 2) * 0.06  # GPT-5.2 ~10x more expensive
    old_image_cost = ai_images * 0.08  # GPT Image 1 ~4x more expensive
    old_total = old_text_cost + old_image_cost + narration_cost
    savings_pct = round((1 - total_cost / old_total) * 100) if old_total > 0 else 0

    return {
        "estimate": {
            "totalSlides": total_slides,
            "contentSlides": content_slides,
            "aiImages": ai_images,
            "storyboardBatches": storyboard_batches,
            "narrationEnabled": config.get("narrationEnabled", False),
            "costs": {
                "text": round(text_cost, 3),
                "images": round(image_cost, 3),
                "narration": round(narration_cost, 3),
                "total": round(total_cost, 3),
            },
            "comparison": {
                "oldTotal": round(old_total, 3),
                "newTotal": round(total_cost, 3),
                "savingsPercent": savings_pct,
            },
            "models": {
                "text": "Gemini 3 Flash",
                "images": "Gemini Nano Banana",
                "narration": "ElevenLabs" if config.get("narrationEnabled") else "N/A",
            }
        }
    }


@router.post("/agent/sessions/{session_id}/generate-course")
async def agent_generate_course(session_id: str):
    """Step 5: Generate actual Scormfy project from storyboard with media (background thread)."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    if not s.get("storyboard"):
        raise HTTPException(400, "Storyboard not generated yet")

    # If already generating, don't start another
    if s.get("step") == "generating_course":
        return {"status": "processing", "message": "Geração do curso já em andamento..."}
    # If already done
    if s.get("step") == "generated" and s.get("projectId"):
        return {"status": "already_done", "projectId": s["projectId"]}

    # Mark as processing
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"step": "generating_course", "courseProgress": {"message": "Iniciando geração..."}, "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )

    def _sync_generate_course():
        import asyncio as _asyncio
        loop = _asyncio.new_event_loop()
        _asyncio.set_event_loop(loop)
        project = None
        try:
            from services.ai_agent import generate_course_from_storyboard
            from motor.motor_asyncio import AsyncIOMotorClient as _MotorClient

            _client = _MotorClient(os.environ.get("MONGO_URL"), serverSelectionTimeoutMS=30000, connectTimeoutMS=30000)
            _db = _client[os.environ.get("DB_NAME")]

            _s = loop.run_until_complete(_db.agent_sessions.find_one({"id": session_id}, {"_id": 0}))
            config = _s.get("config", {})
            media_config = _s.get("mediaConfig", {})
            bg_config = _s.get("bgConfig", {})
            global_text_color = _s.get("globalTextColor", "")
            global_font_size = _s.get("globalFontSize", "")
            global_animation = _s.get("globalAnimation", "")
            title = config.get("title", _s.get("analysis", {}).get("title", "Curso Gerado por IA"))
            desc = config.get("description", _s.get("analysis", {}).get("summary", ""))

            from models import Project
            project = Project(name=title, description=desc)
            project_dir = PROJECTS_DIR / project.id
            (project_dir / "assets").mkdir(parents=True, exist_ok=True)

            # Save project early with "generating" status so images aren't lost on failure
            project.course.metadata.title = title
            project.course.metadata.description = desc
            project_dict = project.model_dump()
            project_dict["createdAt"] = project.createdAt.isoformat()
            project_dict["updatedAt"] = project.updatedAt.isoformat()
            project_dict["course"]["createdAt"] = project.course.createdAt.isoformat()
            project_dict["course"]["updatedAt"] = project.course.updatedAt.isoformat()
            project_dict["course"]["slides"] = []
            project_dict["createdByAgent"] = True
            project_dict["agentSessionId"] = session_id
            project_dict["status"] = "generating"
            project_dict["userId"] = _s.get("userId")
            project_dict["companyId"] = _s.get("companyId")
            loop.run_until_complete(_db.projects.insert_one(project_dict))
            logger.info(f"Early project save: {project.id} (status=generating)")

            # Update progress
            loop.run_until_complete(_db.agent_sessions.update_one(
                {"id": session_id},
                {"$set": {"courseProgress": {"message": "Gerando slides com IA..."}, "projectId": project.id, "updatedAt": datetime.now(timezone.utc).isoformat()}}
            ))

            course_data = loop.run_until_complete(generate_course_from_storyboard(
                session_id, _s["storyboard"], config,
                project_dir=str(PROJECTS_DIR), project_id=project.id,
                media_config=media_config, bg_config=bg_config,
                global_text_color=global_text_color, global_font_size=global_font_size, global_animation=global_animation,
                design_template_id=config.get("designTemplateId", "")
            ))

            loop.run_until_complete(_db.agent_sessions.update_one(
                {"id": session_id},
                {"$set": {"courseProgress": {"message": "Salvando projeto..."}, "updatedAt": datetime.now(timezone.utc).isoformat()}}
            ))

            heygen_pending = course_data.get("heygenPending", [])
            
            # Update project with final course data (project was already created early)
            loop.run_until_complete(_db.projects.update_one(
                {"id": project.id},
                {"$set": {
                    "course.slides": course_data["slides"],
                    "course.updatedAt": datetime.now(timezone.utc).isoformat(),
                    "status": "draft",
                    "updatedAt": datetime.now(timezone.utc).isoformat(),
                    "heygenPending": heygen_pending if heygen_pending else []
                }}
            ))

            # Save quiz questions
            for q in course_data.get("quizQuestions", []):
                q["projectId"] = project.id
                q["createdAt"] = datetime.now(timezone.utc).isoformat()
                q["updatedAt"] = datetime.now(timezone.utc).isoformat()
                loop.run_until_complete(_db.questions.insert_one({**q, "_id": q["id"]}))

            # Handle narration (global toggle with per-slide control)
            narration_enabled = config.get("narrationEnabled", False)
            narration_voice = config.get("narrationVoiceId", "")
            narration_slides_map = config.get("narrationSlides", {})
            narration_count = 0
            if narration_enabled and narration_voice and ELEVENLABS_API_KEY:
                import re as _re
                narration_tasks = []
                for si, slide_data in enumerate(course_data["slides"]):
                    # Check per-slide toggle: if map exists, only narrate enabled slides
                    if narration_slides_map:
                        if not narration_slides_map.get(str(si), False):
                            continue
                    texts = []
                    for el in slide_data.get("elements", []):
                        html_content = el.get("htmlContent", "")
                        if html_content:
                            clean = _re.sub(r'<[^>]+>', '', html_content)
                            clean = _re.sub(r'\s+', ' ', clean).strip()
                            if len(clean) > 30:
                                texts.append(clean)
                    if texts:
                        narration_tasks.append({
                            "slideIndex": si,
                            "slideId": slide_data.get("id", ""),
                            "slideTitle": slide_data.get("title", f"Slide {si+1}"),
                            "text": " ".join(texts)[:2000],
                            "status": "pending",
                        })
                if narration_tasks:
                    narration_count = len(narration_tasks)
                    loop.run_until_complete(_db.projects.update_one(
                        {"id": project.id},
                        {"$set": {"narrationPending": narration_tasks, "narrationVoiceId": narration_voice}}
                    ))

            # Handle per-slide narration from media config
            per_slide_narration = course_data.get("narrationPending", [])
            if per_slide_narration and ELEVENLABS_API_KEY:
                per_slide_tasks = []
                for nr in per_slide_narration:
                    per_slide_tasks.append({
                        "slideIndex": nr["slideIndex"],
                        "slideId": nr["slideId"],
                        "slideTitle": f"Slide {nr['slideIndex'] + 1}",
                        "text": nr["script"],
                        "status": "pending",
                    })
                if per_slide_tasks:
                    narration_count += len(per_slide_tasks)
                    existing = loop.run_until_complete(_db.projects.find_one({"id": project.id}, {"_id": 0, "narrationPending": 1}))
                    existing_tasks = existing.get("narrationPending", []) if existing else []
                    all_tasks = existing_tasks + per_slide_tasks
                    per_slide_voice = per_slide_narration[0]["voiceId"]
                    loop.run_until_complete(_db.projects.update_one(
                        {"id": project.id},
                        {"$set": {"narrationPending": all_tasks, "narrationVoiceId": per_slide_voice}}
                    ))

            # Update session with result
            result_data = {
                "projectId": project.id,
                "projectName": title,
                "slidesCount": len(course_data["slides"]),
                "quizCount": len(course_data.get("quizQuestions", [])),
                "heygenPending": len(heygen_pending),
                "narrationPending": narration_count,
            }
            loop.run_until_complete(_db.agent_sessions.update_one(
                {"id": session_id},
                {"$set": {
                    "projectId": project.id,
                    "step": "generated",
                    "courseProgress": None,
                    "courseResult": result_data,
                    "updatedAt": datetime.now(timezone.utc).isoformat()
                }}
            ))

            # Trigger background tasks (HeyGen, narration, suggestions) in their own threads
            if heygen_pending and HEYGEN_API_KEY:
                import threading
                threading.Thread(target=lambda: asyncio.run(_trigger_heygen_videos(project.id, heygen_pending)), daemon=True).start()
            if narration_count > 0:
                all_narr = []
                if narration_enabled and narration_voice and ELEVENLABS_API_KEY:
                    all_narr = narration_tasks if 'narration_tasks' in dir() else []
                if per_slide_narration and ELEVENLABS_API_KEY:
                    all_narr += per_slide_tasks if 'per_slide_tasks' in dir() else []
                if all_narr:
                    voice = narration_voice or (per_slide_narration[0]["voiceId"] if per_slide_narration else "")
                    import threading
                    threading.Thread(target=lambda: asyncio.run(_generate_narrations(project.id, all_narr, voice)), daemon=True).start()

            # Suggestions in background
            import threading
            threading.Thread(target=lambda: asyncio.run(_generate_improvement_suggestions(session_id, project.id)), daemon=True).start()

            # Log usage for reporting
            try:
                slides_count = len(course_data.get("slides", []))
                ai_images_count = sum(1 for s in course_data.get("slides", []) for e in s.get("elements", []) if e.get("type") == "image" and "ai_img" in e.get("src", ""))
                usage_log = {
                    "id": str(uuid.uuid4()),
                    "type": "course_generation",
                    "userId": _s.get("userId"),
                    "companyId": _s.get("companyId"),
                    "projectId": project.id,
                    "sessionId": session_id,
                    "details": {
                        "slides": slides_count,
                        "aiImages": ai_images_count,
                        "narrations": narration_count + len(per_slide_narration) if per_slide_narration else narration_count,
                        "heygenVideos": len(heygen_pending) if heygen_pending else 0,
                    },
                    "estimatedCost": {
                        "textGeneration": round((len(_s.get("storyboard", {}).get("slides", [])) / 5 + 2) * 0.006, 4),
                        "imageGeneration": round(ai_images_count * 0.02, 4),
                        "narration": round(narration_count * 0.01, 4) if narration_count else 0,
                        "currency": "USD"
                    },
                    "createdAt": datetime.now(timezone.utc).isoformat()
                }
                loop.run_until_complete(_db.usage_logs.insert_one(usage_log))
            except Exception as e:
                logger.warning(f"Failed to log usage (non-fatal): {e}")

            _client.close()
            logger.info(f"Course generation completed for session {session_id}, project {project.id}")

        except Exception as e:
            err_msg = str(e)
            logger.error(f"Course generation error: {err_msg}")
            error_detail = "Erro ao gerar curso."
            if "Budget" in err_msg or "budget" in err_msg:
                error_detail = "Orçamento da chave LLM excedido. Acesse Profile > Universal Key > Add Balance para adicionar saldo."
            elif "502" in err_msg or "BadGateway" in err_msg:
                error_detail = "Serviço de IA temporariamente indisponível (502). Tente novamente em alguns minutos."
            try:
                from motor.motor_asyncio import AsyncIOMotorClient as _MotorClient2
                _c2 = _MotorClient2(os.environ.get("MONGO_URL"), serverSelectionTimeoutMS=10000)
                _d2 = _c2[os.environ.get("DB_NAME")]
                loop.run_until_complete(_d2.agent_sessions.update_one(
                    {"id": session_id},
                    {"$set": {"step": "media_configured", "error": error_detail, "courseProgress": None, "updatedAt": datetime.now(timezone.utc).isoformat()}}
                ))
                # Mark early-saved project as failed (but keep images)
                if project:
                    loop.run_until_complete(_d2.projects.update_one(
                        {"id": project.id},
                        {"$set": {"status": "error", "error": error_detail, "updatedAt": datetime.now(timezone.utc).isoformat()}}
                    ))
                    logger.info(f"Project {project.id} marked as error - images preserved")
                _c2.close()
            except Exception:
                pass
        finally:
            loop.close()

    import threading
    thread = threading.Thread(target=_sync_generate_course, daemon=True)
    thread.start()
    return {"status": "processing", "message": "Geração do curso iniciada..."}


@router.get("/agent/sessions/{session_id}/course-status")
async def agent_course_status(session_id: str):
    """Poll for course generation status."""
    s = await db.agent_sessions.find_one(
        {"id": session_id},
        {"_id": 0, "step": 1, "courseProgress": 1, "courseResult": 1, "error": 1, "projectId": 1}
    )
    if not s:
        raise HTTPException(404, "Session not found")

    if s.get("step") == "generated" and (s.get("courseResult") or s.get("projectId")):
        result = s.get("courseResult", {"projectId": s.get("projectId"), "status": "done"})
        return {"status": "done", **result}
    elif s.get("step") == "generating_course":
        progress = s.get("courseProgress", {})
        return {"status": "processing", "message": progress.get("message", "Gerando...")}
    elif s.get("error"):
        return {"status": "error", "error": s["error"]}
    else:
        return {"status": "unknown", "step": s.get("step")}



async def _generate_improvement_suggestions(session_id: str, project_id: str):
    """Background task: analyze the course creation process and generate improvement suggestions."""
    try:
        from motor.motor_asyncio import AsyncIOMotorClient as _MotorClient
        _client = _MotorClient(os.environ.get("MONGO_URL"), serverSelectionTimeoutMS=30000, connectTimeoutMS=30000)
        _db = _client[os.environ.get("DB_NAME")]

        session = await _db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
        if not session:
            _client.close()
            return

        # Build analysis context
        config = session.get("config", {})
        analysis = session.get("analysis", {})
        storyboard = session.get("storyboard", {})
        media_config = session.get("mediaConfig", {})
        bg_config = session.get("bgConfig", {})
        slides_count = len(storyboard.get("slides", []))
        modules_count = len(storyboard.get("modules", []))

        # Collect slide types
        slide_types = {}
        for s in storyboard.get("slides", []):
            t = s.get("type", "unknown")
            slide_types[t] = slide_types.get(t, 0) + 1

        # Media usage
        media_types = {}
        for v in media_config.values():
            t = v.get("type", "none")
            media_types[t] = media_types.get(t, 0) + 1

        # Custom backgrounds
        custom_bg_count = sum(1 for v in bg_config.values() if v.get("type", "default") != "default")

        context = f"""
ANÁLISE DO PROCESSO DE CRIAÇÃO DO CURSO:
- Título: {config.get('title', 'N/A')}
- Template: {config.get('template', 'N/A')}
- Tom: {config.get('tone', 'N/A')}
- Idioma: {config.get('language', 'pt-BR')}
- Módulos: {modules_count}
- Total de slides: {slides_count}
- Tipos de slides: {slide_types}
- Mídias configuradas: {media_types}
- Fundos personalizados: {custom_bg_count}
- Narração habilitada: {config.get('narrationEnabled', False)}
- HeyGen habilitado: {any(v.get('type') == 'heygen' for v in media_config.values())}
- Conteúdo fonte: {analysis.get('contentType', 'texto')}
- Palavras no conteúdo fonte: {analysis.get('wordCount', 'N/A')}
- Tópicos identificados: {', '.join(analysis.get('topics', [])[:5])}
- Nível de complexidade: {analysis.get('complexity', 'N/A')}

ESTRUTURA DOS MÓDULOS:
{chr(10).join(f"- {m.get('title', '?')}: {m.get('slidesCount', '?')} slides" for m in storyboard.get('modules', [])[:10])}
"""

        prompt = f"""Você é um consultor especialista em design instrucional e plataformas de e-learning.
Analise o seguinte processo de criação de curso e gere sugestões de melhoria detalhadas.

{context}

Gere sugestões organizadas nas seguintes categorias. Para cada sugestão, inclua:
- "title": título curto (max 60 chars)
- "description": descrição detalhada (2-3 frases)
- "priority": "alta", "media" ou "baixa"
- "impact": breve descrição do impacto esperado

Responda APENAS em JSON válido com esta estrutura exata:
{{
  "platform_ux": [sugestões para melhorar a experiência do usuário na plataforma Scormfy],
  "platform_features": [sugestões de novas funcionalidades para a plataforma],
  "platform_performance": [sugestões para melhorar performance e confiabilidade],
  "course_content": [sugestões para melhorar a qualidade do conteúdo do curso gerado],
  "course_design": [sugestões para melhorar o design visual do curso],
  "course_pedagogy": [sugestões para melhorar a metodologia pedagógica do curso],
  "course_interactivity": [sugestões de simuladores, jogos educativos e elementos interativos HTML+JS para adicionar ao curso. Tipos de jogos: Forca Educacional (adivinhação de termos), Bate-pênalti com perguntas (quiz esportivo), Jogo de acerto ao alvo (certo x errado), Jogo da memória educativa (cartas pergunta+resposta), Quiz gamificado com barra de energia, Calculadoras temáticas, Flashcards, Drag-and-drop, Linha do tempo interativa. Foco: fixação de conteúdo, engajamento emocional, repetição ativa, feedback imediato]
}}

Gere 2-3 sugestões por categoria, totalizando 14-21 sugestões. Seja específico e actionable."""

        from emergentintegrations.llm.chat import LlmChat, UserMessage
        emergent_key = os.environ.get("EMERGENT_LLM_KEY", "")

        async def _try_llm(provider, model_name):
            chat = LlmChat(
                api_key=emergent_key,
                session_id=f"suggestions_{session_id}",
                system_message="Você é um consultor especialista em design instrucional e plataformas de e-learning. Responda sempre em JSON válido.",
            ).with_model(provider, model_name)
            return await chat.send_message(UserMessage(text=prompt))

        try:
            raw = await _try_llm("gemini", "gemini-3-flash-preview")
        except Exception as llm_err:
            logger.warning(f"Suggestions: gemini-3-flash failed ({str(llm_err)[:60]}), falling back to gpt-4o")
            raw = await _try_llm("openai", "gpt-4o")

        import json
        import re
        # Extract JSON from possible markdown blocks
        json_match = re.search(r"```json\s*([\s\S]*?)```", raw)
        if json_match:
            suggestions = json.loads(json_match.group(1).strip())
        else:
            suggestions = json.loads(raw.strip())

        # Store suggestions
        await _db.agent_sessions.update_one(
            {"id": session_id},
            {"$set": {
                "suggestions": suggestions,
                "suggestionsGeneratedAt": datetime.now(timezone.utc).isoformat(),
                "updatedAt": datetime.now(timezone.utc).isoformat(),
            }}
        )
        logger.info(f"Improvement suggestions generated for session {session_id}")

        _client.close()
    except Exception as e:
        logger.error(f"Failed to generate suggestions for session {session_id}: {e}")
        try:
            await _db.agent_sessions.update_one(
                {"id": session_id},
                {"$set": {"suggestionsError": str(e), "updatedAt": datetime.now(timezone.utc).isoformat()}}
            )
            _client.close()
        except Exception:
            pass


@router.get("/agent/sessions/{session_id}/suggestions")
async def agent_get_suggestions(session_id: str):
    """Get improvement suggestions for a session."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0, "suggestions": 1, "suggestionsGeneratedAt": 1, "suggestionsError": 1})
    if not s:
        raise HTTPException(404, "Session not found")
    if s.get("suggestions"):
        return {"status": "ready", "suggestions": s["suggestions"], "generatedAt": s.get("suggestionsGeneratedAt")}
    if s.get("suggestionsError"):
        return {"status": "error", "error": s["suggestionsError"]}
    return {"status": "pending"}


@router.post("/agent/sessions/{session_id}/suggestions/regenerate")
async def agent_regenerate_suggestions(session_id: str, background_tasks: BackgroundTasks):
    """Regenerate improvement suggestions on demand."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0, "projectId": 1, "step": 1})
    if not s:
        raise HTTPException(404, "Session not found")
    if s.get("step") != "generated":
        raise HTTPException(400, "Course not generated yet")
    # Clear old suggestions
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$unset": {"suggestions": "", "suggestionsError": ""}, "$set": {"updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    background_tasks.add_task(_generate_improvement_suggestions, session_id, s.get("projectId", ""))
    return {"status": "regenerating"}


async def _trigger_heygen_videos(project_id: str, pending_list: list):
    """Background task: trigger HeyGen video generation for each pending slide."""
    for item in pending_list:
        try:
            used_transparent = False
            async with httpx.AsyncClient(timeout=60.0) as http_client:
                response = None

                # Strategy 1: Try transparent WebM first (v1/video.webm)
                webm_payload = {
                    "avatar_pose_id": item["avatar_id"],
                    "avatar_style": "normal",
                    "input_text": item["script"],
                    "voice_id": item["voice_id"],
                    "dimension": {"width": 1280, "height": 720},
                }
                response = await http_client.post(
                    f"{HEYGEN_BASE_URL}/v1/video.webm",
                    headers=HEYGEN_HEADERS,
                    json=webm_payload,
                )
                if response.status_code == 200:
                    used_transparent = True
                    logger.info(f"HeyGen WebM transparent succeeded for slide {item['slideId']}")
                else:
                    logger.warning(f"WebM transparent failed for slide {item['slideId']} ({response.status_code}): {response.text[:200]}")
                    response = None

                # Strategy 2: Fallback to standard v2 endpoint
                if response is None:
                    fallback_payload = {
                        "video_inputs": [{
                            "character": {
                                "type": "avatar",
                                "avatar_id": item["avatar_id"],
                                "avatar_style": "normal"
                            },
                            "voice": {
                                "type": "text",
                                "input_text": item["script"],
                                "voice_id": item["voice_id"]
                            }
                        }],
                        "dimension": {"width": 1280, "height": 720},
                        "title": f"Agent-{project_id}-{item['slideId']}"
                    }
                    response = await http_client.post(
                        f"{HEYGEN_BASE_URL}/v2/video/generate",
                        headers=HEYGEN_HEADERS,
                        json=fallback_payload,
                    )

                if response.status_code == 200:
                    data = response.json()
                    video_id = data.get("data", {}).get("video_id")
                    if video_id:
                        await db.heygen_videos.insert_one({
                            "video_id": video_id,
                            "avatar_id": item["avatar_id"],
                            "voice_id": item["voice_id"],
                            "script": item["script"],
                            "title": item["title"],
                            "status": "processing",
                            "transparent": used_transparent,
                            "project_id": project_id,
                            "slide_id": item["slideId"],
                            "created_at": now_utc(),
                        })
                        await db.projects.update_one(
                            {"id": project_id, "heygenPending.slideId": item["slideId"]},
                            {"$set": {"heygenPending.$.videoId": video_id, "heygenPending.$.status": "processing"}}
                        )
                        logger.info(f"HeyGen video triggered: {video_id} (transparent={used_transparent}) for slide {item['slideId']}")
                else:
                    logger.error(f"HeyGen video generation failed: {response.status_code} - {response.text}")
                    await db.projects.update_one(
                        {"id": project_id, "heygenPending.slideId": item["slideId"]},
                        {"$set": {"heygenPending.$.status": "failed", "heygenPending.$.error": response.text[:200]}}
                    )
        except Exception as e:
            logger.error(f"HeyGen video trigger error: {e}")
            await db.projects.update_one(
                {"id": project_id, "heygenPending.slideId": item["slideId"]},
                {"$set": {"heygenPending.$.status": "failed", "heygenPending.$.error": str(e)[:200]}}
            )


@router.get("/agent/projects/{project_id}/heygen-status")
async def agent_heygen_status(project_id: str):
    """Check HeyGen video generation status for a project."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0, "id": 1, "heygenPending": 1})
    if not project:
        raise HTTPException(404, "Project not found")
    pending = project.get("heygenPending", [])
    if not pending:
        return {"status": "no_heygen", "videos": []}

    results = []
    all_done = True
    for item in pending:
        video_id = item.get("videoId")
        status = item.get("status", "pending")
        if video_id and status == "processing":
            # Check HeyGen API for status
            try:
                async with httpx.AsyncClient(timeout=30.0) as http_client:
                    resp = await http_client.get(
                        f"{HEYGEN_BASE_URL}/v1/video_status.get?video_id={video_id}",
                        headers=HEYGEN_HEADERS
                    )
                    if resp.status_code == 200:
                        vdata = resp.json().get("data", {})
                        status = vdata.get("status", "processing")
                        video_url = vdata.get("video_url")

                        # Update project pending status
                        update_fields = {"heygenPending.$.status": status}
                        if video_url:
                            update_fields["heygenPending.$.videoUrl"] = video_url
                        await db.projects.update_one(
                            {"id": project_id, "heygenPending.videoId": video_id},
                            {"$set": update_fields}
                        )

                        # If completed, update the slide element
                        if status == "completed" and video_url:
                            await _update_slide_with_heygen_video(project_id, item["slideId"], video_url)
                            await db.heygen_videos.update_one(
                                {"video_id": video_id},
                                {"$set": {"status": "completed", "video_url": video_url}}
                            )
            except Exception as e:
                logger.error(f"HeyGen status check error: {e}")

        if status not in ("completed", "failed"):
            all_done = False

        results.append({
            "slideId": item.get("slideId"),
            "slideIndex": item.get("slideIndex"),
            "title": item.get("title"),
            "videoId": item.get("videoId"),
            "status": status,
            "videoUrl": item.get("videoUrl"),
        })

    return {"status": "all_done" if all_done else "processing", "videos": results}


async def _update_slide_with_heygen_video(project_id: str, slide_id: str, video_url: str):
    """Replace the HeyGen processing element with actual video player."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        return

    slides = project.get("course", {}).get("slides", [])
    for si, slide in enumerate(slides):
        if slide.get("id") == slide_id:
            # Find and replace the heygen processing element
            for ei, el in enumerate(slide.get("elements", [])):
                html = el.get("htmlContent", "")
                if f'data-heygen-slide="{slide_id}"' in html:
                    from models import generate_id
                    new_html = f'''<div style="width:100%;height:100%;border-radius:12px;overflow:hidden;background:#000;position:relative;">
<video src="{video_url}" style="width:100%;height:100%;object-fit:cover;" controls autoplay muted></video>
<div style="position:absolute;bottom:0;left:0;right:0;padding:6px 12px;background:linear-gradient(transparent,rgba(0,0,0,0.7));">
<span style="color:rgba(255,255,255,0.6);font-size:11px;">Avatar HeyGen</span>
</div>
</div>'''
                    await db.projects.update_one(
                        {"id": project_id},
                        {"$set": {f"course.slides.{si}.elements.{ei}.htmlContent": new_html}}
                    )
                    logger.info(f"Updated slide {slide_id} with HeyGen video")
                    return


@router.post("/agent/projects/{project_id}/generate-narration")
async def agent_generate_narration(project_id: str, background_tasks: BackgroundTasks):
    """Trigger narration generation for all content slides of a project."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(404, "Project not found")

    # Get narration config from the agent session
    session_id = project.get("agentSessionId", "")
    session = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0, "config": 1}) if session_id else None
    config = session.get("config", {}) if session else {}
    voice_id = config.get("narrationVoiceId", "")

    if not voice_id:
        raise HTTPException(400, "No narration voice configured")
    if not ELEVENLABS_API_KEY:
        raise HTTPException(400, "ElevenLabs API key not configured")

    # Collect slides that need narration
    slides = project.get("course", {}).get("slides", [])
    narration_tasks = []
    for si, slide in enumerate(slides):
        # Skip slides that already have audio
        if slide.get("audio") and len(slide["audio"]) > 0:
            continue
        # Extract text from HTML elements
        texts = []
        for el in slide.get("elements", []):
            html_content = el.get("htmlContent", "")
            if html_content:
                import re as _re
                clean = _re.sub(r'<[^>]+>', '', html_content)
                clean = _re.sub(r'\s+', ' ', clean).strip()
                if len(clean) > 30:
                    texts.append(clean)
        if texts:
            narration_tasks.append({
                "slideIndex": si,
                "slideId": slide.get("id", ""),
                "slideTitle": slide.get("title", f"Slide {si+1}"),
                "text": " ".join(texts)[:2000],
                "status": "pending",
            })

    if not narration_tasks:
        return {"status": "no_slides", "message": "No slides need narration"}

    # Store narration tasks in the project
    await db.projects.update_one(
        {"id": project_id},
        {"$set": {"narrationPending": narration_tasks, "narrationVoiceId": voice_id}}
    )

    # Trigger background generation
    background_tasks.add_task(_generate_narrations, project_id, narration_tasks, voice_id)

    return {"status": "ok", "slides": len(narration_tasks), "voiceId": voice_id}


async def _generate_narrations(project_id: str, tasks: list, voice_id: str):
    """Background task: generate ElevenLabs narration for each slide."""
    try:
        import aiofiles as _aiofiles
        from elevenlabs import ElevenLabs
        from elevenlabs.types import VoiceSettings
        from motor.motor_asyncio import AsyncIOMotorClient as _MotorClient

        # Create own DB connection for this thread's event loop
        _client = _MotorClient(os.environ.get("MONGO_URL"), serverSelectionTimeoutMS=30000, connectTimeoutMS=30000)
        _db = _client[os.environ.get("DB_NAME")]

        el_client = ElevenLabs(api_key=ELEVENLABS_API_KEY)
        voice_settings = VoiceSettings(stability=0.5, similarity_boost=0.75, style=0.0, use_speaker_boost=True)

        for task in tasks:
            try:
                # Generate TTS
                audio_generator = el_client.text_to_speech.convert(
                    text=task["text"][:1500],
                    voice_id=voice_id,
                    model_id="eleven_multilingual_v2",
                    voice_settings=voice_settings
                )
                audio_data = b""
                for chunk in audio_generator:
                    audio_data += chunk

                # Save audio file
                audio_id = str(uuid.uuid4())
                filename = f"narration_{project_id}_{task['slideIndex']}.mp3"
                audio_path = STORAGE_DIR / "audio" / filename
                audio_path.parent.mkdir(exist_ok=True)

                async with _aiofiles.open(audio_path, "wb") as f:
                    await f.write(audio_data)

                # Persist audio data to MongoDB for production durability
                await _db.tts_generations.update_one(
                    {"filename": filename},
                    {"$set": {
                        "filename": filename,
                        "audio_data": base64.b64encode(audio_data).decode(),
                        "file_size": len(audio_data),
                        "project_id": project_id,
                        "type": "narration",
                        "created_at": now_utc().isoformat()
                    }},
                    upsert=True
                )

                # Add audio to the slide
                slide_audio = {
                    "id": audio_id,
                    "url": f"/api/audio/{filename}",
                    "filename": filename,
                    "type": "narration",
                    "volume": 1.0,
                    "startTime": 0,
                    "duration": len(audio_data) / 16000,  # Approximate
                }
                await _db.projects.update_one(
                    {"id": project_id},
                    {"$push": {f"course.slides.{task['slideIndex']}.audio": slide_audio}}
                )

                # Update status
                await _db.projects.update_one(
                    {"id": project_id, "narrationPending.slideId": task["slideId"]},
                    {"$set": {"narrationPending.$.status": "completed", "narrationPending.$.audioUrl": f"/api/audio/{filename}"}}
                )
                logger.info(f"Narration generated for slide {task['slideIndex']} of project {project_id}")

            except Exception as e:
                logger.error(f"Narration error for slide {task['slideIndex']}: {e}")
                await _db.projects.update_one(
                    {"id": project_id, "narrationPending.slideId": task["slideId"]},
                    {"$set": {"narrationPending.$.status": "failed", "narrationPending.$.error": str(e)[:200]}}
                )

        _client.close()
    except Exception as e:
        logger.error(f"Narration generation failed: {e}")


@router.get("/agent/projects/{project_id}/narration-status")
async def agent_narration_status(project_id: str):
    """Check narration generation status for a project."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0, "id": 1, "narrationPending": 1})
    if not project:
        raise HTTPException(404, "Project not found")
    pending = project.get("narrationPending", [])
    if not pending:
        return {"status": "no_narration", "slides": []}

    completed = sum(1 for t in pending if t.get("status") == "completed")
    failed = sum(1 for t in pending if t.get("status") == "failed")
    total = len(pending)
    all_done = (completed + failed) == total

    return {
        "status": "all_done" if all_done else "processing",
        "total": total,
        "completed": completed,
        "failed": failed,
        "slides": [
            {"slideIndex": t.get("slideIndex"), "title": t.get("slideTitle"), "status": t.get("status", "pending"), "audioUrl": t.get("audioUrl")}
            for t in pending
        ],
    }

@router.post("/agent/sessions/{session_id}/chat")
async def agent_chat_endpoint(session_id: str, data: AgentChatMessage):
    """Chat with the agent for adjustments."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    from services.ai_agent import agent_chat
    context = {"structure": s.get("structure"), "config": s.get("config"), "analysis": s.get("analysis")}
    response = await agent_chat(session_id, data.message, context)

    # Save to chat history
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$push": {"chatHistory": {"role": "user", "text": data.message, "ts": datetime.now(timezone.utc).isoformat()}}}
    )
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$push": {"chatHistory": {"role": "agent", "text": response, "ts": datetime.now(timezone.utc).isoformat()}}}
    )

    return {"response": response}


# --- Agent: Templates ---

@router.get("/agent/templates")
async def agent_list_templates():
    """List available course templates."""
    from services.ai_agent import get_templates
    return get_templates()


@router.get("/agent/design-templates")
async def agent_list_design_templates():
    """List available visual/design templates for course styling."""
    from services.ai_agent import get_design_templates
    return get_design_templates()



@router.post("/agent/sessions/{session_id}/generate-structure-from-template")
async def agent_generate_structure_from_template(session_id: str, data: dict):
    """Generate course structure using a template as base."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    template_id = data.get("templateId")
    if not template_id:
        raise HTTPException(400, "templateId is required")

    from services.ai_agent import generate_structure_from_template
    structure = await generate_structure_from_template(session_id, s["contentText"], s.get("config", {}), template_id)
    if not structure:
        raise HTTPException(500, "Failed to generate structure from template")

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"structure": structure, "step": "structured", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return structure


# --- Agent: Course Editing ---

@router.get("/agent/courses")
async def agent_list_courses():
    """List all courses available for agent analysis (agent-created + imported)."""
    projects = await db.projects.find(
        {}, {"_id": 0}
    ).sort("createdAt", -1).to_list(500)
    result = []
    for p in projects:
        slides = p.get("course", {}).get("slides", [])
        is_agent = bool(p.get("createdByAgent")) or bool(p.get("agentSessionId"))
        result.append({
            "id": p.get("id"),
            "name": p.get("name", ""),
            "description": p.get("description", ""),
            "createdAt": p.get("createdAt"),
            "updatedAt": p.get("updatedAt"),
            "agentSessionId": p.get("agentSessionId"),
            "createdByAgent": p.get("createdByAgent"),
            "slidesCount": len(slides),
            "source": "agent" if is_agent else "imported",
        })
    return result


class AgentImprovementsApply(BaseModel):
    improvements: list
    selectedNewSlides: Optional[list] = None
    previewId: Optional[str] = None


def _build_improved_elements(ai_elements: list, generate_id_fn) -> list:
    """Build properly positioned HTML elements from AI response."""
    new_html_elements = []
    current_y = 80
    for elem in ai_elements:
        if not isinstance(elem, dict):
            continue
        
        elem_type = elem.get("type", "text")
        elem_height = elem.get("height", 400)
        
        # Handle HTML/simulator elements (full HTML documents)
        if elem_type == "html" and elem.get("htmlContent"):
            html_content = elem["htmlContent"]
            if isinstance(html_content, str) and ("<!DOCTYPE" in html_content or "<html" in html_content or "<script" in html_content):
                el = {
                    "id": generate_id_fn(),
                    "type": "html",
                    "x": 0,
                    "y": 0,
                    "width": elem.get("width", 960),
                    "height": elem.get("height", 540),
                    "content": "",
                    "htmlContent": html_content,
                    "style": {},
                    "startTime": 0,
                    "animations": [],
                }
                new_html_elements.append(el)
                continue
        
        # Standard text/content elements
        raw_content = elem.get("htmlContent") or elem.get("content", "")
        if isinstance(raw_content, dict):
            raw_content = raw_content.get("html", "") or raw_content.get("text", "") or json.dumps(raw_content, ensure_ascii=False)
        elif isinstance(raw_content, list):
            raw_content = " ".join(str(c) for c in raw_content)
        elif not isinstance(raw_content, str):
            raw_content = str(raw_content) if raw_content else ""

        el = {
            "id": generate_id_fn(),
            "type": "html",
            "x": 80,
            "y": current_y,
            "width": elem.get("width", 1760),
            "height": elem_height,
            "content": "",
            "htmlContent": raw_content,
            "style": {"fontSize": 18, "fontFamily": "Inter, sans-serif", "fontColor": "#333333"},
            "startTime": 0,
            "animations": [],
        }
        new_html_elements.append(el)
        current_y += elem_height + 20
    return new_html_elements


def _apply_ai_result_to_slides(slides: list, result: dict, generate_id_fn) -> int:
    """Apply AI improvement result to slides in-place. Returns count of new slides added."""
    import copy
    for upd in result.get("updatedSlides", []):
        if not isinstance(upd, dict):
            continue
        idx = upd.get("slideIndex")
        if idx is not None and 0 <= idx < len(slides):
            title = upd.get("title")
            if title and isinstance(title, str):
                slides[idx]["title"] = title
            notes = upd.get("notes")
            if notes and isinstance(notes, str):
                slides[idx]["notes"] = notes
            narration = upd.get("narrationScript")
            if narration and isinstance(narration, str):
                slides[idx]["librasScript"] = narration
            libras = upd.get("librasScript")
            if libras and isinstance(libras, str):
                slides[idx]["librasScript"] = libras
            if upd.get("elements") and isinstance(upd["elements"], list):
                existing_elements = slides[idx].get("elements", [])
                preserved = [e for e in existing_elements if e.get("type") not in ("html", "text")]
                header = [e for e in existing_elements if e.get("type") == "html" and e.get("y", 0) == 0 and e.get("height", 0) <= 60]
                new_html = _build_improved_elements(upd["elements"], generate_id_fn)
                slides[idx]["elements"] = header + new_html + preserved

    new_slides_added = 0
    for ns in sorted(result.get("newSlides", []), key=lambda x: x.get("afterIndex", 999), reverse=True):
        if not isinstance(ns, dict):
            continue
        after_idx = ns.get("afterIndex", len(slides) - 1)
        insert_at = min(after_idx + 1, len(slides))
        new_elements = _build_improved_elements(ns.get("elements", []) if isinstance(ns.get("elements"), list) else [], generate_id_fn)
        ns_title = ns.get("title", "Novo Slide")
        if not isinstance(ns_title, str):
            ns_title = str(ns_title)
        new_slide = {
            "id": generate_id_fn(),
            "title": ns_title,
            "order": insert_at,
            "width": 1920,
            "height": 820,
            "background": ns.get("background", "#FFFFFF"),
            "elements": new_elements,
            "annotations": [],
            "transition": {"type": "fade", "duration": 0.5},
            "audio": [],
            "notes": "",
            "librasScript": ns.get("librasScript", ""),
            "duration": 5.0,
        }
        slides.insert(insert_at, new_slide)
        new_slides_added += 1

    for i, s in enumerate(slides):
        s["order"] = i

    return new_slides_added


@router.post("/agent/courses/{project_id}/analyze")
async def agent_analyze_course(project_id: str):
    """Analyze an existing agent-created course and suggest improvements."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(404, "Project not found")

    session_id = project.get("agentSessionId") or str(uuid.uuid4())
    from services.ai_agent import analyze_existing_course
    analysis = await analyze_existing_course(session_id, project)
    return analysis


@router.post("/agent/courses/{project_id}/preview-improvements")
async def agent_preview_improvements(project_id: str, data: AgentImprovementsApply):
    """Preview improvements without saving. Returns before/after slide data."""
    import copy
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(404, "Project not found")

    session_id = project.get("agentSessionId") or str(uuid.uuid4())
    from services.ai_agent import apply_course_improvements
    from models import generate_id

    result = await apply_course_improvements(session_id, project, data.improvements, data.selectedNewSlides)

    original_slides = project.get("course", {}).get("slides", [])

    # Build preview of "after" slides without modifying originals
    after_slides = copy.deepcopy(original_slides)
    _apply_ai_result_to_slides(after_slides, result, generate_id)

    # Build before/after comparison for affected slides only
    affected_indices = set()
    for upd in result.get("updatedSlides", []):
        idx = upd.get("slideIndex")
        if idx is not None and 0 <= idx < len(original_slides):
            affected_indices.add(idx)

    def _strip_html(html_str: str) -> str:
        """Strip HTML tags and return plain text."""
        if not html_str:
            return ""
        clean = re.sub(r'<[^>]+>', ' ', html_str)
        clean = re.sub(r'&nbsp;', ' ', clean)
        clean = re.sub(r'&amp;', '&', clean)
        clean = re.sub(r'&lt;', '<', clean)
        clean = re.sub(r'&gt;', '>', clean)
        clean = re.sub(r'&#\d+;', '', clean)
        clean = re.sub(r'\s+', ' ', clean).strip()
        return clean

    def _extract_text(slide):
        texts = []
        for el in slide.get("elements", []):
            if el.get("type") not in ("html", "text"):
                continue
            c = el.get("htmlContent") or el.get("content") or ""
            if c and isinstance(c, str):
                plain = _strip_html(c)
                if plain:
                    texts.append(plain)
        return texts

    comparisons = []
    for idx in sorted(affected_indices):
        comparisons.append({
            "slideIndex": idx,
            "title": {
                "before": original_slides[idx].get("title", ""),
                "after": after_slides[idx].get("title", ""),
            },
            "contentBefore": _extract_text(original_slides[idx]),
            "contentAfter": _extract_text(after_slides[idx]),
        })

    # New slides preview
    new_slide_previews = []
    for ns in result.get("newSlides", []):
        new_slide_previews.append({
            "afterIndex": ns.get("afterIndex", 0),
            "title": ns.get("title", "Novo Slide"),
            "content": [_strip_html(e.get("content", "")) for e in ns.get("elements", [])],
        })

    # Cache the AI result for later apply
    preview_id = str(uuid.uuid4())
    await db.improvement_previews.insert_one({
        "id": preview_id,
        "projectId": project_id,
        "aiResult": result,
        "createdAt": datetime.now(timezone.utc).isoformat(),
    })

    return {
        "previewId": preview_id,
        "comparisons": comparisons,
        "newSlides": new_slide_previews,
        "updatedCount": len(comparisons),
        "newCount": len(new_slide_previews),
    }


@router.post("/agent/courses/{project_id}/apply-improvements")
async def agent_apply_improvements(project_id: str, data: AgentImprovementsApply):
    """Apply selected improvements to an existing course. Uses cached preview if previewId provided."""
    import copy
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(404, "Project not found")

    from models import generate_id

    # Save snapshot for undo
    original_course = copy.deepcopy(project.get("course", {}))
    await db.course_snapshots.update_one(
        {"projectId": project_id},
        {"$set": {
            "projectId": project_id,
            "course": original_course,
            "createdAt": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True,
    )

    # Use cached AI result if preview was done, otherwise call AI fresh
    if data.previewId:
        preview = await db.improvement_previews.find_one({"id": data.previewId, "projectId": project_id}, {"_id": 0})
        if preview:
            result = preview["aiResult"]
            await db.improvement_previews.delete_one({"id": data.previewId})
        else:
            raise HTTPException(400, "Preview expired or not found. Please preview again.")
    else:
        session_id = project.get("agentSessionId") or str(uuid.uuid4())
        from services.ai_agent import apply_course_improvements
        result = await apply_course_improvements(session_id, project, data.improvements, data.selectedNewSlides)

    slides = project.get("course", {}).get("slides", [])
    new_slides_added = _apply_ai_result_to_slides(slides, result, generate_id)

    # Save
    course = project.get("course", {})
    course["slides"] = slides
    await update_project(project_id, {"course": course})

    # Detect avatar scenes and trigger background generation
    avatar_scene_pending = []
    for ns in result.get("newSlides", []):
        avatar_scene = ns.get("avatarScene")
        if avatar_scene and isinstance(avatar_scene, dict):
            # Find the slide that was just added
            ns_title = ns.get("title", "")
            target_slide_id = None
            for s in slides:
                if s.get("title") == ns_title:
                    target_slide_id = s.get("id")
                    break
            if not target_slide_id:
                # Try to find by afterIndex
                after_idx = ns.get("afterIndex", 0)
                insert_at = min(after_idx + 1, len(slides) - 1)
                if 0 <= insert_at < len(slides):
                    target_slide_id = slides[insert_at].get("id")

            if target_slide_id:
                avatar_settings = project.get("avatarSceneSettings", {})
                avatar_scene_pending.append({
                    "slideId": target_slide_id,
                    "slideTitle": ns.get("title", "Avatar Scene"),
                    "narrationScript": avatar_scene.get("narrationScript", ""),
                    "backgroundPrompt": avatar_scene.get("backgroundPrompt", ""),
                    "avatarPosition": avatar_scene.get("avatarPosition", "left"),
                    "avatarId": avatar_settings.get("defaultAvatarId", ""),
                    "voiceId": avatar_settings.get("defaultVoiceId", ""),
                    "bgStatus": "pending",
                    "audioStatus": "pending",
                    "heygenStatus": "pending",
                    "createdAt": datetime.now(timezone.utc).isoformat(),
                })

    # Also check updatedSlides for avatar scenes
    for upd in result.get("updatedSlides", []):
        avatar_scene = upd.get("avatarScene")
        if avatar_scene and isinstance(avatar_scene, dict):
            idx = upd.get("slideIndex")
            if idx is not None and 0 <= idx < len(slides):
                target_slide_id = slides[idx].get("id")
                if target_slide_id:
                    avatar_settings = project.get("avatarSceneSettings", {})
                    avatar_scene_pending.append({
                        "slideId": target_slide_id,
                        "slideTitle": upd.get("title", f"Slide {idx + 1}"),
                        "narrationScript": avatar_scene.get("narrationScript", ""),
                        "backgroundPrompt": avatar_scene.get("backgroundPrompt", ""),
                        "avatarPosition": avatar_scene.get("avatarPosition", "left"),
                        "avatarId": avatar_settings.get("defaultAvatarId", ""),
                        "voiceId": avatar_settings.get("defaultVoiceId", ""),
                        "bgStatus": "pending",
                        "audioStatus": "pending",
                        "heygenStatus": "pending",
                        "createdAt": datetime.now(timezone.utc).isoformat(),
                    })

    avatar_generation_started = False
    if avatar_scene_pending:
        await db.projects.update_one(
            {"id": project_id},
            {"$set": {"avatarScenePending": avatar_scene_pending}}
        )
        # Trigger background generation
        import threading
        threading.Thread(
            target=lambda: asyncio.run(_trigger_avatar_scene_generation(project_id, avatar_scene_pending)),
            daemon=True,
        ).start()
        avatar_generation_started = True
        logger.info(f"Avatar scene generation triggered for {len(avatar_scene_pending)} scenes in project {project_id}")

    return {
        "status": "ok",
        "updatedSlides": len(result.get("updatedSlides", [])),
        "newSlides": new_slides_added,
        "totalSlides": len(slides),
        "canUndo": True,
        "avatarScenesTriggered": len(avatar_scene_pending),
        "avatarGenerationStarted": avatar_generation_started,
    }


@router.post("/agent/courses/{project_id}/undo-improvements")
async def agent_undo_improvements(project_id: str):
    """Undo the last applied improvements by restoring the saved snapshot."""
    snapshot = await db.course_snapshots.find_one({"projectId": project_id}, {"_id": 0})
    if not snapshot:
        raise HTTPException(404, "Nenhum snapshot encontrado para desfazer.")

    await update_project(project_id, {"course": snapshot["course"]})
    await db.course_snapshots.delete_one({"projectId": project_id})

    slides = snapshot["course"].get("slides", [])
    return {
        "status": "ok",
        "message": "Melhorias desfeitas com sucesso.",
        "totalSlides": len(slides),
    }


# =============================================================================
# AVATAR SCENE SETTINGS & GENERATION
# =============================================================================

class AvatarSceneSettings(BaseModel):
    maxScenes: int = 3
    defaultAvatarId: Optional[str] = None
    defaultVoiceId: Optional[str] = None


@router.get("/agent/projects/{project_id}/avatar-settings")
async def get_avatar_settings(project_id: str):
    """Get avatar scene settings for a project."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0, "id": 1, "avatarSceneSettings": 1})
    if not project:
        raise HTTPException(404, "Project not found")
    settings = project.get("avatarSceneSettings", {"maxScenes": 3, "defaultAvatarId": None, "defaultVoiceId": None})
    return settings


@router.put("/agent/projects/{project_id}/avatar-settings")
async def update_avatar_settings(project_id: str, settings: AvatarSceneSettings):
    """Update avatar scene settings for a project."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0, "id": 1})
    if not project:
        raise HTTPException(404, "Project not found")
    await db.projects.update_one(
        {"id": project_id},
        {"$set": {
            "avatarSceneSettings": settings.model_dump(),
            "updatedAt": datetime.now(timezone.utc).isoformat(),
        }}
    )
    return {"status": "ok", "settings": settings.model_dump()}


@router.get("/agent/projects/{project_id}/avatar-generation-status")
async def get_avatar_generation_status(project_id: str):
    """Get status of avatar scene generation (background image and HeyGen video)."""
    project = await db.projects.find_one(
        {"id": project_id},
        {"_id": 0, "avatarScenePending": 1, "id": 1}
    )
    if not project:
        raise HTTPException(404, "Project not found")
    pending = project.get("avatarScenePending", [])
    if not pending:
        return {"status": "no_pending", "scenes": []}

    # Check HeyGen video status for any pending items
    for scene in pending:
        if scene.get("heygenVideoId") and scene.get("heygenStatus") == "processing":
            try:
                async with httpx.AsyncClient(timeout=30.0) as http_client:
                    resp = await http_client.get(
                        f"{HEYGEN_BASE_URL}/v1/video_status.get?video_id={scene['heygenVideoId']}",
                        headers=HEYGEN_HEADERS,
                    )
                    if resp.status_code == 200:
                        vdata = resp.json().get("data", {})
                        status = vdata.get("status", "")
                        if status == "completed":
                            video_url = vdata.get("video_url", "")
                            scene["heygenStatus"] = "completed"
                            scene["heygenVideoUrl"] = video_url
                            # Update the slide with the video element
                            await _apply_heygen_video_to_slide(
                                project_id, scene["slideId"], video_url,
                                scene.get("avatarPosition", "left")
                            )
                        elif status == "failed":
                            scene["heygenStatus"] = "failed"
                            scene["heygenError"] = vdata.get("error", {}).get("message", "Unknown error")
            except Exception as e:
                logger.warning(f"HeyGen status check failed: {e}")

    # Persist updated status
    await db.projects.update_one(
        {"id": project_id},
        {"$set": {"avatarScenePending": pending}}
    )

    all_done = all(
        scene.get("bgStatus") in ("completed", "failed", "skipped") and
        scene.get("audioStatus") in ("completed", "failed", "skipped") and
        scene.get("heygenStatus") in ("completed", "failed", "skipped", None, "")
        for scene in pending
    )

    return {
        "status": "completed" if all_done else "processing",
        "scenes": [{k: v for k, v in s.items() if k != "_id"} for s in pending],
    }


async def _apply_heygen_video_to_slide(project_id: str, slide_id: str, video_url: str, position: str):
    """Add HeyGen video element to a slide after generation completes."""
    from models import generate_id
    project = await db.projects.find_one({"id": project_id}, {"_id": 0, "course.slides": 1})
    if not project:
        return
    slides = project.get("course", {}).get("slides", [])
    for slide in slides:
        if slide.get("id") == slide_id:
            # Position the video element based on avatarPosition
            if position == "right":
                vx, vy, vw, vh = 1320, 120, 520, 520
            elif position == "center":
                vx, vy, vw, vh = 700, 100, 520, 520
            else:  # left
                vx, vy, vw, vh = 80, 120, 520, 520

            video_el = {
                "id": generate_id(),
                "type": "video",
                "x": vx, "y": vy,
                "width": vw, "height": vh,
                "src": video_url,
                "autoplay": True,
                "loop": False,
                "style": {},
                "startTime": 0,
                "animations": [],
            }
            slide.setdefault("elements", []).append(video_el)
            break

    await db.projects.update_one(
        {"id": project_id},
        {"$set": {"course.slides": slides, "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    logger.info(f"HeyGen video applied to slide {slide_id} in project {project_id}")


async def _trigger_avatar_scene_generation(project_id: str, scenes: list):
    """Background task: generate background images and HeyGen avatar videos (with native TTS) for avatar scenes."""
    from motor.motor_asyncio import AsyncIOMotorClient as _MotorClient
    _client = _MotorClient(os.environ.get("MONGO_URL"), serverSelectionTimeoutMS=30000)
    _db = _client[os.environ.get("DB_NAME")]

    for scene in scenes:
        slide_id = scene.get("slideId", "")
        try:
            # 1. Generate background image via Gemini Nano Banana
            bg_prompt = scene.get("backgroundPrompt", "")
            bg_url = None
            if bg_prompt:
                try:
                    from emergentintegrations.llm.chat import LlmChat, UserMessage
                    emergent_key = os.environ.get("EMERGENT_LLM_KEY", "")
                    chat = LlmChat(
                        api_key=emergent_key,
                        session_id=f"avatar_bg_{uuid.uuid4().hex[:8]}",
                        system_message="You are an image generator.",
                    ).with_model("gemini", "gemini-3-pro-image-preview").with_params(modalities=["image", "text"])

                    full_prompt = f"Professional photorealistic background for educational video scene: {bg_prompt}. No text, no watermarks, high quality, cinematic lighting, suitable as video call background. 16:9 aspect ratio."
                    text_resp, images = await chat.send_message_multimodal_response(UserMessage(text=full_prompt))
                    if images and len(images) > 0:
                        img_bytes = base64.b64decode(images[0]['data'])
                        import hashlib
                        seed = hashlib.md5(bg_prompt.encode()).hexdigest()[:10]
                        fname = f"avatar_bg_{seed}.png"
                        fpath = os.path.join(str(PROJECTS_DIR), project_id, "assets", fname)
                        os.makedirs(os.path.dirname(fpath), exist_ok=True)
                        with open(fpath, "wb") as f:
                            f.write(img_bytes)
                        bg_url = f"/api/projects/{project_id}/assets/{fname}"

                        # Persist in MongoDB
                        try:
                            from services.asset_store import store_asset_sync
                            import threading
                            threading.Thread(
                                target=store_asset_sync,
                                args=(os.environ.get("MONGO_URL"), os.environ["DB_NAME"], project_id, fname, fpath),
                                daemon=True,
                            ).start()
                        except Exception:
                            pass

                        # Apply background to slide
                        await _db.projects.update_one(
                            {"id": project_id, "course.slides.id": slide_id},
                            {"$set": {"course.slides.$.backgroundImage": bg_url}}
                        )
                        scene["bgStatus"] = "completed"
                        scene["bgUrl"] = bg_url
                        logger.info(f"Avatar BG generated for slide {slide_id}: {bg_url}")
                    else:
                        scene["bgStatus"] = "failed"
                        scene["bgError"] = "No image generated"
                except Exception as e:
                    logger.error(f"Avatar BG generation failed for slide {slide_id}: {e}")
                    scene["bgStatus"] = "failed"
                    scene["bgError"] = str(e)[:200]
            else:
                scene["bgStatus"] = "skipped"

            # 2. Trigger HeyGen video generation (using HeyGen's native TTS, not ElevenLabs)
            narration_script = scene.get("narrationScript", "")
            avatar_id = scene.get("avatarId", "")
            if not avatar_id:
                proj = await _db.projects.find_one({"id": project_id}, {"_id": 0, "avatarSceneSettings": 1, "id": 1})
                avatar_id = (proj or {}).get("avatarSceneSettings", {}).get("defaultAvatarId", "")

            if narration_script and avatar_id and HEYGEN_API_KEY:
                scene["audioStatus"] = "skipped"  # No separate audio step — HeyGen handles TTS
                try:
                    # Get the HeyGen voice_id from scene or project settings
                    heygen_voice_id = scene.get("voiceId", "")
                    if not heygen_voice_id:
                        proj = await _db.projects.find_one({"id": project_id}, {"_id": 0, "avatarSceneSettings": 1})
                        heygen_voice_id = (proj or {}).get("avatarSceneSettings", {}).get("defaultVoiceId", "")

                    # If still no voice, fetch a PT-BR voice from HeyGen API
                    if not heygen_voice_id:
                        try:
                            async with httpx.AsyncClient(timeout=15.0) as vc:
                                vresp = await vc.get(
                                    f"{HEYGEN_BASE_URL}/v2/voices",
                                    headers=HEYGEN_HEADERS,
                                )
                                if vresp.status_code == 200:
                                    all_voices = vresp.json().get("data", {}).get("voices", [])
                                    # Priority: pt-BR > any Portuguese > first available
                                    pt_br_voices = [v for v in all_voices if "brazil" in v.get("language", "").lower() or "brasil" in v.get("language", "").lower()]
                                    pt_voices = [v for v in all_voices if "portuguese" in v.get("language", "").lower() or "português" in v.get("language", "").lower()]
                                    if pt_br_voices:
                                        heygen_voice_id = pt_br_voices[0].get("voice_id", "")
                                        logger.info(f"Auto-selected PT-BR HeyGen voice: {pt_br_voices[0].get('name', '')} ({heygen_voice_id})")
                                    elif pt_voices:
                                        heygen_voice_id = pt_voices[0].get("voice_id", "")
                                        logger.info(f"Auto-selected Portuguese HeyGen voice: {pt_voices[0].get('name', '')} ({heygen_voice_id})")
                                    elif all_voices:
                                        heygen_voice_id = all_voices[0].get("voice_id", "")
                                        logger.warning(f"No PT-BR voice found, using first available: {all_voices[0].get('name', '')} ({heygen_voice_id})")
                        except Exception as ve:
                            logger.warning(f"Failed to fetch HeyGen voices for fallback: {ve}")

                    if not heygen_voice_id:
                        scene["heygenStatus"] = "failed"
                        scene["heygenError"] = "Nenhuma voz HeyGen válida encontrada"
                        logger.error("No valid HeyGen voice found for avatar scene")
                        # Update pending status before continuing
                        await _db.projects.update_one(
                            {"id": project_id, "avatarScenePending.slideId": slide_id},
                            {"$set": {
                                "avatarScenePending.$.bgStatus": scene.get("bgStatus", "pending"),
                                "avatarScenePending.$.bgUrl": scene.get("bgUrl"),
                                "avatarScenePending.$.audioStatus": scene.get("audioStatus", "skipped"),
                                "avatarScenePending.$.heygenStatus": "failed",
                                "avatarScenePending.$.heygenError": scene.get("heygenError"),
                            }}
                        )
                        continue

                    logger.info(f"HeyGen using native TTS: voice_id={heygen_voice_id}")
                    used_transparent = False

                    async with httpx.AsyncClient(timeout=60.0) as http_client:
                        response = None

                        # Strategy 1: Try transparent WebM (v1/video.webm) — same as manual creation
                        webm_payload = {
                            "avatar_pose_id": avatar_id,
                            "avatar_style": "normal",
                            "input_text": narration_script[:5000],
                            "voice_id": heygen_voice_id,
                            "dimension": {"width": 1280, "height": 720},
                        }
                        logger.info(f"HeyGen trying WebM transparent: avatar={avatar_id}, voice_id={heygen_voice_id}")
                        try:
                            response = await http_client.post(
                                f"{HEYGEN_BASE_URL}/v1/video.webm",
                                headers=HEYGEN_HEADERS,
                                json=webm_payload,
                            )
                            if response.status_code == 200:
                                used_transparent = True
                                logger.info("HeyGen WebM transparent succeeded")
                            else:
                                logger.warning(f"WebM transparent failed ({response.status_code}): {response.text[:200]}")
                                response = None
                        except Exception as webm_err:
                            logger.warning(f"WebM request error: {webm_err}")
                            response = None

                        # Strategy 2: Fallback to standard v2 (same as manual creation flow)
                        if response is None:
                            fallback_payload = {
                                "video_inputs": [{
                                    "character": {
                                        "type": "avatar",
                                        "avatar_id": avatar_id,
                                        "avatar_style": "normal"
                                    },
                                    "voice": {
                                        "type": "text",
                                        "input_text": narration_script[:5000],
                                        "voice_id": heygen_voice_id,
                                    },
                                }],
                                "dimension": {"width": 1280, "height": 720},
                                "title": f"AvatarScene-{project_id[:8]}-{slide_id[:8]}"
                            }
                            logger.info(f"HeyGen fallback v2: avatar={avatar_id}, voice_id={heygen_voice_id}")
                            response = await http_client.post(
                                f"{HEYGEN_BASE_URL}/v2/video/generate",
                                headers=HEYGEN_HEADERS,
                                json=fallback_payload,
                            )

                        if response.status_code == 200:
                            data = response.json()
                            video_id = data.get("data", {}).get("video_id")
                            if video_id:
                                await _db.heygen_videos.insert_one({
                                    "video_id": video_id,
                                    "avatar_id": avatar_id,
                                    "voice_id": heygen_voice_id,
                                    "script": narration_script[:500],
                                    "title": f"AvatarScene-{slide_id[:8]}",
                                    "status": "processing",
                                    "transparent": used_transparent,
                                    "project_id": project_id,
                                    "slide_id": slide_id,
                                    "created_at": datetime.now(timezone.utc),
                                })
                                scene["heygenVideoId"] = video_id
                                scene["heygenStatus"] = "processing"
                                scene["transparent"] = used_transparent
                                logger.info(f"HeyGen avatar video triggered: {video_id} (transparent={used_transparent}) for slide {slide_id}")
                            else:
                                scene["heygenStatus"] = "failed"
                                scene["heygenError"] = "No video_id in response"
                        else:
                            resp_text = response.text[:300]
                            scene["heygenStatus"] = "failed"
                            scene["heygenError"] = resp_text
                            logger.error(f"HeyGen avatar generation failed ({response.status_code}): {resp_text}")
                except Exception as e:
                    scene["heygenStatus"] = "failed"
                    scene["heygenError"] = str(e)[:200]
                    logger.error(f"HeyGen avatar trigger error: {e}")
            else:
                scene["audioStatus"] = "skipped"
                scene["heygenStatus"] = "skipped" if not narration_script else "skipped_no_avatar"

            # Update pending status
            await _db.projects.update_one(
                {"id": project_id, "avatarScenePending.slideId": slide_id},
                {"$set": {
                    "avatarScenePending.$.bgStatus": scene.get("bgStatus", "pending"),
                    "avatarScenePending.$.bgUrl": scene.get("bgUrl"),
                    "avatarScenePending.$.audioStatus": scene.get("audioStatus", "pending"),
                    "avatarScenePending.$.audioUrl": scene.get("audioUrl"),
                    "avatarScenePending.$.heygenVideoId": scene.get("heygenVideoId"),
                    "avatarScenePending.$.heygenStatus": scene.get("heygenStatus", "pending"),
                }}
            )

        except Exception as e:
            logger.error(f"Avatar scene generation error for slide {slide_id}: {e}")

    _client.close()



@router.post("/agent/sessions/{session_id}/add-scenario")
async def add_scenario_to_project(session_id: str, data: dict):
    """Add a scenario slide to an existing project via agent chat request."""
    session = await db.agent_sessions.find_one({"id": session_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    project_id = data.get("projectId") or session.get("projectId")
    if not project_id:
        raise HTTPException(status_code=400, detail="No project associated")

    theme = data.get("theme", "")
    objectives = data.get("objectives", "")
    audience = data.get("audience", "")

    # Use course config as defaults
    config = session.get("config", {})
    if not theme:
        theme = config.get("title", "Cenário interativo")
    if not objectives:
        objectives = config.get("objectives", "Praticar tomada de decisão")
    if not audience:
        audience = config.get("audience", "")

    # Generate scenario via AI
    from services.scenario_service import generate_scenario_with_ai
    try:
        scenario_data = await generate_scenario_with_ai({
            "theme": theme,
            "objectives": objectives,
            "audience": audience,
            "complexity": "intermediate",
            "industry": "",
            "duration_minutes": 10,
            "language": "pt-BR",
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Falha na geração do cenário: {str(e)}")

    # Save scenario to DB
    from models import generate_id
    from datetime import datetime, timezone
    scenario_id = generate_id()
    now_str = datetime.now(timezone.utc).isoformat()

    scenario_doc = {
        "id": scenario_id,
        "project_id": project_id,
        "title": scenario_data.get("title", "Cenário"),
        "description": scenario_data.get("description", ""),
        "context": scenario_data.get("context", ""),
        "characters": scenario_data.get("characters", []),
        "learning_objectives": scenario_data.get("learning_objectives", []),
        "competencies_evaluated": scenario_data.get("competencies_evaluated", []),
        "nodes": scenario_data.get("nodes", []),
        "start_node_id": scenario_data["nodes"][0]["id"] if scenario_data.get("nodes") else None,
        "config": {"theme": theme, "objectives": objectives, "audience": audience},
        "created_at": now_str,
        "updated_at": now_str,
    }
    await db.scenarios.insert_one(scenario_doc)
    scenario_doc.pop("_id", None)

    # Create a new slide with the scenario
    project = await db.projects.find_one({"id": project_id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    slides = project.get("course", {}).get("slides", [])
    new_slide_id = generate_id()

    # Get background from last slide or use default
    last_bg = "#0f172a"
    if slides:
        last_bg = slides[-1].get("background", "#0f172a")

    new_slide = {
        "id": new_slide_id,
        "background": last_bg,
        "elements": [
            {
                "id": generate_id(), "type": "scenario",
                "x": 0, "y": 0, "width": 1920, "height": 820,
                "content": scenario_id,
                "scenarioData": scenario_doc,
                "style": {}, "startTime": 0, "animations": [],
            }
        ],
        "notes": f"Cenário interativo: {scenario_doc['title']}",
        "duration": 5.0,
    }

    slides.append(new_slide)
    await db.projects.update_one(
        {"id": project_id},
        {"$set": {"course.slides": slides, "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )

    return {
        "success": True,
        "message": f"Cenário '{scenario_doc['title']}' adicionado ao curso!",
        "scenario": scenario_doc,
        "slideId": new_slide_id,
        "slideIndex": len(slides) - 1,
    }
