"""
Test new features for iteration 49:
1. Enhanced content input (URL scraping, DOCX native extraction)
2. ElevenLabs narration config + generation
3. HeyGen preview functionality

Using existing project ID: cebb110f-ced1-4e62-8478-7fb6bd99943d
"""

import pytest
import requests
import os
from docx import Document
from io import BytesIO

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')

@pytest.fixture
def api_client():
    """Shared requests session"""
    session = requests.Session()
    session.headers.update({"Content-Type": "application/json"})
    return session


class TestHealthAndBasics:
    """Basic API health and session tests"""

    def test_health_endpoint(self, api_client):
        """Test health endpoint"""
        response = api_client.get(f"{BASE_URL}/api/health")
        assert response.status_code == 200
        data = response.json()
        assert data.get("status") == "healthy"
        print(f"PASS: Health endpoint returns {data}")

    def test_create_agent_session(self, api_client):
        """Test session creation"""
        response = api_client.post(f"{BASE_URL}/api/agent/sessions", json={})
        assert response.status_code == 200
        data = response.json()
        assert "id" in data
        assert data.get("step") == "created"
        print(f"PASS: Created session {data['id']}")
        return data["id"]

    def test_get_templates(self, api_client):
        """Test templates endpoint returns 6 templates"""
        response = api_client.get(f"{BASE_URL}/api/agent/templates")
        assert response.status_code == 200
        data = response.json()
        assert isinstance(data, list)
        assert len(data) == 6, f"Expected 6 templates, got {len(data)}"
        print(f"PASS: Templates endpoint returns {len(data)} templates")
        for t in data:
            print(f"  - {t.get('id')}: {t.get('name')}")


class TestURLContentExtraction:
    """Test URL content extraction feature"""

    def test_upload_url_content(self, api_client):
        """Test URL scraping with python.org/about"""
        # First create a session
        resp = api_client.post(f"{BASE_URL}/api/agent/sessions", json={})
        assert resp.status_code == 200
        session_id = resp.json()["id"]
        print(f"Created session: {session_id}")

        # Submit URL for content extraction
        url = "https://www.python.org/about/"
        resp = api_client.post(
            f"{BASE_URL}/api/agent/sessions/{session_id}/upload",
            data={"url": url},
            headers={"Content-Type": "application/x-www-form-urlencoded"}
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data.get("status") == "ok"
        assert "contentLength" in data
        assert data["contentLength"] > 100, "Expected substantial content from URL"
        assert data.get("fileName") == url
        print(f"PASS: URL content extracted - {data['contentLength']} chars from {url}")

    def test_upload_invalid_url(self, api_client):
        """Test URL scraping with invalid URL"""
        resp = api_client.post(f"{BASE_URL}/api/agent/sessions", json={})
        session_id = resp.json()["id"]

        # Submit invalid URL
        resp = api_client.post(
            f"{BASE_URL}/api/agent/sessions/{session_id}/upload",
            data={"url": "https://invalid-url-that-does-not-exist-12345.com/"},
            headers={"Content-Type": "application/x-www-form-urlencoded"}
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data.get("status") == "ok"
        # Content should indicate an error accessing the URL
        print(f"PASS: Invalid URL handled - content length: {data.get('contentLength')}")


class TestDOCXContentExtraction:
    """Test DOCX native content extraction"""

    def test_upload_docx_file(self, api_client):
        """Test DOCX native extraction"""
        # First create a session
        resp = api_client.post(f"{BASE_URL}/api/agent/sessions", json={})
        assert resp.status_code == 200
        session_id = resp.json()["id"]
        print(f"Created session: {session_id}")

        # Create a test DOCX file in memory
        doc = Document()
        doc.add_heading("Test Document for Scormfy", level=1)
        doc.add_paragraph("This is a test paragraph with some content for the course builder.")
        doc.add_paragraph("Second paragraph with more educational content about Python programming.")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Topic"
        table.rows[0].cells[1].text = "Description"
        table.rows[1].cells[0].text = "Python"
        table.rows[1].cells[1].text = "A programming language"
        
        # Save to bytes
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Upload the DOCX file
        files = {"file": ("test_document.docx", docx_buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        resp = requests.post(
            f"{BASE_URL}/api/agent/sessions/{session_id}/upload",
            files=files
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data.get("status") == "ok"
        assert "contentLength" in data
        assert data["contentLength"] > 50, "Expected content from DOCX"
        assert data.get("fileName") == "test_document.docx"
        print(f"PASS: DOCX extracted - {data['contentLength']} chars from {data['fileName']}")


class TestElevenLabsVoices:
    """Test ElevenLabs voice listing"""

    def test_list_elevenlabs_voices(self, api_client):
        """Test GET /api/elevenlabs/voices"""
        response = api_client.get(f"{BASE_URL}/api/elevenlabs/voices")
        assert response.status_code == 200
        data = response.json()
        assert "voices" in data
        assert isinstance(data["voices"], list)
        assert len(data["voices"]) > 0, "Expected at least one voice"
        
        # Check voice structure
        voice = data["voices"][0]
        assert "voice_id" in voice
        assert "name" in voice
        print(f"PASS: ElevenLabs voices - {len(data['voices'])} voices available")
        
        # Sample voices
        for v in data["voices"][:5]:
            print(f"  - {v['name']} ({v.get('gender', 'N/A')}) - ID: {v['voice_id']}")

    def test_filter_elevenlabs_voices_by_gender(self, api_client):
        """Test filtering ElevenLabs voices by gender"""
        response = api_client.get(f"{BASE_URL}/api/elevenlabs/voices?gender=female")
        assert response.status_code == 200
        data = response.json()
        assert "voices" in data
        # All returned voices should be female
        for v in data["voices"]:
            if v.get("gender"):
                assert v["gender"] == "female", f"Expected female, got {v['gender']}"
        print(f"PASS: Filtered to {len(data['voices'])} female voices")


class TestNarrationStatus:
    """Test narration status and generation endpoints"""
    
    EXISTING_PROJECT_ID = "cebb110f-ced1-4e62-8478-7fb6bd99943d"

    def test_narration_status_existing_project(self, api_client):
        """Test narration status for existing project"""
        response = api_client.get(f"{BASE_URL}/api/agent/projects/{self.EXISTING_PROJECT_ID}/narration-status")
        assert response.status_code == 200
        data = response.json()
        assert "status" in data
        # Status can be 'no_narration', 'processing', or 'all_done'
        print(f"PASS: Narration status for existing project: {data}")

    def test_narration_status_invalid_project(self, api_client):
        """Test narration status for non-existent project"""
        response = api_client.get(f"{BASE_URL}/api/agent/projects/invalid-project-id/narration-status")
        assert response.status_code == 404
        print("PASS: Narration status returns 404 for invalid project")

    def test_generate_narration_no_voice_configured(self, api_client):
        """Test narration generation without configured voice"""
        # Create a new session first
        resp = api_client.post(f"{BASE_URL}/api/agent/sessions", json={})
        session_id = resp.json()["id"]
        
        # Try to generate narration for the existing project (may not have voice configured)
        response = api_client.post(f"{BASE_URL}/api/agent/projects/{self.EXISTING_PROJECT_ID}/generate-narration")
        # Either 400 (no voice) or 200 (voice configured)
        assert response.status_code in [200, 400]
        data = response.json()
        print(f"PASS: Generate narration response: {response.status_code} - {data}")


class TestHeyGenIntegration:
    """Test HeyGen avatar and voice endpoints"""

    EXISTING_PROJECT_ID = "cebb110f-ced1-4e62-8478-7fb6bd99943d"

    def test_heygen_avatars(self, api_client):
        """Test GET /api/heygen/avatars?limit=5"""
        response = api_client.get(f"{BASE_URL}/api/heygen/avatars?limit=5")
        assert response.status_code == 200
        data = response.json()
        assert "avatars" in data
        assert len(data["avatars"]) <= 5
        if len(data["avatars"]) > 0:
            avatar = data["avatars"][0]
            assert "avatar_id" in avatar
            assert "avatar_name" in avatar
            print(f"PASS: HeyGen avatars - {len(data['avatars'])} avatars returned")
            for av in data["avatars"]:
                print(f"  - {av.get('avatar_name')} (ID: {av.get('avatar_id')})")

    def test_heygen_portuguese_voices(self, api_client):
        """Test GET /api/heygen/voices?language=portuguese"""
        response = api_client.get(f"{BASE_URL}/api/heygen/voices?language=portuguese")
        assert response.status_code == 200
        data = response.json()
        assert "voices" in data
        print(f"PASS: HeyGen Portuguese voices - {len(data['voices'])} voices")
        for v in data["voices"][:5]:
            print(f"  - {v.get('name')} ({v.get('gender', 'N/A')})")

    def test_heygen_status_existing_project(self, api_client):
        """Test GET /api/agent/projects/{id}/heygen-status"""
        response = api_client.get(f"{BASE_URL}/api/agent/projects/{self.EXISTING_PROJECT_ID}/heygen-status")
        assert response.status_code == 200
        data = response.json()
        assert "status" in data
        # Status can be 'no_heygen', 'processing', 'all_done'
        print(f"PASS: HeyGen status for existing project: {data}")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
