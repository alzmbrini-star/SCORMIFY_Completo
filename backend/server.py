from fastapi import FastAPI, APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks, Request
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient, AsyncIOMotorGridFSBucket
import os
import re
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import shutil
import aiofiles
import httpx
import io
import asyncio
import json
import base64

from models import (
    Project, ProjectCreate, ProjectUpdate, Course, CourseMetadata,
    Slide, SlideCreate, SlideUpdate, SlideElement, ElementCreate, ElementUpdate,
    Animation, Annotation, AnnotationCreate, SlideAudio, GlobalAudio,
    JobStatus, ReorderSlidesRequest,
    QuizQuestion, QuizQuestionCreate, QuizQuestionUpdate, QuizAlternative,
    QuizConfig, QuizAttempt, QuizGenerateRequest, QuizSubmitRequest
)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env', override=False)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Check and install system dependencies (LibreOffice, poppler-utils)
# Deferred to background startup task for faster server boot
logger.info("Starting Scormify API server...")

# Storage directories
STORAGE_DIR = ROOT_DIR / "storage"
UPLOADS_DIR = STORAGE_DIR / "uploads"
PROJECTS_DIR = STORAGE_DIR / "projects"
EXPORTS_DIR = STORAGE_DIR / "exports"

# Create directories
STORAGE_DIR.mkdir(exist_ok=True)
UPLOADS_DIR.mkdir(exist_ok=True)
PROJECTS_DIR.mkdir(exist_ok=True)
EXPORTS_DIR.mkdir(exist_ok=True)

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# GridFS bucket for persistent export storage
exports_bucket = AsyncIOMotorGridFSBucket(db, bucket_name="exports")

# HeyGen API Configuration
HEYGEN_API_KEY = os.environ.get('HEYGEN_API_KEY', '')
HEYGEN_BASE_URL = "https://api.heygen.com"
HEYGEN_HEADERS = {
    "X-Api-Key": HEYGEN_API_KEY,
    "Accept": "application/json",
    "Content-Type": "application/json"
}

# ElevenLabs API Configuration
ELEVENLABS_API_KEY = os.environ.get('ELEVENLABS_API_KEY', '')

# HeyGen Credits Cache (to avoid slow repeated API calls)
heygen_credits_cache: Dict[str, Any] = {
    "data": None,
    "timestamp": None,
    "ttl": 60  # Cache for 60 seconds
}

# SSE Event Store for HeyGen webhook notifications
# Maps video_id to list of asyncio.Event objects waiting for updates
heygen_sse_subscribers: Dict[str, List[asyncio.Queue]] = {}

# Create the main app
app = FastAPI(title="Scormify API", version="1.0.0")

# Create routers
api_router = APIRouter(prefix="/api")

# Import and setup auth routes
from routes import auth as auth_routes
from routes import companies as companies_routes
from routes import users as users_routes

# Set database connections for route modules
auth_routes.set_db(db)
companies_routes.set_db(db)
users_routes.set_db(db)

# Include routers
api_router.include_router(auth_routes.router)
api_router.include_router(companies_routes.router)
api_router.include_router(users_routes.router)

# Job tracking (in-memory for simplicity)
jobs = {}

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Helper functions
def now_utc():
    return datetime.now(timezone.utc)

def serialize_doc(doc: dict) -> dict:
    """Serialize MongoDB document for JSON response"""
    if doc is None:
        return None
    doc.pop('_id', None)
    for key, value in doc.items():
        if isinstance(value, datetime):
            doc[key] = value.isoformat()
    return doc

async def get_project_by_id(project_id: str) -> Optional[dict]:
    """Get project from database"""
    doc = await db.projects.find_one({"id": project_id}, {"_id": 0})
    return doc

async def update_project(project_id: str, update_data: dict):
    """Update project in database"""
    update_data['updatedAt'] = now_utc().isoformat()
    await db.projects.update_one({"id": project_id}, {"$set": update_data})

# Background task for PPT processing
def process_ppt_upload(job_id: str, file_path: str, project_id: str):
    """Process uploaded PPT file in background using high-fidelity parser"""
    from pymongo import MongoClient
    from services.ppt_image_parser import parse_pptx_high_fidelity
    
    try:
        jobs[job_id]['status'] = 'processing'
        jobs[job_id]['message'] = 'Converting PowerPoint slides to images...'
        jobs[job_id]['progress'] = 10
        
        # Use high-fidelity parser that renders slides as images
        course = parse_pptx_high_fidelity(file_path, project_id, str(PROJECTS_DIR))
        
        jobs[job_id]['progress'] = 80
        jobs[job_id]['message'] = 'Saving course data...'
        
        # Use synchronous MongoDB client for background task
        sync_client = MongoClient(mongo_url)
        sync_db = sync_client[os.environ['DB_NAME']]
        
        course_dict = course.model_dump()
        course_dict['createdAt'] = course.createdAt.isoformat()
        course_dict['updatedAt'] = course.updatedAt.isoformat()
        
        sync_db.projects.update_one(
            {"id": project_id},
            {"$set": {
                "course": course_dict,
                "status": "ready",
                "updatedAt": now_utc().isoformat()
            }}
        )
        
        sync_client.close()
        
        jobs[job_id]['status'] = 'completed'
        jobs[job_id]['progress'] = 100
        jobs[job_id]['message'] = 'Processing complete - slides rendered with high fidelity'
        jobs[job_id]['result'] = {'projectId': project_id}
        
    except Exception as e:
        logger.error(f"Error processing PPT: {e}")
        jobs[job_id]['status'] = 'failed'
        jobs[job_id]['message'] = str(e)
    finally:
        # Clean up uploaded file
        try:
            os.remove(file_path)
        except OSError:
            pass

# API Routes

@api_router.get("/")
async def root():
    return {"message": "Scormify API v1.0"}

@api_router.get("/health")
async def health():
    return {"status": "healthy"}

# VLibras CORS proxy - resolves cross-origin issues with VLibras government servers
# Handles both dicionario2.vlibras.gov.br and traducao2.vlibras.gov.br
_VLIBRAS_DOMAINS = {
    "dicionario2": "https://dicionario2.vlibras.gov.br",
    "traducao2": "https://traducao2.vlibras.gov.br",
}
_VLIBRAS_BROWSER_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "*/*",
    "Accept-Language": "pt-BR,pt;q=0.9,en;q=0.8",
    "Referer": "https://www.vlibras.gov.br/",
    "Origin": "https://www.vlibras.gov.br",
}

async def _vlibras_proxy_request(domain_key: str, path: str, request: Request):
    """Generic VLibras proxy handler for GET and POST requests."""
    import httpx
    from fastapi.responses import Response
    base_url = _VLIBRAS_DOMAINS.get(domain_key)
    if not base_url:
        raise HTTPException(status_code=400, detail="Invalid VLibras domain key")
    url = f"{base_url}/{path}" if path else base_url
    query = str(request.query_params)
    if query:
        url += f"?{query}"
    try:
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            if request.method == "POST":
                body = await request.body()
                content_type = request.headers.get("content-type", "application/json")
                resp = await client.post(url, content=body, headers={**_VLIBRAS_BROWSER_HEADERS, "Content-Type": content_type})
            else:
                resp = await client.get(url, headers=_VLIBRAS_BROWSER_HEADERS)
        return Response(
            content=resp.content,
            status_code=resp.status_code,
            media_type=resp.headers.get("content-type", "application/octet-stream"),
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
                "Access-Control-Allow-Headers": "Content-Type",
                "Cache-Control": "public, max-age=86400" if request.method == "GET" else "no-cache",
            },
        )
    except Exception as e:
        logger.warning(f"VLibras proxy error ({domain_key}/{path}): {e}")
        raise HTTPException(status_code=502, detail=f"VLibras proxy error: {str(e)}")

@api_router.api_route("/vlibras-proxy/dicionario2/{path:path}", methods=["GET", "OPTIONS"])
async def vlibras_dicionario_proxy(path: str, request: Request):
    if request.method == "OPTIONS":
        from fastapi.responses import Response
        return Response(status_code=204, headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type",
            "Access-Control-Max-Age": "86400",
        })
    return await _vlibras_proxy_request("dicionario2", path, request)

@api_router.api_route("/vlibras-proxy/traducao2/{path:path}", methods=["GET", "POST", "OPTIONS"])
async def vlibras_traducao_proxy(path: str, request: Request):
    if request.method == "OPTIONS":
        from fastapi.responses import Response
        return Response(status_code=204, headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type",
            "Access-Control-Max-Age": "86400",
        })
    return await _vlibras_proxy_request("traducao2", path, request)

# Keep old endpoint for backwards compatibility
@api_router.get("/vlibras-dict/{path:path}")
async def vlibras_dict_proxy_legacy(path: str, request: Request):
    return await _vlibras_proxy_request("dicionario2", path, request)

from starlette.responses import HTMLResponse as StarletteHTMLResponse

@api_router.get("/vlibras-test/{path:path}")
async def vlibras_test_assets(path: str):
    test_dir = Path(__file__).parent / "static_test" / "scorm_live"
    file_path = test_dir / path
    if file_path.exists() and file_path.is_file():
        media_types = {'.js': 'application/javascript', '.css': 'text/css', '.json': 'application/json', '.png': 'image/png', '.jpg': 'image/jpeg', '.mp3': 'audio/mpeg', '.html': 'text/html'}
        return FileResponse(str(file_path), media_type=media_types.get(file_path.suffix.lower(), 'application/octet-stream'))
    raise HTTPException(status_code=404, detail="File not found")

@api_router.get("/vlibras-test")
async def vlibras_test():
    # Rewrite relative paths to be absolute for the test server
    test_dir = Path(__file__).parent / "static_test" / "scorm_live"
    html_path = test_dir / "index.html"
    html = html_path.read_text()
    # Add base tag to resolve all relative URLs correctly
    html = html.replace('<head>', '<head>\n    <base href="/api/vlibras-test/">')
    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=html)

@api_router.get("/vlibras-simple-test")
async def vlibras_simple_test():
    """Simple VLibras test page - NO proxy, direct browser access"""
    test_path = Path(__file__).parent / "static_test" / "vlibras_simple_test.html"
    html = test_path.read_text()
    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=html)

@api_router.get("/vlibras-proxy-test")
async def vlibras_proxy_test(request: Request):
    """VLibras test page WITH CORS proxy - tests the full proxy flow"""
    test_path = Path(__file__).parent / "static_test" / "vlibras_proxy_test.html"
    html = test_path.read_text()
    # Determine external URL - use origin header, x-forwarded-host, or read from frontend .env
    origin = request.headers.get('origin', '')
    if not origin:
        fwd_host = request.headers.get('x-forwarded-host', '')
        scheme = request.headers.get('x-forwarded-proto', 'https')
        if fwd_host:
            origin = f"{scheme}://{fwd_host}"
        else:
            # Read from frontend .env as fallback
            try:
                env_path = Path(__file__).parent.parent / "frontend" / ".env"
                for line in env_path.read_text().splitlines():
                    if line.startswith("REACT_APP_BACKEND_URL="):
                        origin = line.split("=", 1)[1].strip()
                        break
            except Exception:
                host = request.headers.get('host', '')
                origin = f"https://{host}" if host else ''
    proxy_base = origin.rstrip('/') + '/api/vlibras-proxy'
    html = html.replace('__PROXY_BASE__', proxy_base)
    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=html)

# Project Routes

@api_router.get("/projects", response_model=List[dict])
async def list_projects():
    """List all projects"""
    projects = await db.projects.find({}, {"_id": 0}).sort("createdAt", -1).to_list(100)
    return projects

@api_router.post("/projects", response_model=dict)
async def create_project(data: ProjectCreate):
    """Create a new project"""
    project = Project(
        name=data.name,
        description=data.description
    )
    
    # Set course metadata title to project name
    project.course.metadata.title = data.name
    
    # Create default first slide
    default_slide = Slide(
        title="Slide 1",
        order=0,
        background="#FFFFFF"
    )
    project.course.slides = [default_slide]
    
    project_dict = project.model_dump()
    project_dict['createdAt'] = project.createdAt.isoformat()
    project_dict['updatedAt'] = project.updatedAt.isoformat()
    project_dict['course']['createdAt'] = project.course.createdAt.isoformat()
    project_dict['course']['updatedAt'] = project.course.updatedAt.isoformat()
    
    await db.projects.insert_one(project_dict)
    
    # Create project directory
    project_dir = PROJECTS_DIR / project.id
    (project_dir / "assets").mkdir(parents=True, exist_ok=True)
    
    return serialize_doc(project_dict)

@api_router.get("/projects/{project_id}", response_model=dict)
async def get_project(project_id: str):
    """Get project by ID"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project

@api_router.put("/projects/{project_id}", response_model=dict)
async def update_project_endpoint(project_id: str, data: ProjectUpdate):
    """Update project"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    update_data = data.model_dump(exclude_unset=True)
    
    # If project name is being updated, also update course metadata title
    if 'name' in update_data:
        update_data['course.metadata.title'] = update_data['name']
    
    await update_project(project_id, update_data)
    
    return await get_project_by_id(project_id)

@api_router.delete("/projects/{project_id}")
async def delete_project(project_id: str):
    """Delete project"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    await db.projects.delete_one({"id": project_id})
    
    # Delete project files
    project_dir = PROJECTS_DIR / project_id
    if project_dir.exists():
        shutil.rmtree(project_dir)
    
    return {"message": "Project deleted"}

# PPT Upload

@api_router.post("/ppt/upload")
async def upload_ppt(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    project_name: Optional[str] = None
):
    """Upload and process a PPT/PPTX file"""
    # Validate file type
    if not file.filename.lower().endswith(('.ppt', '.pptx')):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PPT/PPTX files are allowed.")
    
    # Validate file size (max 50MB)
    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 50MB.")
    
    # Create project
    project_name = project_name or Path(file.filename).stem
    project = Project(name=project_name)
    
    project_dict = project.model_dump()
    project_dict['createdAt'] = project.createdAt.isoformat()
    project_dict['updatedAt'] = project.updatedAt.isoformat()
    project_dict['course']['createdAt'] = project.course.createdAt.isoformat()
    project_dict['course']['updatedAt'] = project.course.updatedAt.isoformat()
    project_dict['status'] = 'processing'
    
    await db.projects.insert_one(project_dict)
    
    # Create project directory
    project_dir = PROJECTS_DIR / project.id
    (project_dir / "assets").mkdir(parents=True, exist_ok=True)
    
    # Save uploaded file
    upload_path = UPLOADS_DIR / f"{project.id}_{file.filename}"
    async with aiofiles.open(upload_path, 'wb') as f:
        await f.write(content)
    
    # Create job
    job_id = str(uuid.uuid4())
    jobs[job_id] = {
        'id': job_id,
        'status': 'pending',
        'progress': 0,
        'message': 'Upload received, starting processing...',
        'result': None
    }
    
    # Start background processing
    background_tasks.add_task(process_ppt_upload, job_id, str(upload_path), project.id)
    
    return {
        "jobId": job_id,
        "projectId": project.id,
        "message": "File uploaded, processing started"
    }

@api_router.get("/job/{job_id}", response_model=JobStatus)
async def get_job_status(job_id: str):
    """Get job status"""
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    return JobStatus(**jobs[job_id])

# Course Routes

@api_router.get("/course/{project_id}")
async def get_course(project_id: str):
    """Get course data for a project"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project.get('course', {})

@api_router.post("/course/{project_id}/save")
async def save_course(project_id: str, course_data: dict):
    """Save course data"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    await update_project(project_id, {"course": course_data})
    return {"message": "Course saved"}

# Slide Routes

@api_router.post("/projects/{project_id}/slides")
async def create_slide(project_id: str, data: SlideCreate):
    """Add a new slide to the project"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    # Get dimensions from first slide if not provided, to maintain consistency
    first_slide = slides[0] if slides else None
    slide_width = data.width or (first_slide.get('width') if first_slide else 1280)
    slide_height = data.height or (first_slide.get('height') if first_slide else 720)
    
    new_slide = Slide(
        title=data.title,
        background=data.background,
        width=slide_width,
        height=slide_height,
        order=len(slides)
    )
    
    slides.append(new_slide.model_dump())
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return new_slide.model_dump()

@api_router.put("/projects/{project_id}/slides/{slide_id}")
async def update_slide(project_id: str, slide_id: str, data: SlideUpdate):
    """Update a slide"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    update_data = data.model_dump(exclude_unset=True)
    slides[slide_index].update(update_data)
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return slides[slide_index]

@api_router.delete("/projects/{project_id}/slides/{slide_id}")
async def delete_slide(project_id: str, slide_id: str):
    """Delete a slide"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slides = [s for s in slides if s.get('id') != slide_id]
    
    # Re-order slides
    for i, slide in enumerate(slides):
        slide['order'] = i
    
    course['slides'] = slides
    await update_project(project_id, {"course": course})
    
    return {"message": "Slide deleted"}

@api_router.post("/projects/{project_id}/slides/{slide_id}/duplicate")
async def duplicate_slide(project_id: str, slide_id: str):
    """Duplicate a slide"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    # Deep copy the slide
    import copy
    new_slide = copy.deepcopy(slides[slide_index])
    new_slide['id'] = str(uuid.uuid4())
    new_slide['title'] = f"{new_slide.get('title', 'Slide')} (copy)"
    new_slide['order'] = slide_index + 1
    
    # Insert after original
    slides.insert(slide_index + 1, new_slide)
    
    # Re-order subsequent slides
    for i in range(slide_index + 2, len(slides)):
        slides[i]['order'] = i
    
    course['slides'] = slides
    await update_project(project_id, {"course": course})
    
    return new_slide

@api_router.post("/projects/{project_id}/normalize-dimensions")
async def normalize_slide_dimensions(project_id: str, target_width: int = 1536, target_height: int = 864):
    """Normalize all slides to the same dimensions, scaling elements proportionally"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    normalized_count = 0
    
    for slide in slides:
        current_width = slide.get('width', 1536)
        current_height = slide.get('height', 864)
        
        # Skip if already the target dimensions
        if current_width == target_width and current_height == target_height:
            continue
        
        # Calculate scale factors
        scale_x = target_width / current_width
        scale_y = target_height / current_height
        
        # Update slide dimensions
        slide['width'] = target_width
        slide['height'] = target_height
        
        # Scale all elements proportionally
        for element in slide.get('elements', []):
            # Scale position
            if 'x' in element:
                element['x'] = element['x'] * scale_x
            if 'y' in element:
                element['y'] = element['y'] * scale_y
            
            # Scale size
            if 'width' in element:
                element['width'] = element['width'] * scale_x
            if 'height' in element:
                element['height'] = element['height'] * scale_y
        
        # Scale annotations if present
        for annotation in slide.get('annotations', []):
            if 'points' in annotation:
                for point in annotation['points']:
                    if 'x' in point:
                        point['x'] = point['x'] * scale_x
                    if 'y' in point:
                        point['y'] = point['y'] * scale_y
        
        normalized_count += 1
    
    # Save updated course
    course['slides'] = slides
    await update_project(project_id, {"course": course})
    
    return {
        "message": f"Normalized {normalized_count} slides to {target_width}x{target_height}",
        "normalized_count": normalized_count,
        "target_dimensions": {"width": target_width, "height": target_height}
    }

@api_router.post("/projects/{project_id}/slides/reorder")
async def reorder_slides(project_id: str, data: ReorderSlidesRequest):
    """Reorder slides"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    # Create mapping
    slide_map = {s['id']: s for s in slides}
    
    # Reorder based on provided IDs
    new_slides = []
    for i, slide_id in enumerate(data.slideIds):
        if slide_id in slide_map:
            slide = slide_map[slide_id]
            slide['order'] = i
            new_slides.append(slide)
    
    course['slides'] = new_slides
    await update_project(project_id, {"course": course})
    
    return {"message": "Slides reordered"}

# Element Routes

@api_router.post("/projects/{project_id}/slides/{slide_id}/elements")
async def add_element(project_id: str, slide_id: str, data: ElementCreate):
    """Add element to slide"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    element_data = data.model_dump(exclude_unset=True)
    # Ensure style is a dict if not provided
    if 'style' not in element_data or element_data.get('style') is None:
        element_data['style'] = {}
    element = SlideElement(**element_data)
    elements = slides[slide_index].get('elements', [])
    element_dict = element.model_dump()
    element_dict['zIndex'] = len(elements)
    elements.append(element_dict)
    
    slides[slide_index]['elements'] = elements
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return element_dict

@api_router.put("/projects/{project_id}/slides/{slide_id}/elements/{element_id}")
async def update_element(project_id: str, slide_id: str, element_id: str, data: ElementUpdate):
    """Update element"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    elements = slides[slide_index].get('elements', [])
    elem_index = next((i for i, e in enumerate(elements) if e.get('id') == element_id), None)
    if elem_index is None:
        raise HTTPException(status_code=404, detail="Element not found")
    
    update_data = data.model_dump(exclude_unset=True)
    elements[elem_index].update(update_data)
    
    slides[slide_index]['elements'] = elements
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return elements[elem_index]

@api_router.delete("/projects/{project_id}/slides/{slide_id}/elements/{element_id}")
async def delete_element(project_id: str, slide_id: str, element_id: str):
    """Delete element"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    elements = slides[slide_index].get('elements', [])
    elements = [e for e in elements if e.get('id') != element_id]
    
    slides[slide_index]['elements'] = elements
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return {"message": "Element deleted"}

# Media Upload

@api_router.post("/projects/{project_id}/media")
async def upload_media(project_id: str, file: UploadFile = File(...)):
    """Upload media file (image, audio, video) with automatic image optimization"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Validate file type
    allowed_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg', '.mp3', '.wav', '.ogg', '.mp4', '.webm'}
    ext = Path(file.filename).suffix.lower()
    if ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Invalid file type. Allowed: {', '.join(allowed_extensions)}")
    
    # Read and validate size
    content = await file.read()
    original_size = len(content)
    max_size = 100 * 1024 * 1024  # 100MB
    if original_size > max_size:
        raise HTTPException(status_code=400, detail="File too large")
    
    # Optimize images automatically
    image_extensions = {'.png', '.jpg', '.jpeg', '.webp'}
    optimized = False
    final_content = content
    
    if ext in image_extensions:
        try:
            from PIL import Image
            # Open image with Pillow
            img = Image.open(io.BytesIO(content))
            original_width, original_height = img.size
            
            # Only optimize if image is large (> 500KB or dimensions > 1920px)
            should_optimize = original_size > 500 * 1024 or img.width > 1920 or img.height > 1080
            
            if should_optimize:
                # Convert RGBA to RGB for JPEG (remove alpha channel)
                if img.mode == 'RGBA' and ext in {'.jpg', '.jpeg'}:
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[3])
                    img = background
                elif img.mode == 'RGBA':
                    pass  # Keep alpha for PNG/WebP
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Calculate max dimensions (Full HD)
                max_width = 1920
                max_height = 1080
                
                # Resize if image is larger than max dimensions
                if img.width > max_width or img.height > max_height:
                    ratio = min(max_width / img.width, max_height / img.height)
                    new_size = (int(img.width * ratio), int(img.height * ratio))
                    img = img.resize(new_size, Image.Resampling.LANCZOS)
                    logger.info(f"Resized image from {original_width}x{original_height} to {new_size[0]}x{new_size[1]}")
                
                # Save optimized image to buffer
                output = io.BytesIO()
                
                if ext == '.png':
                    # For PNG, use optimize and reduce colors if possible
                    img.save(output, format='PNG', optimize=True)
                elif ext == '.webp':
                    # WebP with quality compression
                    img.save(output, format='WEBP', quality=85, method=6)
                else:
                    # JPEG with quality compression
                    img.save(output, format='JPEG', quality=85, optimize=True)
                
                optimized_content = output.getvalue()
                
                # Only use optimized version if it's actually smaller
                if len(optimized_content) < original_size:
                    final_content = optimized_content
                    optimized = True
                    logger.info(f"Image optimized: {original_size} bytes -> {len(final_content)} bytes ({100 - (len(final_content)/original_size*100):.1f}% reduction)")
                else:
                    logger.info(f"Optimization skipped: optimized size ({len(optimized_content)}) >= original ({original_size})")
            
        except Exception as e:
            logger.warning(f"Image optimization failed, using original: {e}")
            # Use original content if optimization fails
    
    # Save file
    file_id = str(uuid.uuid4())
    filename = f"{file_id}{ext}"
    file_path = PROJECTS_DIR / project_id / "assets" / filename
    file_path.parent.mkdir(parents=True, exist_ok=True)
    
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(final_content)
    
    # Persist in MongoDB for production environments with ephemeral storage (non-blocking)
    import threading
    try:
        from services.asset_store import store_asset_sync
        threading.Thread(
            target=store_asset_sync,
            args=(mongo_url, os.environ['DB_NAME'], project_id, filename, str(file_path)),
            daemon=True
        ).start()
    except Exception as e:
        logger.warning(f"Failed to persist media in MongoDB (non-fatal): {e}")
    
    return {
        "id": file_id,
        "filename": filename,
        "url": f"/api/projects/{project_id}/assets/{filename}",
        "size": len(final_content),
        "originalSize": original_size,
        "optimized": optimized,
        "type": ext[1:]
    }

# Audio Recording

@api_router.post("/projects/{project_id}/slides/{slide_id}/audio")
async def upload_slide_audio(
    project_id: str,
    slide_id: str,
    file: UploadFile = File(...),
    audio_type: str = Form("narration")
):
    """Upload audio for a slide"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Validate
    if not file.filename.lower().endswith(('.mp3', '.wav', '.ogg', '.webm')):
        raise HTTPException(status_code=400, detail="Invalid audio format")
    
    content = await file.read()
    
    # Save file
    file_id = str(uuid.uuid4())
    ext = Path(file.filename).suffix.lower()
    filename = f"audio_{file_id}{ext}"
    file_path = PROJECTS_DIR / project_id / "assets" / filename
    file_path.parent.mkdir(parents=True, exist_ok=True)
    
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(content)
    
    # Persist in MongoDB for production environments with ephemeral storage (non-blocking)
    import threading
    try:
        from services.asset_store import store_asset_sync
        threading.Thread(
            target=store_asset_sync,
            args=(mongo_url, os.environ['DB_NAME'], project_id, filename, str(file_path)),
            daemon=True
        ).start()
    except Exception as e:
        logger.warning(f"Failed to persist audio in MongoDB (non-fatal): {e}")
    
    # Update slide with audio
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    audio_data = {
        "id": file_id,
        "type": audio_type,
        "src": f"/api/projects/{project_id}/assets/{filename}",
        "filename": filename,
        "duration": 0,
        "volume": 1.0
    }
    
    audio_list = slides[slide_index].get('audio', [])
    audio_list.append(audio_data)
    slides[slide_index]['audio'] = audio_list
    
    course['slides'] = slides
    await update_project(project_id, {"course": course})
    
    return audio_data

# Global Audio (Soundtrack)

@api_router.post("/projects/{project_id}/global-audio")
async def set_global_audio(project_id: str, file: UploadFile = File(...)):
    """Set global soundtrack for the course"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    if not file.filename.lower().endswith(('.mp3', '.wav', '.ogg')):
        raise HTTPException(status_code=400, detail="Invalid audio format")
    
    content = await file.read()
    
    file_id = str(uuid.uuid4())
    ext = Path(file.filename).suffix.lower()
    filename = f"global_audio_{file_id}{ext}"
    file_path = PROJECTS_DIR / project_id / "assets" / filename
    file_path.parent.mkdir(parents=True, exist_ok=True)
    
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(content)
    
    # Persist in MongoDB for production environments with ephemeral storage (non-blocking)
    import threading
    try:
        from services.asset_store import store_asset_sync
        threading.Thread(
            target=store_asset_sync,
            args=(mongo_url, os.environ['DB_NAME'], project_id, filename, str(file_path)),
            daemon=True
        ).start()
    except Exception as e:
        logger.warning(f"Failed to persist global audio in MongoDB (non-fatal): {e}")
    
    global_audio = {
        "id": file_id,
        "src": f"/api/projects/{project_id}/assets/{filename}",
        "filename": filename,
        "duration": 0,
        "volume": 0.5,
        "loop": True
    }
    
    course = project.get('course', {})
    course['globalAudio'] = global_audio
    
    await update_project(project_id, {"course": course})
    
    return global_audio


@api_router.delete("/projects/{project_id}/global-audio")
async def remove_global_audio(project_id: str):
    """Remove global audio from project"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    
    # Remove global audio file if exists
    if course.get('globalAudio'):
        old_file = course['globalAudio'].get('filename')
        if old_file:
            assets_dir = Path(f"storage/projects/{project_id}/assets")
            old_path = assets_dir / old_file
            if old_path.exists():
                old_path.unlink()
    
    course['globalAudio'] = None
    await update_project(project_id, {"course": course})
    
    return {"message": "Global audio removed"}


@api_router.put("/projects/{project_id}/global-audio/volume")
async def update_global_audio_volume(project_id: str, volume: float):
    """Update global audio volume"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    if not course.get('globalAudio'):
        raise HTTPException(status_code=404, detail="No global audio set")
    
    # Clamp volume between 0 and 1
    volume = max(0.0, min(1.0, volume))
    course['globalAudio']['volume'] = volume
    
    await update_project(project_id, {"course": course})
    
    return course['globalAudio']


@api_router.delete("/projects/{project_id}/slides/{slide_id}/audio/{audio_id}")
async def remove_slide_audio(project_id: str, slide_id: str, audio_id: str):
    """Remove audio from slide"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    audio_list = slides[slide_index].get('audio', [])
    
    # Find and remove audio
    audio_to_remove = next((a for a in audio_list if a.get('id') == audio_id), None)
    if not audio_to_remove:
        raise HTTPException(status_code=404, detail="Audio not found")
    
    # Remove audio file
    if audio_to_remove.get('filename'):
        assets_dir = Path(f"storage/projects/{project_id}/assets")
        audio_path = assets_dir / audio_to_remove['filename']
        if audio_path.exists():
            audio_path.unlink()
    
    # Update slide
    slides[slide_index]['audio'] = [a for a in audio_list if a.get('id') != audio_id]
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return {"message": "Audio removed from slide"}


@api_router.put("/projects/{project_id}/slides/{slide_id}/audio/{audio_id}/volume")
async def update_slide_audio_volume(project_id: str, slide_id: str, audio_id: str, volume: float):
    """Update slide audio volume"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    audio_list = slides[slide_index].get('audio', [])
    audio_index = next((i for i, a in enumerate(audio_list) if a.get('id') == audio_id), None)
    if audio_index is None:
        raise HTTPException(status_code=404, detail="Audio not found")
    
    # Clamp volume between 0 and 1
    volume = max(0.0, min(1.0, volume))
    audio_list[audio_index]['volume'] = volume
    slides[slide_index]['audio'] = audio_list
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return audio_list[audio_index]


@api_router.put("/projects/{project_id}/slides/{slide_id}/audio/{audio_id}/timing")
async def update_slide_audio_timing(project_id: str, slide_id: str, audio_id: str, data: dict):
    """Update slide audio timing (startTime and duration for trimming)"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    audio_list = slides[slide_index].get('audio', [])
    audio_index = next((i for i, a in enumerate(audio_list) if a.get('id') == audio_id), None)
    if audio_index is None:
        raise HTTPException(status_code=404, detail="Audio not found")
    
    # Store original duration if not already stored
    if 'originalDuration' not in audio_list[audio_index]:
        audio_list[audio_index]['originalDuration'] = audio_list[audio_index].get('duration', 10)
    
    # Update timing
    if data.get('startTime') is not None:
        audio_list[audio_index]['startTime'] = max(0, data['startTime'])
    if data.get('duration') is not None:
        audio_list[audio_index]['duration'] = max(0.5, data['duration'])
    
    slides[slide_index]['audio'] = audio_list
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return audio_list[audio_index]


# Annotations

@api_router.post("/projects/{project_id}/slides/{slide_id}/annotations")
async def add_annotation(project_id: str, slide_id: str, data: AnnotationCreate):
    """Add annotation to slide"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    annotation = Annotation(**data.model_dump())
    annotations = slides[slide_index].get('annotations', [])
    annotations.append(annotation.model_dump())
    
    slides[slide_index]['annotations'] = annotations
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return annotation.model_dump()

@api_router.put("/projects/{project_id}/slides/{slide_id}/annotations/{annotation_id}")
async def update_annotation(project_id: str, slide_id: str, annotation_id: str, update_data: dict):
    """Update annotation (for timeline settings)"""
    from models import AnnotationUpdate
    
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    annotations = slides[slide_index].get('annotations', [])
    annotation_index = next((i for i, a in enumerate(annotations) if a.get('id') == annotation_id), None)
    if annotation_index is None:
        raise HTTPException(status_code=404, detail="Annotation not found")
    
    # Update annotation with new data
    for key, value in update_data.items():
        if value is not None:
            annotations[annotation_index][key] = value
    
    slides[slide_index]['annotations'] = annotations
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return annotations[annotation_index]

@api_router.delete("/projects/{project_id}/slides/{slide_id}/annotations/{annotation_id}")
async def delete_annotation(project_id: str, slide_id: str, annotation_id: str):
    """Delete annotation"""
    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    course = project.get('course', {})
    slides = course.get('slides', [])
    
    slide_index = next((i for i, s in enumerate(slides) if s.get('id') == slide_id), None)
    if slide_index is None:
        raise HTTPException(status_code=404, detail="Slide not found")
    
    annotations = slides[slide_index].get('annotations', [])
    annotations = [a for a in annotations if a.get('id') != annotation_id]
    
    slides[slide_index]['annotations'] = annotations
    course['slides'] = slides
    
    await update_project(project_id, {"course": course})
    
    return {"message": "Annotation deleted"}

# SCORM Export

def _cleanup_old_exports(exports_dir: str, max_age_hours: int = 24):
    """Delete export files older than max_age_hours to prevent disk exhaustion."""
    import time
    exports_path = Path(exports_dir)
    if not exports_path.exists():
        return
    cutoff = time.time() - max_age_hours * 3600
    removed = 0
    for f in exports_path.iterdir():
        if f.is_file() and f.stat().st_mtime < cutoff:
            try:
                f.unlink()
                removed += 1
            except Exception:
                pass
    if removed:
        logger.info(f"Cleaned up {removed} old export files from {exports_dir}")


async def save_export_to_gridfs(file_path: str, filename: str):
    """Save an export file to MongoDB GridFS for persistence across deploys."""
    try:
        # Delete any existing file with same name
        cursor = exports_bucket.find({"filename": filename})
        async for grid_file in cursor:
            await exports_bucket.delete(grid_file._id)
        # Upload the new file
        with open(file_path, "rb") as f:
            data = f.read()
        await exports_bucket.upload_from_stream(filename, io.BytesIO(data))
        logger.info(f"Saved export to GridFS: {filename} ({len(data)} bytes)")
    except Exception as e:
        logger.warning(f"GridFS save failed (non-fatal): {e}")


async def get_export_from_gridfs(filename: str, dest_path: str) -> bool:
    """Retrieve an export file from GridFS and write to dest_path. Returns True if found."""
    try:
        cursor = exports_bucket.find({"filename": filename}, limit=1).sort("uploadDate", -1)
        grid_file = await cursor.next()
        stream = await exports_bucket.open_download_stream(grid_file._id)
        data = await stream.read()
        Path(dest_path).parent.mkdir(parents=True, exist_ok=True)
        with open(dest_path, "wb") as f:
            f.write(data)
        logger.info(f"Restored export from GridFS: {filename} ({len(data)} bytes)")
        return True
    except StopAsyncIteration:
        return False
    except Exception as e:
        logger.warning(f"GridFS retrieve failed: {e}")
        return False


@api_router.post("/course/{project_id}/export-scorm")
async def export_scorm(project_id: str, request: Request, background_tasks: BackgroundTasks):
    """Export project as SCORM 1.2 package"""
    project_doc = await get_project_by_id(project_id)
    if not project_doc:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Create job
    job_id = str(uuid.uuid4())
    jobs[job_id] = {
        'id': job_id,
        'status': 'processing',
        'progress': 0,
        'message': 'Generating SCORM package...',
        'result': None
    }
    
    try:
        # Convert dict to Project model
        project = Project(**project_doc)
        
        # Collect all question IDs from quiz elements in the course
        question_ids = set()
        for slide in project.course.slides:
            for element in slide.elements:
                if element.type == 'quiz' and hasattr(element, 'quizConfig'):
                    quiz_config = element.quizConfig
                    if quiz_config and isinstance(quiz_config, dict):
                        question_ids.update(quiz_config.get('questionIds', []))
        
        # Load questions from database if there are quiz elements
        questions = []
        if question_ids:
            question_docs = await db.questions.find(
                {"id": {"$in": list(question_ids)}},
                {"_id": 0}
            ).to_list(500)
            questions = question_docs
            logger.info(f"Loaded {len(questions)} questions for SCORM export")
        
        # Load tutor settings
        tutor_settings = None
        try:
            settings_doc = await db.settings.find_one({"key": "tutor"}, {"_id": 0})
            if settings_doc and settings_doc.get("enabled"):
                # Detect the backend URL from the current request
                # This works automatically in both dev and production
                origin = request.headers.get('origin', '')
                if origin:
                    backend_url = origin
                else:
                    fwd_host = request.headers.get('x-forwarded-host', '')
                    scheme = request.headers.get('x-forwarded-proto', 'https')
                    if fwd_host:
                        backend_url = f"{scheme}://{fwd_host}"
                    else:
                        host = request.headers.get('host', '')
                        backend_url = f"{scheme}://{host}" if host else ''
                
                tutor_settings = {
                    'enabled': True,
                    'apiUrl': backend_url,
                    'tutorName': settings_doc.get('tutorName', 'Tutor IA'),
                    'messageLimit': settings_doc.get('messageLimit', 50),
                    'suggestedQuestions': settings_doc.get('suggestedQuestions', []),
                    'courseTopic': project.course.metadata.title or project.name
                }
        except Exception as e:
            logger.warning(f"Tutor settings load failed (non-fatal): {e}")

        # Generate package with questions
        # Run in thread pool to avoid blocking the async event loop
        # (export_scorm_package is a CPU+IO intensive synchronous function)
        from services.scorm_exporter import export_scorm_package
        import asyncio
        
        # Detect backend URL for VLibras dictionary proxy
        # Priority: origin header (from frontend) > x-forwarded-host > frontend .env fallback
        origin = request.headers.get('origin', '')
        if origin:
            scorm_backend_url = origin
        else:
            fwd_host = request.headers.get('x-forwarded-host', '')
            scheme = request.headers.get('x-forwarded-proto', 'https')
            if fwd_host:
                scorm_backend_url = f"{scheme}://{fwd_host}"
            else:
                # Read from frontend .env as last resort
                scorm_backend_url = ''
                try:
                    env_path = Path(__file__).parent.parent / "frontend" / ".env"
                    for line in env_path.read_text().splitlines():
                        if line.startswith("REACT_APP_BACKEND_URL="):
                            scorm_backend_url = line.split("=", 1)[1].strip()
                            break
                except Exception:
                    pass
        
        zip_path = await asyncio.to_thread(
            export_scorm_package,
            project,
            str(PROJECTS_DIR),
            str(EXPORTS_DIR),
            questions=questions,
            tutor_config=tutor_settings,
            backend_url=scorm_backend_url
        )
        
        # Clean up old exports to prevent disk space exhaustion (keep last 24h)
        try:
            await asyncio.to_thread(_cleanup_old_exports, str(EXPORTS_DIR))
        except Exception as cleanup_err:
            logger.warning(f"Export cleanup failed (non-fatal): {cleanup_err}")
        
        # Persist to GridFS in background (don't block the response)
        export_filename = Path(zip_path).name
        background_tasks.add_task(save_export_to_gridfs, zip_path, export_filename)
        
        jobs[job_id]['status'] = 'completed'
        jobs[job_id]['progress'] = 100
        jobs[job_id]['message'] = 'SCORM package ready'
        jobs[job_id]['result'] = {
            'downloadUrl': f"/api/exports/{export_filename}"
        }
        
        return {
            "jobId": job_id,
            "downloadUrl": f"/api/exports/{export_filename}"
        }
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        logger.error(f"SCORM export error: {e}")
        logger.error(f"SCORM export traceback: {error_details}")
        jobs[job_id]['status'] = 'failed'
        jobs[job_id]['message'] = str(e)
        raise HTTPException(status_code=500, detail=f"SCORM export failed: {str(e)}")

# HTML Standalone Export

@api_router.post("/course/{project_id}/export-html")
async def export_html(project_id: str, background_tasks: BackgroundTasks):
    """Export project as standalone HTML file"""
    project_doc = await get_project_by_id(project_id)
    if not project_doc:
        raise HTTPException(status_code=404, detail="Project not found")
    
    try:
        from services.html_exporter import generate_standalone_html
        # Get assets directory
        assets_dir = str(PROJECTS_DIR / project_id / "assets")
        
        # Get base URL for external assets
        base_url = os.environ.get('REACT_APP_BACKEND_URL', '')
        
        # Collect question IDs from quiz elements
        question_ids = set()
        course_data = project_doc.get('course', {})
        for slide in course_data.get('slides', []):
            for element in slide.get('elements', []):
                if element.get('type') == 'quiz' and element.get('quizConfig'):
                    quiz_config = element.get('quizConfig')
                    if quiz_config and isinstance(quiz_config, dict):
                        question_ids.update(quiz_config.get('questionIds', []))
        
        # Load questions from database if there are quiz elements
        questions = []
        if question_ids:
            question_docs = await db.questions.find(
                {"id": {"$in": list(question_ids)}},
                {"_id": 0}
            ).to_list(500)
            questions = question_docs
            logger.info(f"Loaded {len(questions)} questions for HTML export")
        
        # Generate HTML with questions
        html_content = await generate_standalone_html(
            project_doc,
            assets_dir,
            base_url,
            questions=questions,
            backend_url=base_url
        )
        
        # Save HTML file
        project_name = project_doc.get('name', 'course')
        safe_name = re.sub(r'[^\w\s-]', '', project_name).replace(' ', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_name}_{timestamp}.html"
        
        html_path = EXPORTS_DIR / filename
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Persist to GridFS in background (don't block the response)
        background_tasks.add_task(save_export_to_gridfs, str(html_path), filename)
        
        return {
            "downloadUrl": f"/api/exports/{filename}",
            "filename": filename,
            "message": "HTML standalone file generated successfully"
        }
        
    except Exception as e:
        logger.error(f"HTML export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Video Export

@api_router.post("/course/{project_id}/export-video")
async def export_video_endpoint(project_id: str, request: Request, background_tasks: BackgroundTasks):
    """Export project as video (MP4 or WebM)"""
    from services.video_exporter import export_video as export_video_func, is_ffmpeg_available
    # Check if video export is available
    if not is_ffmpeg_available():
        raise HTTPException(
            status_code=503,
            detail="A exportação de vídeo não está disponível neste ambiente. FFmpeg não está instalado. Use exportação SCORM ou HTML como alternativa."
        )
    
    project_doc = await get_project_by_id(project_id)
    if not project_doc:
        raise HTTPException(status_code=404, detail="Project not found")

    body = await request.json()
    video_format = body.get('format', 'mp4')
    if video_format not in ('mp4', 'webm'):
        video_format = 'mp4'
    default_duration = float(body.get('default_duration', 5.0))

    # Create job for tracking
    job_id = str(uuid.uuid4())
    jobs[job_id] = {
        'id': job_id,
        'status': 'processing',
        'progress': 0,
        'message': 'Iniciando exportação de vídeo...',
        'result': None
    }

    async def run_export():
        try:
            def on_progress(progress, message):
                jobs[job_id]['progress'] = progress
                jobs[job_id]['message'] = message

            output_path = await export_video_func(
                project_doc,
                str(PROJECTS_DIR),
                str(STORAGE_DIR),
                str(EXPORTS_DIR),
                video_format=video_format,
                default_duration=default_duration,
                on_progress=on_progress
            )

            filename = Path(output_path).name
            # Persist to GridFS
            await save_export_to_gridfs(output_path, filename)
            jobs[job_id]['status'] = 'completed'
            jobs[job_id]['progress'] = 100
            jobs[job_id]['message'] = 'Vídeo exportado com sucesso!'
            jobs[job_id]['result'] = {
                'downloadUrl': f"/api/exports/{filename}",
                'filename': filename
            }
        except Exception as e:
            logger.error(f"Video export error: {e}")
            jobs[job_id]['status'] = 'failed'
            jobs[job_id]['message'] = f"Erro na exportação: {str(e)}"

    # Run in background
    asyncio.create_task(run_export())

    return {
        "jobId": job_id,
        "message": f"Exportação de vídeo {video_format.upper()} iniciada"
    }

# Static file serving

@api_router.get("/projects/{project_id}/assets/{filename}")
async def serve_asset(project_id: str, filename: str):
    """Serve project asset - falls back to MongoDB if local file is missing"""
    file_path = PROJECTS_DIR / project_id / "assets" / filename
    if file_path.exists():
        return FileResponse(file_path)
    
    # Fallback: try to restore from MongoDB (production ephemeral storage)
    try:
        from services.asset_store import retrieve_asset_async
        data, content_type = await retrieve_asset_async(db, project_id, filename)
        if data:
            # Restore to filesystem for future requests
            file_path.parent.mkdir(parents=True, exist_ok=True)
            with open(file_path, 'wb') as f:
                f.write(data)
            logger.info(f"Restored asset from MongoDB: {project_id}/{filename}")
            return FileResponse(file_path)
    except Exception as e:
        logger.warning(f"MongoDB asset fallback failed: {e}")
    
    raise HTTPException(status_code=404, detail="File not found")

@api_router.get("/exports/{filename}")
async def serve_export(filename: str, preview: str = None):
    """Serve exported files (SCORM zip or HTML) with forced download.
    Falls back to MongoDB GridFS if the file is not on local disk."""
    file_path = EXPORTS_DIR / filename
    
    # If not on local disk, try to restore from GridFS
    if not file_path.exists():
        restored = await get_export_from_gridfs(filename, str(file_path))
        if not restored:
            raise HTTPException(status_code=404, detail="File not found")
    
    # Determine media type based on file extension
    if filename.endswith('.html'):
        media_type = 'text/html'
    elif filename.endswith('.zip'):
        media_type = 'application/zip'
    elif filename.endswith('.mp4'):
        media_type = 'video/mp4'
    elif filename.endswith('.webm'):
        media_type = 'video/webm'
    else:
        media_type = 'application/octet-stream'
    
    # If preview mode, serve inline (for HTML files)
    if preview and filename.endswith('.html'):
        return FileResponse(
            file_path,
            media_type=media_type,
        )
    
    # Always force download for export files by setting filename
    # This adds Content-Disposition: attachment header
    return FileResponse(
        file_path,
        media_type=media_type,
        filename=filename,
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename}\""
        }
    )

@api_router.get("/assets/{filename}")
async def serve_global_asset(filename: str):
    """Serve global assets (AI generated images, etc.)"""
    file_path = STORAGE_DIR / "assets" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path)

# ============================================
# HeyGen API Endpoints
# ============================================

@api_router.get("/heygen/avatars")
async def list_heygen_avatars(limit: int = 200, gender: Optional[str] = None):
    """List available HeyGen avatars with optional gender filter"""
    if not HEYGEN_API_KEY:
        raise HTTPException(status_code=500, detail="HeyGen API key not configured")
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            response = await http_client.get(
                f"{HEYGEN_BASE_URL}/v2/avatars",
                headers=HEYGEN_HEADERS
            )
            
            if response.status_code != 200:
                logger.error(f"HeyGen avatars error: {response.status_code} - {response.text}")
                raise HTTPException(status_code=response.status_code, detail="Failed to fetch avatars from HeyGen")
            
            data = response.json()
            avatars = data.get("data", {}).get("avatars", [])
            
            # Filter by gender if specified
            if gender and gender.lower() != 'all':
                avatars = [a for a in avatars if a.get("gender", "").lower() == gender.lower()]
            
            # Format avatars for frontend
            formatted_avatars = []
            for avatar in avatars[:limit]:
                formatted_avatars.append({
                    "avatar_id": avatar.get("avatar_id"),
                    "avatar_name": avatar.get("avatar_name"),
                    "preview_image_url": avatar.get("preview_image_url"),
                    "preview_video_url": avatar.get("preview_video_url"),
                    "gender": avatar.get("gender"),
                })
            
            # Get unique genders for filter options
            all_genders = list(set(a.get("gender", "unknown") for a in data.get("data", {}).get("avatars", []) if a.get("gender")))
            
            return {
                "avatars": formatted_avatars, 
                "total": len(avatars),
                "available_genders": sorted(all_genders)
            }
    except httpx.RequestError as e:
        logger.error(f"HeyGen request error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to connect to HeyGen: {str(e)}")

@api_router.get("/heygen/voices")
async def list_heygen_voices(language: Optional[str] = None, gender: Optional[str] = None):
    """List available HeyGen voices with optional language and gender filters"""
    if not HEYGEN_API_KEY:
        raise HTTPException(status_code=500, detail="HeyGen API key not configured")
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            response = await http_client.get(
                f"{HEYGEN_BASE_URL}/v2/voices",
                headers=HEYGEN_HEADERS
            )
            
            if response.status_code != 200:
                logger.error(f"HeyGen voices error: {response.status_code} - {response.text}")
                raise HTTPException(status_code=response.status_code, detail="Failed to fetch voices from HeyGen")
            
            data = response.json()
            all_voices = data.get("data", {}).get("voices", [])
            voices = all_voices.copy()
            
            # Filter by language if specified
            if language and language.lower() != 'all':
                # Support specific language codes like "pt-BR", "pt-PT", "en-US"
                if '-' in language:
                    # Exact match for language code
                    voices = [v for v in voices if v.get("language", "").lower() == language.lower()]
                else:
                    # Partial match (e.g., "portuguese" matches "Portuguese (Brazil)")
                    voices = [v for v in voices if language.lower() in v.get("language", "").lower()]
            
            # Filter by gender if specified
            if gender and gender.lower() != 'all':
                voices = [v for v in voices if v.get("gender", "").lower() == gender.lower()]
            
            # Format voices for frontend
            formatted_voices = []
            for voice in voices:
                lang = voice.get("language", "")
                voice_name = voice.get("name", "")
                # Determine language code and country
                lang_code = ""
                country_flag = ""
                
                # Check if Portuguese and determine variant
                if "portuguese" in lang.lower():
                    # Check voice name for Brazil indicators
                    if "brazil" in voice_name.lower() or "brasil" in voice_name.lower():
                        lang_code = "pt-BR"
                        country_flag = "🇧🇷"
                    elif "portugal" in voice_name.lower():
                        lang_code = "pt-PT"
                        country_flag = "🇵🇹"
                    else:
                        # Default Portuguese to Brazil (most common in LatAm)
                        lang_code = "pt"
                        country_flag = "🇧🇷"
                elif "brazil" in lang.lower() or "brasileiro" in lang.lower():
                    lang_code = "pt-BR"
                    country_flag = "🇧🇷"
                elif "portugal" in lang.lower() or "português" in lang.lower():
                    lang_code = "pt-PT"
                    country_flag = "🇵🇹"
                elif "english" in lang.lower():
                    if "us" in lang.lower() or "american" in lang.lower():
                        lang_code = "en-US"
                        country_flag = "🇺🇸"
                    elif "uk" in lang.lower() or "british" in lang.lower():
                        lang_code = "en-GB"
                        country_flag = "🇬🇧"
                    else:
                        lang_code = "en"
                        country_flag = "🇺🇸"
                elif "spanish" in lang.lower() or "español" in lang.lower():
                    lang_code = "es"
                    country_flag = "🇪🇸"
                elif "french" in lang.lower():
                    lang_code = "fr"
                    country_flag = "🇫🇷"
                elif "german" in lang.lower():
                    lang_code = "de"
                    country_flag = "🇩🇪"
                elif "italian" in lang.lower():
                    lang_code = "it"
                    country_flag = "🇮🇹"
                
                formatted_voices.append({
                    "voice_id": voice.get("voice_id"),
                    "name": voice_name,
                    "language": lang,
                    "language_code": lang_code,
                    "country_flag": country_flag,
                    "gender": voice.get("gender"),
                    "preview_audio": voice.get("preview_audio"),
                    "support_pause": voice.get("support_pause", False),
                })
            
            # Get unique languages for filter options
            language_options = []
            seen_languages = set()
            for v in all_voices:
                lang = v.get("language", "")
                if lang and lang not in seen_languages:
                    seen_languages.add(lang)
                    # Create a simplified label
                    if "brazil" in lang.lower():
                        language_options.append({"value": lang, "label": "🇧🇷 Português (Brasil)", "code": "pt-BR"})
                    elif "portugal" in lang.lower():
                        language_options.append({"value": lang, "label": "🇵🇹 Português (Portugal)", "code": "pt-PT"})
                    elif "english" in lang.lower():
                        if "us" in lang.lower() or "american" in lang.lower():
                            language_options.append({"value": lang, "label": "🇺🇸 English (US)", "code": "en-US"})
                        elif "uk" in lang.lower() or "british" in lang.lower():
                            language_options.append({"value": lang, "label": "🇬🇧 English (UK)", "code": "en-GB"})
                        else:
                            language_options.append({"value": lang, "label": "🇺🇸 English", "code": "en"})
                    elif "spanish" in lang.lower():
                        language_options.append({"value": lang, "label": "🇪🇸 Español", "code": "es"})
                    elif "french" in lang.lower():
                        language_options.append({"value": lang, "label": "🇫🇷 Français", "code": "fr"})
                    elif "german" in lang.lower():
                        language_options.append({"value": lang, "label": "🇩🇪 Deutsch", "code": "de"})
                    elif "italian" in lang.lower():
                        language_options.append({"value": lang, "label": "🇮🇹 Italiano", "code": "it"})
                    else:
                        language_options.append({"value": lang, "label": lang, "code": ""})
            
            # Sort by label
            language_options.sort(key=lambda x: x["label"])
            
            # Get unique genders
            available_genders = list(set(v.get("gender", "unknown") for v in all_voices if v.get("gender")))
            
            return {
                "voices": formatted_voices,
                "total": len(voices),
                "available_languages": language_options,
                "available_genders": sorted(available_genders)
            }
    except httpx.RequestError as e:
        logger.error(f"HeyGen request error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to connect to HeyGen: {str(e)}")

from pydantic import BaseModel, ConfigDict

class HeyGenVideoRequest(BaseModel):
    avatar_id: str
    voice_id: str
    script: str
    title: Optional[str] = "Generated Video"
    aspect_ratio: Optional[str] = "16:9"
    transparent_background: Optional[bool] = True  # Default to transparent
    project_id: Optional[str] = None  # Associate video with a project

@api_router.post("/heygen/generate-video")
async def generate_heygen_video(request: HeyGenVideoRequest):
    """Generate a video using HeyGen API"""
    if not HEYGEN_API_KEY:
        raise HTTPException(status_code=500, detail="HeyGen API key not configured")
    
    if len(request.script) > 5000:
        raise HTTPException(status_code=400, detail="Script exceeds 5000 character limit")
    
    try:
        async with httpx.AsyncClient(timeout=60.0) as http_client:
            
            # For transparent background, try the WebM endpoint first (v1/video.webm)
            if request.transparent_background:
                payload = {
                    "avatar_pose_id": request.avatar_id,
                    "avatar_style": "normal",
                    "input_text": request.script,
                    "voice_id": request.voice_id,
                    "dimension": {
                        "width": 1280 if request.aspect_ratio == "16:9" else 720,
                        "height": 720 if request.aspect_ratio == "16:9" else 1280
                    }
                }
                
                response = await http_client.post(
                    f"{HEYGEN_BASE_URL}/v1/video.webm",
                    headers=HEYGEN_HEADERS,
                    json=payload
                )
                
                # If WebM endpoint fails (avatar not supported), fall back to standard endpoint
                if response.status_code != 200:
                    error_data = response.json()
                    error_code = error_data.get("data", {}).get("error", {}).get("code", "")
                    
                    # Check if it's an avatar compatibility issue
                    if "AVATAR_NOT_FOUND" in str(error_code) or "avatar" in str(error_data).lower():
                        logger.warning("Avatar not compatible with WebM, falling back to standard video")
                        # Fall back to standard endpoint without transparent background
                        request.transparent_background = False
                    else:
                        logger.error(f"HeyGen WebM error: {response.status_code} - {response.text}")
                        error_msg = error_data.get("error", {}).get("message", "")
                        if not error_msg:
                            error_msg = error_data.get("data", {}).get("error", {}).get("message", response.text)
                        raise HTTPException(status_code=response.status_code, detail=f"HeyGen error: {error_msg}")
            
            # Standard video (either requested or fallback from WebM)
            if not request.transparent_background:
                payload = {
                    "video_inputs": [
                        {
                            "character": {
                                "type": "avatar",
                                "avatar_id": request.avatar_id,
                                "avatar_style": "normal"
                            },
                            "voice": {
                                "type": "text",
                                "input_text": request.script,
                                "voice_id": request.voice_id
                            }
                        }
                    ],
                    "dimension": {
                        "width": 1280 if request.aspect_ratio == "16:9" else 720,
                        "height": 720 if request.aspect_ratio == "16:9" else 1280
                    },
                    "title": request.title
                }
                
                response = await http_client.post(
                    f"{HEYGEN_BASE_URL}/v2/video/generate",
                    headers=HEYGEN_HEADERS,
                    json=payload
                )
            
            if response.status_code != 200:
                logger.error(f"HeyGen generate error: {response.status_code} - {response.text}")
                error_data = response.json()
                error_msg = error_data.get("error", {}).get("message", "")
                if not error_msg:
                    error_msg = error_data.get("data", {}).get("error", {}).get("message", response.text)
                raise HTTPException(status_code=response.status_code, detail=f"HeyGen error: {error_msg}")
            
            data = response.json()
            video_id = data.get("data", {}).get("video_id")
            
            if not video_id:
                raise HTTPException(status_code=500, detail="No video ID returned from HeyGen")
            
            # Store video generation request in database
            await db.heygen_videos.insert_one({
                "video_id": video_id,
                "avatar_id": request.avatar_id,
                "voice_id": request.voice_id,
                "script": request.script,
                "title": request.title,
                "status": "processing",
                "transparent": request.transparent_background,
                "project_id": request.project_id,
                "created_at": now_utc()
            })
            
            return {
                "video_id": video_id,
                "status": "processing",
                "message": "Video generation started. Poll status endpoint for updates."
            }
    except httpx.RequestError as e:
        logger.error(f"HeyGen request error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to connect to HeyGen: {str(e)}")

# AI Script Generation (lazy imports for fast startup)

class GenerateScriptRequest(BaseModel):
    topic: str
    style: Optional[str] = "educational"  # educational, conversational, formal, friendly
    duration: Optional[str] = "medium"  # short (30s), medium (1-2min), long (3-5min)
    language: Optional[str] = "português brasileiro"

@api_router.post("/ai/generate-script")
async def generate_ai_script(request: GenerateScriptRequest):
    """Generate a video script using AI"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    emergent_key = os.environ.get('EMERGENT_LLM_KEY', '')
    
    if not emergent_key:
        raise HTTPException(status_code=500, detail="AI key not configured")
    
    # Define duration guidelines
    duration_guide = {
        "short": "30 segundos a 1 minuto (aproximadamente 75-150 palavras)",
        "medium": "1 a 2 minutos (aproximadamente 150-300 palavras)",
        "long": "3 a 5 minutos (aproximadamente 450-750 palavras)"
    }
    
    style_guide = {
        "educational": "tom educativo e didático, explicando conceitos de forma clara",
        "conversational": "tom conversacional e descontraído, como se estivesse falando com um amigo",
        "formal": "tom formal e profissional, adequado para ambientes corporativos",
        "friendly": "tom amigável e acolhedor, criando conexão com o espectador"
    }
    
    try:
        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"script-gen-{uuid.uuid4()}",
            system_message=f"""Você é um roteirista profissional especializado em criar scripts para vídeos com avatares de IA.

Suas diretrizes:
1. Escreva em {request.language}
2. Use {style_guide.get(request.style, style_guide['educational'])}
3. O script deve ter {duration_guide.get(request.duration, duration_guide['medium'])}
4. Escreva de forma natural e fluida, como se fosse uma pessoa real falando
5. Use pausas naturais (vírgulas, pontos) para dar ritmo ao texto
6. Evite jargões técnicos complexos, a menos que sejam explicados
7. Comece com uma saudação ou gancho que prenda a atenção
8. Termine com uma conclusão clara ou call-to-action

IMPORTANTE: Retorne APENAS o script, sem títulos, numeração de cenas ou instruções de direção."""
        ).with_model("openai", "gpt-4o")
        
        user_message = UserMessage(
            text=f"Crie um script de vídeo sobre o seguinte tema:\n\n{request.topic}"
        )
        
        response = await chat.send_message(user_message)
        
        return {
            "script": response,
            "topic": request.topic,
            "style": request.style,
            "duration": request.duration,
            "language": request.language
        }
    except Exception as e:
        logger.error(f"AI script generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate script: {str(e)}")

class GenerateNarrationRequest(BaseModel):
    slide_content: Optional[str] = ""
    style: Optional[str] = "educational"
    language: Optional[str] = "português brasileiro"

@api_router.post("/projects/{project_id}/slides/{slide_id}/generate-narration")
async def generate_slide_narration(project_id: str, slide_id: str, request: GenerateNarrationRequest):
    """Generate 3 narration text options for a slide using Gemini 3 with vision (OCR for images)"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContent
    emergent_key = os.environ.get('EMERGENT_LLM_KEY', '')
    if not emergent_key:
        raise HTTPException(status_code=500, detail="AI key not configured")

    project = await get_project_by_id(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    course = project.get('course', {})
    slides = course.get('slides', [])
    slide = next((s for s in slides if s.get('id') == slide_id), None)
    if not slide:
        raise HTTPException(status_code=404, detail="Slide not found")

    # Collect text content and images from the slide
    text_parts = []
    image_files = []  # list of FileContent for Gemini vision

    # Check for backgroundImage (PPT-imported slides store full slide as image)
    bg_image = slide.get('backgroundImage', '')
    if bg_image and bg_image.startswith('/api/projects/'):
        # Extract file path from URL: /api/projects/{id}/assets/{filename}
        parts = bg_image.split('/assets/')
        if len(parts) == 2:
            local_path = PROJECTS_DIR / project_id / "assets" / parts[1]
            if local_path.exists():
                try:
                    img_data = local_path.read_bytes()
                    ext = local_path.suffix.lower()
                    mime = 'image/png' if ext == '.png' else 'image/jpeg'
                    image_files.append(FileContent(
                        content_type=mime,
                        file_content_base64=base64.b64encode(img_data).decode('utf-8')
                    ))
                    logger.info(f"Loaded background image for vision: {local_path.name}")
                except Exception as e:
                    logger.warning(f"Failed to load background image: {e}")

    # Extract content from elements
    for element in slide.get('elements', []):
        el_type = element.get('type', '')
        # Check both 'content' and 'htmlContent' fields (htmlContent is used by rich text editor)
        raw_content = element.get('content') or element.get('htmlContent') or ''
        if el_type in ('text', 'html') and raw_content:
            # Extract text from HTML
            clean = re.sub(r'<[^>]+>', '', raw_content)
            if clean.strip():
                text_parts.append(clean.strip())
            # Also check for inline images in htmlContent (RTF editor embeds images)
            img_matches = re.findall(r'src="(/api/[^"]+)"', raw_content)
            for img_src in img_matches:
                if img_src.startswith('/api/projects/'):
                    parts = img_src.split('/assets/')
                    if len(parts) == 2:
                        local_path = PROJECTS_DIR / project_id / "assets" / parts[1]
                        if local_path.exists():
                            try:
                                img_data = local_path.read_bytes()
                                ext = local_path.suffix.lower()
                                mime = 'image/png' if ext == '.png' else 'image/jpeg' if ext in ('.jpg', '.jpeg') else 'image/webp'
                                image_files.append(FileContent(
                                    content_type=mime,
                                    file_content_base64=base64.b64encode(img_data).decode('utf-8')
                                ))
                                logger.info(f"Loaded inline image for vision: {local_path.name}")
                            except Exception as e:
                                logger.warning(f"Failed to load inline image: {e}")
                elif img_src.startswith('/api/assets/'):
                    # Global assets path
                    asset_name = img_src.split('/api/assets/')[-1]
                    local_path = STORAGE_DIR / "assets" / asset_name
                    if local_path.exists():
                        try:
                            img_data = local_path.read_bytes()
                            ext = local_path.suffix.lower()
                            mime = 'image/png' if ext == '.png' else 'image/jpeg' if ext in ('.jpg', '.jpeg') else 'image/webp'
                            image_files.append(FileContent(
                                content_type=mime,
                                file_content_base64=base64.b64encode(img_data).decode('utf-8')
                            ))
                            logger.info(f"Loaded global asset image for vision: {asset_name}")
                        except Exception as e:
                            logger.warning(f"Failed to load global asset image: {e}")
        elif el_type == 'image' and element.get('src'):
            src = element['src']
            if src.startswith('/api/projects/'):
                parts = src.split('/assets/')
                if len(parts) == 2:
                    local_path = PROJECTS_DIR / project_id / "assets" / parts[1]
                    if local_path.exists():
                        try:
                            img_data = local_path.read_bytes()
                            ext = local_path.suffix.lower()
                            mime = 'image/png' if ext == '.png' else 'image/jpeg' if ext in ('.jpg', '.jpeg') else 'image/webp'
                            image_files.append(FileContent(
                                content_type=mime,
                                file_content_base64=base64.b64encode(img_data).decode('utf-8')
                            ))
                            logger.info(f"Loaded element image for vision: {local_path.name}")
                        except Exception as e:
                            logger.warning(f"Failed to load element image: {e}")
        elif el_type == 'quiz':
            text_parts.append("[Quiz/Atividade interativa presente no slide]")
        elif el_type == 'video':
            text_parts.append("[Vídeo presente no slide]")

    slide_text = "\n".join(text_parts) if text_parts else ""
    slide_title = slide.get('title', '')
    has_images = len(image_files) > 0

    style_guide = {
        "educational": "educativo e didático, explicando conceitos de forma clara e objetiva",
        "conversational": "conversacional e descontraído, como se estivesse falando com um amigo",
        "formal": "formal e profissional, adequado para ambientes corporativos",
        "friendly": "amigável e acolhedor, criando conexão com o espectador"
    }

    try:
        system_msg = f"""Você é um especialista em criar textos de narração para slides de cursos e apresentações.

Suas diretrizes:
1. Escreva em {request.language}
2. Use tom {style_guide.get(request.style, style_guide['educational'])}
3. O texto deve ser adequado para narração em voz alta (TTS)
4. Escreva de forma natural, fluida e envolvente
5. Use pausas naturais (vírgulas, pontos) para dar ritmo
6. O texto deve complementar e explicar o conteúdo visual do slide
7. Cada opção deve ter entre 2 e 5 frases (ideal para 20-60 segundos de narração)
8. Cada opção deve ter uma abordagem ligeiramente diferente"""

        if has_images:
            system_msg += """
9. IMPORTANTE: Analise atentamente as imagens do slide. Leia todo o texto visível nas imagens (OCR).
10. Use o conteúdo visual e textual das imagens como base principal para a narração."""

        system_msg += """

FORMATO DE RESPOSTA OBRIGATÓRIO:
Retorne exatamente 3 opções, separadas por "---". Cada opção deve conter APENAS o texto de narração, sem numeração, títulos ou marcadores. Exemplo:

Texto da primeira opção aqui...
---
Texto da segunda opção aqui...
---
Texto da terceira opção aqui..."""

        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"narration-gen-{uuid.uuid4()}",
            system_message=system_msg
        ).with_model("gemini", "gemini-3-flash-preview")

        prompt = "Crie 3 opções de texto de narração para o seguinte slide:"
        if slide_title:
            prompt += f"\n\nTítulo do slide: {slide_title}"
        if has_images:
            prompt += "\n\nAs imagens do slide estão anexadas. Leia o conteúdo visual e textual delas para criar a narração."
        if slide_text:
            prompt += f"\n\nTexto extraído do slide:\n{slide_text}"
        if request.slide_content:
            prompt += f"\n\nContexto adicional do usuário: {request.slide_content}"

        user_message = UserMessage(
            text=prompt,
            file_contents=image_files if image_files else None
        )
        response = await chat.send_message(user_message)

        # Parse the 3 options
        options = [opt.strip() for opt in response.split("---") if opt.strip()]

        # Ensure we have exactly 3 options
        if len(options) < 3:
            while len(options) < 3:
                options.append(options[-1] if options else "Narração não disponível.")
        options = options[:3]

        return {
            "options": options,
            "slide_id": slide_id,
            "style": request.style
        }
    except Exception as e:
        logger.error(f"AI narration generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Falha ao gerar narração: {str(e)}")



@api_router.get("/heygen/video-status/{video_id}")
async def get_heygen_video_status(video_id: str):
    """Check the status of a HeyGen video generation"""
    if not HEYGEN_API_KEY:
        raise HTTPException(status_code=500, detail="HeyGen API key not configured")
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            response = await http_client.get(
                f"{HEYGEN_BASE_URL}/v1/video_status.get",
                headers=HEYGEN_HEADERS,
                params={"video_id": video_id}
            )
            
            if response.status_code != 200:
                logger.error(f"HeyGen status error: {response.status_code} - {response.text}")
                raise HTTPException(status_code=response.status_code, detail="Failed to get video status")
            
            data = response.json()
            video_data = data.get("data", {})
            
            status = video_data.get("status", "unknown")
            video_url = video_data.get("video_url")
            thumbnail_url = video_data.get("thumbnail_url")
            duration = video_data.get("duration")
            
            # Update database
            await db.heygen_videos.update_one(
                {"video_id": video_id},
                {"$set": {
                    "status": status,
                    "video_url": video_url,
                    "thumbnail_url": thumbnail_url,
                    "duration": duration,
                    "updated_at": now_utc()
                }}
            )
            
            return {
                "video_id": video_id,
                "status": status,
                "video_url": video_url,
                "thumbnail_url": thumbnail_url,
                "duration": duration
            }
    except httpx.RequestError as e:
        logger.error(f"HeyGen request error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to connect to HeyGen: {str(e)}")

@api_router.get("/heygen/videos")
async def list_heygen_videos(project_id: Optional[str] = None):
    """List all generated HeyGen videos, optionally filtered by project"""
    query = {}
    if project_id:
        query["project_id"] = project_id
    
    videos = await db.heygen_videos.find(query).sort("created_at", -1).to_list(100)
    
    formatted_videos = []
    for video in videos:
        formatted_videos.append({
            "video_id": video.get("video_id"),
            "title": video.get("title"),
            "status": video.get("status"),
            "video_url": video.get("video_url"),
            "thumbnail_url": video.get("thumbnail_url"),
            "duration": video.get("duration"),
            "script": video.get("script"),
            "avatar_id": video.get("avatar_id"),
            "voice_id": video.get("voice_id"),
            "project_id": video.get("project_id"),
            "transparent": video.get("transparent"),
            "created_at": video.get("created_at").isoformat() if video.get("created_at") else None
        })
    
    return {"videos": formatted_videos}


@api_router.get("/heygen/videos/{video_id}/refresh")
async def refresh_heygen_video_status(video_id: str):
    """Refresh video status from HeyGen API and update database"""
    if not HEYGEN_API_KEY:
        raise HTTPException(status_code=500, detail="HeyGen API key not configured")
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            response = await http_client.get(
                f"{HEYGEN_BASE_URL}/v1/video_status.get?video_id={video_id}",
                headers=HEYGEN_HEADERS
            )
            
            if response.status_code != 200:
                logger.error(f"HeyGen status error: {response.status_code} - {response.text}")
                raise HTTPException(status_code=response.status_code, detail="Failed to fetch video status")
            
            data = response.json()
            video_data = data.get("data", {})
            status = video_data.get("status", "unknown")
            video_url = video_data.get("video_url")
            thumbnail_url = video_data.get("thumbnail_url")
            duration = video_data.get("duration")
            
            # Update database with fresh status
            update_data = {"status": status}
            if video_url:
                update_data["video_url"] = video_url
            if thumbnail_url:
                update_data["thumbnail_url"] = thumbnail_url
            if duration:
                update_data["duration"] = duration
            
            await db.heygen_videos.update_one(
                {"video_id": video_id},
                {"$set": update_data}
            )
            
            # Get full video data from database
            video_doc = await db.heygen_videos.find_one({"video_id": video_id})
            
            return {
                "video_id": video_id,
                "status": status,
                "video_url": video_url,
                "thumbnail_url": thumbnail_url,
                "duration": duration,
                "title": video_doc.get("title") if video_doc else None,
                "script": video_doc.get("script") if video_doc else None,
                "project_id": video_doc.get("project_id") if video_doc else None,
                "created_at": video_doc.get("created_at").isoformat() if video_doc and video_doc.get("created_at") else None
            }
    except httpx.RequestError as e:
        logger.error(f"HeyGen request error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to connect to HeyGen: {str(e)}")


@api_router.delete("/heygen/videos/{video_id}")
async def delete_heygen_video(video_id: str):
    """Delete a video from the library (database only, not from HeyGen)"""
    # Check if video exists
    video_doc = await db.heygen_videos.find_one({"video_id": video_id})
    if not video_doc:
        raise HTTPException(status_code=404, detail="Video not found")
    
    # Delete from database
    result = await db.heygen_videos.delete_one({"video_id": video_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=500, detail="Failed to delete video")
    
    logger.info(f"Deleted video {video_id} from library")
    
    return {
        "message": "Video deleted successfully",
        "video_id": video_id
    }


# HeyGen Credits/Quota Check
@api_router.get("/heygen/credits")
async def get_heygen_credits(force_refresh: bool = False):
    """Check remaining HeyGen API credits/quota (with caching)"""
    if not HEYGEN_API_KEY:
        raise HTTPException(status_code=500, detail="HeyGen API key not configured")
    
    # Check cache first (unless force refresh requested)
    if not force_refresh and heygen_credits_cache["data"] is not None:
        cache_age = (datetime.now(timezone.utc) - heygen_credits_cache["timestamp"]).total_seconds()
        if cache_age < heygen_credits_cache["ttl"]:
            logger.info(f"Returning cached HeyGen credits (age: {cache_age:.1f}s)")
            return heygen_credits_cache["data"]
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            response = await http_client.get(
                f"{HEYGEN_BASE_URL}/v2/user/remaining_quota",
                headers=HEYGEN_HEADERS
            )
            
            if response.status_code != 200:
                logger.error(f"HeyGen credits error: {response.status_code} - {response.text}")
                raise HTTPException(status_code=response.status_code, detail="Failed to fetch credits from HeyGen")
            
            data = response.json()
            quota_data = data.get("data", {})
            
            result = {
                "remaining_quota": quota_data.get("remaining_quota", 0),
                "used_quota": quota_data.get("used_quota"),
                "plan": quota_data.get("plan"),
                "has_credits": quota_data.get("remaining_quota", 0) > 0
            }
            
            # Update cache
            heygen_credits_cache["data"] = result
            heygen_credits_cache["timestamp"] = datetime.now(timezone.utc)
            
            return result
    except httpx.RequestError as e:
        logger.error(f"HeyGen credits request error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to connect to HeyGen: {str(e)}")

# HeyGen Webhook Endpoint
class HeyGenWebhookPayload(BaseModel):
    model_config = ConfigDict(extra="allow")
    event_type: Optional[str] = None
    video_id: Optional[str] = None
    status: Optional[str] = None
    video_url: Optional[str] = None
    thumbnail_url: Optional[str] = None
    duration: Optional[float] = None

@api_router.post("/heygen/webhook")
async def heygen_webhook(request: Request):
    """Receive webhook notifications from HeyGen when video processing completes"""
    try:
        # Get raw body for signature verification if needed
        body = await request.json()
        logger.info(f"HeyGen webhook received: {body}")
        
        # Extract relevant data
        video_id = body.get("video_id") or body.get("data", {}).get("video_id")
        status = body.get("status") or body.get("data", {}).get("status")
        video_url = body.get("video_url") or body.get("data", {}).get("video_url")
        thumbnail_url = body.get("thumbnail_url") or body.get("data", {}).get("thumbnail_url")
        duration = body.get("duration") or body.get("data", {}).get("duration")
        
        if video_id:
            # Update video status in database
            update_data = {
                "webhook_received": True,
                "webhook_received_at": now_utc(),
                "updated_at": now_utc()
            }
            
            if status:
                update_data["status"] = status
            if video_url:
                update_data["video_url"] = video_url
            if thumbnail_url:
                update_data["thumbnail_url"] = thumbnail_url
            if duration:
                update_data["duration"] = duration
            
            result = await db.heygen_videos.update_one(
                {"video_id": video_id},
                {"$set": update_data}
            )
            
            logger.info(f"HeyGen webhook processed: video_id={video_id}, status={status}, matched={result.matched_count}")
            
            # Notify SSE subscribers waiting for this video
            if video_id in heygen_sse_subscribers:
                event_data = {
                    "event": "video_update",
                    "video_id": video_id,
                    "status": status,
                    "video_url": video_url,
                    "thumbnail_url": thumbnail_url,
                    "duration": duration
                }
                for queue in heygen_sse_subscribers[video_id]:
                    try:
                        queue.put_nowait(event_data)
                    except asyncio.QueueFull:
                        pass
                logger.info(f"Notified {len(heygen_sse_subscribers[video_id])} SSE subscribers for video {video_id}")
            
            # Invalidate credits cache after video generation
            heygen_credits_cache["data"] = None
            heygen_credits_cache["timestamp"] = None
            
            return {
                "success": True,
                "video_id": video_id,
                "status": status,
                "message": "Webhook processed successfully"
            }
        else:
            logger.warning(f"HeyGen webhook received without video_id: {body}")
            return {"success": True, "message": "Webhook received but no video_id found"}
            
    except Exception as e:
        logger.error(f"HeyGen webhook error: {e}")
        # Always return 200 to acknowledge receipt (avoid retries)
        return {"success": False, "error": str(e)}

# HeyGen Webhook Configuration Helper
@api_router.get("/heygen/webhook-url")
async def get_heygen_webhook_url(request: Request):
    """Get the webhook URL to configure in HeyGen dashboard"""
    # Get the base URL from the request or environment
    base_url = os.environ.get('REACT_APP_BACKEND_URL', str(request.base_url).rstrip('/'))
    webhook_url = f"{base_url}/api/heygen/webhook"
    
    return {
        "webhook_url": webhook_url,
        "instructions": [
            "1. Acesse o dashboard da HeyGen",
            "2. Vá em Settings > Webhooks",
            "3. Clique em 'Add Webhook Endpoint'",
            "4. Cole a URL acima no campo 'Endpoint URL'",
            "5. Selecione os eventos: 'video.completed', 'video.failed'",
            "6. Salve as configurações"
        ]
    }

# SSE endpoint for real-time video status updates
@api_router.get("/heygen/video-events/{video_id}")
async def heygen_video_events(video_id: str, request: Request):
    """Server-Sent Events stream for real-time video status updates.
    The frontend connects to this endpoint to receive webhook notifications without polling."""
    
    async def event_generator():
        queue = asyncio.Queue(maxsize=10)
        
        # Register subscriber
        if video_id not in heygen_sse_subscribers:
            heygen_sse_subscribers[video_id] = []
        heygen_sse_subscribers[video_id].append(queue)
        logger.info(f"SSE subscriber connected for video {video_id}")
        
        try:
            # Send initial connection confirmation
            yield f"data: {json.dumps({'event': 'connected', 'video_id': video_id})}\n\n"
            
            # Check current status from database immediately
            video_doc = await db.heygen_videos.find_one({"video_id": video_id})
            if video_doc:
                current_status = video_doc.get("status", "processing")
                current_url = video_doc.get("video_url")
                yield f"data: {json.dumps({'event': 'current_status', 'video_id': video_id, 'status': current_status, 'video_url': current_url})}\n\n"
                
                # If already completed, close stream
                if current_status in ["completed", "failed", "error"]:
                    yield f"data: {json.dumps({'event': 'final', 'video_id': video_id, 'status': current_status, 'video_url': current_url})}\n\n"
                    return
            
            # Wait for updates from webhook
            timeout_seconds = 900  # 15 minutes max
            start_time = datetime.now(timezone.utc)
            last_api_check = start_time
            api_check_interval = 10  # Check API every 10 seconds if no webhook
            
            while True:
                # Check if client disconnected
                if await request.is_disconnected():
                    logger.info(f"SSE client disconnected for video {video_id}")
                    break
                
                # Check timeout
                elapsed = (datetime.now(timezone.utc) - start_time).total_seconds()
                if elapsed > timeout_seconds:
                    yield f"data: {json.dumps({'event': 'timeout', 'video_id': video_id})}\n\n"
                    break
                
                try:
                    # Wait for event with timeout
                    event_data = await asyncio.wait_for(queue.get(), timeout=5.0)
                    yield f"data: {json.dumps(event_data)}\n\n"
                    
                    # If video is done, close the stream
                    if event_data.get("status") in ["completed", "failed", "error"]:
                        break
                        
                except asyncio.TimeoutError:
                    # No webhook received, check API directly as fallback
                    time_since_last_check = (datetime.now(timezone.utc) - last_api_check).total_seconds()
                    
                    if time_since_last_check >= api_check_interval:
                        last_api_check = datetime.now(timezone.utc)
                        
                        # Check status from HeyGen API directly
                        try:
                            async with httpx.AsyncClient() as client:
                                headers = {"X-Api-Key": os.getenv("HEYGEN_API_KEY", "")}
                                resp = await client.get(
                                    f"https://api.heygen.com/v1/video_status.get?video_id={video_id}",
                                    headers=headers,
                                    timeout=10.0
                                )
                                if resp.status_code == 200:
                                    api_data = resp.json().get("data", {})
                                    api_status = api_data.get("status", "processing")
                                    api_video_url = api_data.get("video_url")
                                    
                                    # Update database
                                    if api_status in ["completed", "failed", "error"]:
                                        await db.heygen_videos.update_one(
                                            {"video_id": video_id},
                                            {"$set": {
                                                "status": api_status,
                                                "video_url": api_video_url,
                                                "api_checked_at": now_utc()
                                            }}
                                        )
                                        
                                        yield f"data: {json.dumps({'event': 'status_update', 'video_id': video_id, 'status': api_status, 'video_url': api_video_url})}\n\n"
                                        break
                                    else:
                                        # Send progress update
                                        yield f"data: {json.dumps({'event': 'ping', 'elapsed': int(elapsed), 'status': api_status})}\n\n"
                        except Exception as e:
                            logger.error(f"Error checking HeyGen API: {e}")
                            # Send keepalive ping
                            yield f"data: {json.dumps({'event': 'ping', 'elapsed': int(elapsed)})}\n\n"
                    else:
                        # Just send keepalive
                        yield f"data: {json.dumps({'event': 'ping', 'elapsed': int(elapsed)})}\n\n"
                    
        finally:
            # Unregister subscriber
            if video_id in heygen_sse_subscribers:
                try:
                    heygen_sse_subscribers[video_id].remove(queue)
                    if not heygen_sse_subscribers[video_id]:
                        del heygen_sse_subscribers[video_id]
                except ValueError:
                    pass
            logger.info(f"SSE subscriber disconnected for video {video_id}")
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )

# ============================================
# AI Text Generation Endpoint
# ============================================
class AITextGenerateRequest(BaseModel):
    prompt: str
    context: Optional[str] = None
    format: str = "html"  # html, plain, markdown

@api_router.post("/ai/generate-text")
async def generate_text_with_ai(request: AITextGenerateRequest):
    """Generate formatted text content using AI (GPT-4o)"""
    
    emergent_key = os.environ.get('EMERGENT_LLM_KEY')
    if not emergent_key:
        raise HTTPException(status_code=500, detail="AI API key not configured")
    
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        # System message for text generation
        system_message = """Você é um assistente especializado em criar conteúdo educacional formatado em HTML.

Regras:
1. SEMPRE responda em português brasileiro
2. Use tags HTML para formatação: <h1>, <h2>, <h3>, <p>, <strong>, <em>, <ul>, <ol>, <li>, <table>, <tr>, <td>, <th>
3. Crie conteúdo bem estruturado com títulos, parágrafos e listas quando apropriado
4. Use tabelas para dados comparativos ou estruturados
5. Seja conciso mas informativo
6. NÃO inclua tags <html>, <head> ou <body> - apenas o conteúdo interno
7. NÃO use markdown, apenas HTML puro

Exemplo de resposta formatada:
<h2>Título do Tópico</h2>
<p>Parágrafo introdutório explicando o conceito.</p>
<h3>Subtópico</h3>
<ul>
  <li><strong>Item 1:</strong> Descrição do item</li>
  <li><strong>Item 2:</strong> Descrição do item</li>
</ul>"""

        # Initialize chat with GPT-4o
        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"text-gen-{uuid.uuid4()}",
            system_message=system_message
        ).with_model("openai", "gpt-4o")
        
        # Build the prompt
        full_prompt = f"Gere conteúdo formatado sobre: {request.prompt}"
        if request.context:
            full_prompt += f"\n\nContexto adicional: {request.context}"
        
        # Send message and get response
        user_message = UserMessage(text=full_prompt)
        response = await chat.send_message(user_message)
        
        logger.info(f"AI text generated successfully for prompt: {request.prompt[:50]}...")
        
        return {
            "success": True,
            "content": response,
            "format": request.format
        }
        
    except ImportError:
        logger.error("emergentintegrations library not installed")
        raise HTTPException(status_code=500, detail="AI integration library not available")
    except Exception as e:
        logger.error(f"AI text generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate text: {str(e)}")


# AI Image Generation Endpoint
class AIImageGenerateRequest(BaseModel):
    prompt: str
    size: str = "1024x1024"  # 1024x1024, 1792x1024, 1024x1792

@api_router.post("/ai/generate-image")
async def generate_image_with_ai(request: AIImageGenerateRequest):
    """Generate image using AI (GPT Image 1) with optimization"""
    import base64
    from PIL import Image
    import io
    
    emergent_key = os.environ.get('EMERGENT_LLM_KEY')
    if not emergent_key:
        raise HTTPException(status_code=500, detail="AI API key not configured")
    
    try:
        from emergentintegrations.llm.openai.image_generation import OpenAIImageGeneration
        
        logger.info(f"Generating image with prompt: {request.prompt[:50]}...")
        
        # Initialize the image generator
        image_gen = OpenAIImageGeneration(api_key=emergent_key)
        
        # Generate the image
        images = await image_gen.generate_images(
            prompt=request.prompt,
            model="gpt-image-1",
            number_of_images=1
        )
        
        if not images or len(images) == 0:
            raise HTTPException(status_code=500, detail="No image was generated")
        
        # Optimize the image - convert to JPEG with compression
        original_image = Image.open(io.BytesIO(images[0]))
        
        # Convert RGBA to RGB if necessary (JPEG doesn't support transparency)
        if original_image.mode in ('RGBA', 'LA', 'P'):
            # Create white background
            background = Image.new('RGB', original_image.size, (255, 255, 255))
            if original_image.mode == 'P':
                original_image = original_image.convert('RGBA')
            background.paste(original_image, mask=original_image.split()[-1] if original_image.mode == 'RGBA' else None)
            original_image = background
        elif original_image.mode != 'RGB':
            original_image = original_image.convert('RGB')
        
        # Resize if image is too large (max 1200px on longest side for web)
        max_size = 1200
        if max(original_image.size) > max_size:
            ratio = max_size / max(original_image.size)
            new_size = (int(original_image.width * ratio), int(original_image.height * ratio))
            original_image = original_image.resize(new_size, Image.Resampling.LANCZOS)
            logger.info(f"Resized image from original to {new_size}")
        
        # Save as optimized JPEG
        optimized_buffer = io.BytesIO()
        original_image.save(optimized_buffer, format='JPEG', quality=80, optimize=True)
        optimized_data = optimized_buffer.getvalue()
        
        # Log size comparison
        original_size = len(images[0])
        optimized_size = len(optimized_data)
        compression_ratio = (1 - optimized_size / original_size) * 100
        logger.info(f"Image optimized: {original_size/1024:.1f}KB -> {optimized_size/1024:.1f}KB ({compression_ratio:.1f}% reduction)")
        
        # Convert to base64
        image_base64 = base64.b64encode(optimized_data).decode('utf-8')
        
        # Save to storage and return URL
        image_id = str(uuid.uuid4())
        image_filename = f"{image_id}.jpg"  # Changed to .jpg
        
        # Use the general storage assets directory
        assets_dir = STORAGE_DIR / "assets"
        assets_dir.mkdir(exist_ok=True)
        image_path = assets_dir / image_filename
        
        with open(image_path, "wb") as f:
            f.write(optimized_data)
        
        logger.info(f"Image generated successfully: {image_filename}")
        
        return {
            "success": True,
            "imageUrl": f"/api/assets/{image_filename}",
            "imageBase64": f"data:image/jpeg;base64,{image_base64}"
        }
        
    except ImportError as e:
        logger.error(f"emergentintegrations library error: {e}")
        raise HTTPException(status_code=500, detail="AI image integration library not available")
    except Exception as e:
        logger.error(f"AI image generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate image: {str(e)}")


@api_router.post("/migrate-asset-urls")
async def migrate_asset_urls():
    """
    Migrate all absolute asset URLs to relative URLs in the database.
    This fixes issues when the domain changes between sessions (e.g., after a fork).
    Handles both global (/api/assets/) and project-specific (/api/projects/{id}/assets/) URLs.
    Also normalizes image element src attributes.
    """
    import re
    
    migrated_count = 0
    projects = await db.projects.find({}).to_list(1000)
    
    for project in projects:
        updated = False
        course = project.get('course', {})
        slides = course.get('slides', [])
        
        for slide in slides:
            elements = slide.get('elements', [])
            for element in elements:
                # Fix HTML elements with embedded image URLs
                if element.get('type') == 'html' and element.get('htmlContent'):
                    html_content = element['htmlContent']
                    
                    # Strip domain from /api/assets/ URLs
                    new_content = re.sub(
                        r'https?://[^/\s"\']+/api/assets/',
                        '/api/assets/',
                        html_content
                    )
                    # Strip domain from /api/projects/ URLs
                    new_content = re.sub(
                        r'https?://[^/\s"\']+/api/projects/',
                        '/api/projects/',
                        new_content
                    )
                    
                    if new_content != html_content:
                        element['htmlContent'] = new_content
                        updated = True
                        migrated_count += 1
                
                # Fix image element src attributes
                src = element.get('src', '')
                if src and src.startswith('http') and '/api/' in src:
                    new_src = re.sub(r'https?://[^/\s"\']+(/api/.*)', r'\1', src)
                    if new_src != src:
                        element['src'] = new_src
                        updated = True
                        migrated_count += 1
            
            # Fix slide background images
            bg = slide.get('backgroundImage', '')
            if bg and bg.startswith('http') and '/api/' in bg:
                new_bg = re.sub(r'https?://[^/\s"\']+(/api/.*)', r'\1', bg)
                if new_bg != bg:
                    slide['backgroundImage'] = new_bg
                    updated = True
                    migrated_count += 1
        
        if updated:
            await db.projects.update_one(
                {'id': project['id']},
                {'$set': {'course': course}}
            )
            logger.info(f"Migrated {migrated_count} URLs in project {project.get('name', project['id'])}")
    
    return {
        "success": True,
        "message": f"Migrated {migrated_count} elements with asset URLs",
        "migrated_count": migrated_count
    }


# ============================================
# Quiz Generator Endpoints
# ============================================

@api_router.get("/questions")
async def list_questions(project_id: Optional[str] = None, tag: Optional[str] = None):
    """List all questions, optionally filtered by project or tag"""
    query = {}
    if project_id:
        query["projectId"] = project_id
    if tag:
        query["tags"] = tag
    
    questions = await db.questions.find(query, {"_id": 0}).sort("createdAt", -1).to_list(500)
    return questions

@api_router.get("/questions/{question_id}")
async def get_question(question_id: str):
    """Get a single question by ID"""
    question = await db.questions.find_one({"id": question_id}, {"_id": 0})
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
    return question

@api_router.post("/questions")
async def create_question(data: QuizQuestionCreate):
    """Create a new question manually"""
    # Build alternatives with IDs
    alternatives = []
    for alt in data.alternatives:
        alternatives.append(QuizAlternative(
            text=alt.get("text", ""),
            isCorrect=alt.get("isCorrect", False)
        ).model_dump())
    
    question = QuizQuestion(
        projectId=data.projectId,
        type=data.type,
        text=data.text,
        alternatives=alternatives,
        explanation=data.explanation,
        points=data.points,
        tags=data.tags
    )
    
    question_dict = question.model_dump()
    question_dict['createdAt'] = question.createdAt.isoformat()
    question_dict['updatedAt'] = question.updatedAt.isoformat()
    
    await db.questions.insert_one(question_dict)
    return serialize_doc(question_dict)

@api_router.put("/questions/{question_id}")
async def update_question(question_id: str, data: QuizQuestionUpdate):
    """Update an existing question"""
    question = await db.questions.find_one({"id": question_id})
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    update_data = data.model_dump(exclude_unset=True)
    
    # Handle alternatives update
    if "alternatives" in update_data and update_data["alternatives"]:
        alternatives = []
        for alt in update_data["alternatives"]:
            if "id" in alt:
                alternatives.append(alt)
            else:
                alternatives.append(QuizAlternative(
                    text=alt.get("text", ""),
                    isCorrect=alt.get("isCorrect", False)
                ).model_dump())
        update_data["alternatives"] = alternatives
    
    update_data["updatedAt"] = now_utc().isoformat()
    
    await db.questions.update_one({"id": question_id}, {"$set": update_data})
    return await db.questions.find_one({"id": question_id}, {"_id": 0})

@api_router.delete("/questions/{question_id}")
async def delete_question(question_id: str):
    """Delete a question"""
    result = await db.questions.delete_one({"id": question_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Question not found")
    return {"message": "Question deleted"}

@api_router.post("/questions/generate")
async def generate_questions_with_ai(request: QuizGenerateRequest):
    """Generate quiz questions using AI from prompt or document content"""
    
    emergent_key = os.environ.get('EMERGENT_LLM_KEY')
    if not emergent_key:
        raise HTTPException(status_code=500, detail="AI API key not configured")
    
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        # Determine question type instructions
        if request.questionType == "true_false":
            type_instruction = """Gere APENAS questões de Verdadeiro ou Falso.
Cada questão deve ter exatamente 2 alternativas: "Verdadeiro" e "Falso"."""
        elif request.questionType == "multiple_choice":
            type_instruction = """Gere APENAS questões de Múltipla Escolha.
Cada questão deve ter exatamente 4 alternativas, sendo apenas 1 correta."""
        else:  # mixed
            type_instruction = """Gere uma mistura de questões de Múltipla Escolha e Verdadeiro/Falso.
- Questões de múltipla escolha: 4 alternativas, 1 correta
- Questões Verdadeiro/Falso: 2 alternativas ("Verdadeiro" e "Falso")"""
        
        system_message = f"""Você é um especialista em criar questões de quiz educacionais.

{type_instruction}

REGRAS IMPORTANTES:
1. SEMPRE responda em português brasileiro
2. Crie questões claras e objetivas
3. Evite pegadinhas ou ambiguidades
4. Inclua uma breve explicação para cada resposta correta
5. As questões devem testar compreensão, não memorização

FORMATO DE RESPOSTA (JSON válido):
{{
  "questions": [
    {{
      "type": "multiple_choice",
      "text": "Pergunta aqui?",
      "alternatives": [
        {{"text": "Alternativa A", "isCorrect": false}},
        {{"text": "Alternativa B", "isCorrect": true}},
        {{"text": "Alternativa C", "isCorrect": false}},
        {{"text": "Alternativa D", "isCorrect": false}}
      ],
      "explanation": "Explicação de por que B é correta..."
    }},
    {{
      "type": "true_false",
      "text": "Afirmação verdadeira ou falsa?",
      "alternatives": [
        {{"text": "Verdadeiro", "isCorrect": true}},
        {{"text": "Falso", "isCorrect": false}}
      ],
      "explanation": "Explicação..."
    }}
  ]
}}

RESPONDA APENAS COM O JSON, SEM TEXTO ADICIONAL."""

        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"quiz-gen-{uuid.uuid4()}",
            system_message=system_message
        ).with_model("openai", "gpt-4o")
        
        # Build the prompt based on source
        if request.source == "document" and request.documentContent:
            full_prompt = f"""Com base no seguinte conteúdo de documento, gere {request.count} questões de quiz:

CONTEÚDO DO DOCUMENTO:
{request.documentContent}

Gere questões que testem a compreensão deste conteúdo."""
        else:
            full_prompt = f"""Gere {request.count} questões de quiz sobre o seguinte tema:

TEMA: {request.prompt}"""
            if request.context:
                full_prompt += f"\n\nCONTEXTO ADICIONAL: {request.context}"
        
        user_message = UserMessage(text=full_prompt)
        response = await chat.send_message(user_message)
        
        # Parse the JSON response
        try:
            # Clean up response if needed (remove markdown code blocks)
            cleaned_response = response.strip()
            if cleaned_response.startswith("```"):
                cleaned_response = cleaned_response.split("```")[1]
                if cleaned_response.startswith("json"):
                    cleaned_response = cleaned_response[4:]
            if cleaned_response.endswith("```"):
                cleaned_response = cleaned_response[:-3]
            
            parsed = json.loads(cleaned_response.strip())
            questions_data = parsed.get("questions", [])
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response as JSON: {e}")
            logger.error(f"Response was: {response[:500]}")
            raise HTTPException(status_code=500, detail="AI returned invalid format. Please try again.")
        
        # Save questions to database
        saved_questions = []
        for q_data in questions_data:
            # Build alternatives with IDs
            alternatives = []
            for alt in q_data.get("alternatives", []):
                alternatives.append(QuizAlternative(
                    text=alt.get("text", ""),
                    isCorrect=alt.get("isCorrect", False)
                ).model_dump())
            
            question = QuizQuestion(
                projectId=request.projectId,
                type=q_data.get("type", "multiple_choice"),
                text=q_data.get("text", ""),
                alternatives=alternatives,
                explanation=q_data.get("explanation"),
                tags=["ai-generated"]
            )
            
            question_dict = question.model_dump()
            question_dict['createdAt'] = question.createdAt.isoformat()
            question_dict['updatedAt'] = question.updatedAt.isoformat()
            
            await db.questions.insert_one(question_dict)
            saved_questions.append(serialize_doc(question_dict))
        
        logger.info(f"Generated {len(saved_questions)} quiz questions via AI")
        
        return {
            "success": True,
            "questions": saved_questions,
            "count": len(saved_questions)
        }
        
    except ImportError:
        logger.error("emergentintegrations library not installed")
        raise HTTPException(status_code=500, detail="AI integration library not available")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Quiz generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate questions: {str(e)}")

@api_router.post("/questions/parse-doc")
async def parse_doc_file(file: UploadFile = File(...)):
    """Parse a .doc/.docx file and extract text for quiz generation"""
    
    if not file.filename.lower().endswith(('.doc', '.docx')):
        raise HTTPException(status_code=400, detail="Only .doc and .docx files are accepted")
    
    try:
        from docx import Document
        
        content = await file.read()
        
        # Save temporarily
        temp_path = UPLOADS_DIR / f"temp_{uuid.uuid4()}_{file.filename}"
        async with aiofiles.open(temp_path, 'wb') as f:
            await f.write(content)
        
        # Parse the document
        doc = Document(str(temp_path))
        
        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        
        # Clean up temp file
        os.remove(temp_path)
        
        extracted_text = "\n\n".join(full_text)
        
        return {
            "success": True,
            "filename": file.filename,
            "text": extracted_text,
            "wordCount": len(extracted_text.split())
        }
        
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library not installed")
    except Exception as e:
        logger.error(f"Doc parsing error: {e}")
        # Clean up temp file if it exists
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")

@api_router.post("/quiz/submit")
async def submit_quiz(request: QuizSubmitRequest):
    """Submit quiz answers and calculate score"""
    
    # Get all questions for this quiz
    question_ids = [a["questionId"] for a in request.answers]
    questions = await db.questions.find({"id": {"$in": question_ids}}, {"_id": 0}).to_list(100)
    questions_map = {q["id"]: q for q in questions}
    
    # Calculate score
    total_points = 0
    earned_points = 0
    results = []
    
    for answer in request.answers:
        question = questions_map.get(answer["questionId"])
        if not question:
            continue
        
        total_points += question.get("points", 1)
        
        # Find the correct alternative
        correct_alt = None
        selected_alt = None
        for alt in question.get("alternatives", []):
            if alt.get("isCorrect"):
                correct_alt = alt
            if alt.get("id") == answer.get("selectedAlternativeId"):
                selected_alt = alt
        
        is_correct = selected_alt and selected_alt.get("isCorrect", False)
        if is_correct:
            earned_points += question.get("points", 1)
        
        results.append({
            "questionId": answer["questionId"],
            "questionText": question.get("text"),
            "selectedAlternativeId": answer.get("selectedAlternativeId"),
            "selectedText": selected_alt.get("text") if selected_alt else None,
            "correctAlternativeId": correct_alt.get("id") if correct_alt else None,
            "correctText": correct_alt.get("text") if correct_alt else None,
            "isCorrect": is_correct,
            "explanation": question.get("explanation")
        })
    
    # Calculate percentage and final score (0-10)
    percentage = (earned_points / total_points * 100) if total_points > 0 else 0
    final_score = round(percentage / 10, 1)  # Convert to 0-10 scale
    
    # Determine if passed (default passing score is 60%)
    passed = percentage >= 60
    
    # Save attempt to database
    attempt = QuizAttempt(
        quizId=request.quizId,
        projectId="",  # Could be extracted from quiz config
        answers=results,
        score=final_score,
        percentage=round(percentage, 1),
        passed=passed,
        completedAt=now_utc()
    )
    
    attempt_dict = attempt.model_dump()
    attempt_dict['createdAt'] = attempt.createdAt.isoformat()
    attempt_dict['completedAt'] = attempt.completedAt.isoformat() if attempt.completedAt else None
    
    await db.quiz_attempts.insert_one(attempt_dict)
    
    return {
        "success": True,
        "attemptId": attempt.id,
        "score": final_score,
        "percentage": round(percentage, 1),
        "passed": passed,
        "totalQuestions": len(results),
        "correctAnswers": sum(1 for r in results if r["isCorrect"]),
        "results": results
    }

@api_router.get("/quiz/attempts/{project_id}")
async def get_quiz_attempts(project_id: str, quiz_id: Optional[str] = None):
    """Get quiz attempts for a project"""
    query = {"projectId": project_id}
    if quiz_id:
        query["quizId"] = quiz_id
    
    attempts = await db.quiz_attempts.find(query, {"_id": 0}).sort("createdAt", -1).to_list(100)
    return attempts

# =============================================================================
# ElevenLabs Text-to-Speech Endpoints
# =============================================================================

@api_router.get("/elevenlabs/voices")
async def list_elevenlabs_voices(language: Optional[str] = None, gender: Optional[str] = None):
    """
    List available ElevenLabs voices filtered by gender.
    All voices support multiple languages via eleven_multilingual_v2 model.
    Gender: male, female, neutral
    """
    if not ELEVENLABS_API_KEY:
        raise HTTPException(status_code=400, detail="ElevenLabs API key not configured")
    
    try:
        from elevenlabs import ElevenLabs
        
        client = ElevenLabs(api_key=ELEVENLABS_API_KEY)
        voices_response = client.voices.get_all()
        
        voices = []
        available_genders = set()
        
        for voice in voices_response.voices:
            # Extract voice info
            voice_data = {
                "voice_id": voice.voice_id,
                "name": voice.name,
                "description": voice.description or "",
                "preview_url": voice.preview_url,
                "category": voice.category or "premade",
                "labels": voice.labels or {},
                "gender": None,
                "accent": None,
                "age": None,
                "multilingual": True  # All voices support multilingual with eleven_multilingual_v2
            }
            
            # Extract labels
            if voice.labels:
                voice_data["gender"] = voice.labels.get("gender", "").lower()
                voice_data["accent"] = voice.labels.get("accent", "")
                voice_data["age"] = voice.labels.get("age", "")
                
                # Track available values
                if voice_data["gender"]:
                    available_genders.add(voice_data["gender"])
            
            # Apply gender filter
            if gender:
                if voice_data["gender"] != gender.lower():
                    continue
            
            voices.append(voice_data)
        
        # Sort by name
        voices.sort(key=lambda x: x["name"])
        
        # Supported languages (all voices support these via multilingual model)
        supported_languages = [
            {"code": "pt-BR", "label": "🇧🇷 Português (Brasil)"},
            {"code": "en", "label": "🇺🇸 English"},
            {"code": "es", "label": "🇪🇸 Español"}
        ]
        
        return {
            "voices": voices,
            "total": len(voices),
            "supported_languages": supported_languages,
            "available_genders": sorted(list(available_genders)),
            "note": "All voices support Portuguese, English, and Spanish via eleven_multilingual_v2 model. The model auto-detects the language from your text."
        }
        
    except Exception as e:
        logger.error(f"Error fetching ElevenLabs voices: {e}")
        raise HTTPException(status_code=500, detail=f"Error fetching voices: {str(e)}")


class TTSRequest(BaseModel):
    text: str
    voice_id: str
    stability: float = 0.5
    similarity_boost: float = 0.75
    style: float = 0.0
    use_speaker_boost: bool = True


from pydantic import BaseModel as PydanticBaseModel

class TTSRequest(PydanticBaseModel):
    text: str
    voice_id: str
    stability: float = 0.5
    similarity_boost: float = 0.75
    style: float = 0.0
    use_speaker_boost: bool = True


@api_router.post("/elevenlabs/generate-speech")
async def generate_elevenlabs_speech(request: TTSRequest):
    """
    Generate text-to-speech audio using ElevenLabs.
    Returns audio as base64-encoded MP3.
    """
    if not ELEVENLABS_API_KEY:
        raise HTTPException(status_code=400, detail="ElevenLabs API key not configured")
    
    if not request.text.strip():
        raise HTTPException(status_code=400, detail="Text cannot be empty")
    
    try:
        from elevenlabs import ElevenLabs
        from elevenlabs.types import VoiceSettings
        
        client = ElevenLabs(api_key=ELEVENLABS_API_KEY)
        
        # Configure voice settings
        voice_settings = VoiceSettings(
            stability=request.stability,
            similarity_boost=request.similarity_boost,
            style=request.style,
            use_speaker_boost=request.use_speaker_boost
        )
        
        # Generate audio
        audio_generator = client.text_to_speech.convert(
            text=request.text,
            voice_id=request.voice_id,
            model_id="eleven_multilingual_v2",
            voice_settings=voice_settings
        )
        
        # Collect audio data
        audio_data = b""
        for chunk in audio_generator:
            audio_data += chunk
        
        # Convert to base64
        audio_b64 = base64.b64encode(audio_data).decode()
        
        # Generate unique filename
        audio_id = str(uuid.uuid4())
        filename = f"tts_{audio_id}.mp3"
        
        # Save to storage for later use
        audio_path = STORAGE_DIR / "audio" / filename
        audio_path.parent.mkdir(exist_ok=True)
        
        async with aiofiles.open(audio_path, "wb") as f:
            await f.write(audio_data)
        
        # Save metadata to database
        tts_record = {
            "id": audio_id,
            "voice_id": request.voice_id,
            "text": request.text[:500],  # Store first 500 chars
            "filename": filename,
            "file_path": str(audio_path),
            "file_size": len(audio_data),
            "created_at": now_utc().isoformat()
        }
        await db.tts_generations.insert_one(tts_record)
        
        return {
            "success": True,
            "audio_id": audio_id,
            "audio_url": f"/api/audio/{filename}",
            "audio_base64": f"data:audio/mpeg;base64,{audio_b64}",
            "text": request.text,
            "voice_id": request.voice_id,
            "file_size": len(audio_data)
        }
        
    except Exception as e:
        logger.error(f"Error generating TTS: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating speech: {str(e)}")


@api_router.get("/audio/{filename}")
async def get_audio_file(filename: str):
    """Serve generated audio file"""
    audio_path = STORAGE_DIR / "audio" / filename
    
    if not audio_path.exists():
        raise HTTPException(status_code=404, detail="Audio file not found")
    
    return FileResponse(
        audio_path,
        media_type="audio/mpeg",
        filename=filename
    )


@api_router.get("/elevenlabs/voices/recommended")
async def get_recommended_voices():
    """
    Get recommended voices for Portuguese (Brazil), English, and Spanish.
    Curated list of high-quality voices for narration.
    """
    if not ELEVENLABS_API_KEY:
        raise HTTPException(status_code=400, detail="ElevenLabs API key not configured")
    
    try:
        from elevenlabs import ElevenLabs
        
        client = ElevenLabs(api_key=ELEVENLABS_API_KEY)
        voices_response = client.voices.get_all()
        
        # Keywords for each language/gender combination
        filters = {
            "pt-BR": {
                "keywords": ["brazilian", "brazil", "português", "portuguese"],
                "voices": []
            },
            "en": {
                "keywords": ["english", "american", "british"],
                "voices": []
            },
            "es": {
                "keywords": ["spanish", "español", "latin"],
                "voices": []
            }
        }
        
        for voice in voices_response.voices:
            voice_text = f"{voice.name} {voice.description or ''}".lower()
            accent = (voice.labels or {}).get("accent", "").lower()
            voice_text += f" {accent}"
            
            gender = (voice.labels or {}).get("gender", "unknown").lower()
            
            voice_info = {
                "voice_id": voice.voice_id,
                "name": voice.name,
                "description": voice.description or "",
                "preview_url": voice.preview_url,
                "gender": gender,
                "accent": accent
            }
            
            for lang, data in filters.items():
                if any(kw in voice_text for kw in data["keywords"]):
                    data["voices"].append(voice_info)
        
        # Organize by language and gender
        recommended = {}
        for lang, data in filters.items():
            recommended[lang] = {
                "male": [v for v in data["voices"] if v["gender"] == "male"][:5],
                "female": [v for v in data["voices"] if v["gender"] == "female"][:5]
            }
        
        return {
            "recommended": recommended,
            "languages": [
                {"code": "pt-BR", "label": "🇧🇷 Português (Brasil)"},
                {"code": "en", "label": "🇺🇸 English"},
                {"code": "es", "label": "🇪🇸 Español"}
            ]
        }
        
    except Exception as e:
        logger.error(f"Error fetching recommended voices: {e}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


# =============================================================================
# AI Tutor - Chat endpoint for SCORM courses
# =============================================================================

@api_router.get("/admin/tutor-settings")
async def get_tutor_settings():
    """Get global AI tutor settings"""
    settings = await db.settings.find_one({"key": "tutor"}, {"_id": 0})
    if not settings:
        settings = {
            "key": "tutor",
            "enabled": True,
            "messageLimit": 50,
            "suggestedQuestions": [],
            "systemPrompt": "",
            "tutorName": "Tutor IA"
        }
    return settings


@api_router.put("/admin/tutor-settings")
async def update_tutor_settings(request: Request):
    """Update global AI tutor settings"""
    data = await request.json()
    data["key"] = "tutor"
    await db.settings.update_one(
        {"key": "tutor"},
        {"$set": data},
        upsert=True
    )
    return {"status": "ok", "message": "Tutor settings updated"}


@api_router.post("/tutor/chat")
async def tutor_chat(request: Request):
    """AI Tutor chat endpoint - called from SCORM packages running in LMS"""
    data = await request.json()
    user_message = data.get("message", "")
    course_topic = data.get("courseTopic", "")
    course_context = data.get("courseContext", "")
    history = data.get("history", [])
    session_id = data.get("sessionId", str(uuid.uuid4()))

    if not user_message:
        raise HTTPException(status_code=400, detail="Message is required")

    # Get tutor settings
    settings = await db.settings.find_one({"key": "tutor"}, {"_id": 0})
    if not settings or not settings.get("enabled", True):
        raise HTTPException(status_code=403, detail="AI Tutor is disabled")

    message_limit = settings.get("messageLimit", 50)
    tutor_name = settings.get("tutorName", "Tutor IA")
    custom_prompt = settings.get("systemPrompt", "")

    # Check message limit
    msg_count = len([m for m in history if m.get("role") == "user"])
    if msg_count >= message_limit:
        return {
            "response": f"Você atingiu o limite de {message_limit} mensagens para esta sessão. Revise o conteúdo do curso para continuar aprendendo!",
            "limitReached": True
        }

    emergent_key = os.environ.get('EMERGENT_LLM_KEY', '')
    if not emergent_key:
        raise HTTPException(status_code=500, detail="AI service not configured")

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage

        system_msg = f"""Você é o "{tutor_name}", um tutor educacional especializado e amigável.
Seu papel é ajudar alunos a entender o conteúdo do curso sobre: {course_topic}

REGRAS IMPORTANTES:
1. Responda SOMENTE com base no conteúdo fornecido abaixo em "CONTEÚDO DO CURSO". Use citações diretas e referências aos slides quando possível (ex: "Conforme apresentado no Slide 3...")
2. Se o aluno perguntar sobre algo que NÃO está no conteúdo do curso, diga claramente que esse tema não é abordado no curso e redirecione para tópicos que estão no material
3. Use linguagem clara e acessível
4. Dê exemplos práticos baseados no conteúdo real do curso
5. Incentive o aluno a revisar slides específicos quando relevante
6. Responda no mesmo idioma da pergunta do aluno
7. Mantenha respostas concisas (máximo 3 parágrafos)
8. Quando citar informações, mencione de qual slide a informação vem

CONTEÚDO DO CURSO:
{course_context}"""

        if custom_prompt:
            system_msg += f"\n\nINSTRUÇÕES ADICIONAIS DO INSTRUTOR:\n{custom_prompt}"

        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"tutor-{session_id}",
            system_message=system_msg
        ).with_model("gemini", "gemini-3-flash-preview")

        # Build conversation context from history
        for msg in history[-10:]:  # Last 10 messages for context
            if msg.get("role") == "user":
                await chat.send_message(UserMessage(text=msg["content"]))
            # Assistant messages are automatically tracked by LlmChat

        # Send current message
        response = await chat.send_message(UserMessage(text=user_message))

        return {
            "response": response,
            "limitReached": False,
            "messagesUsed": msg_count + 1,
            "messageLimit": message_limit
        }

    except Exception as e:
        logger.error(f"Tutor chat error: {e}")
        raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")


# =============================================================================
# AI INSTRUCTIONAL DESIGN AGENT
# =============================================================================
from pydantic import ConfigDict as PydanticConfigDict

class AgentSessionCreate(BaseModel):
    """Create agent session - optionally with initial text content"""
    contentText: Optional[str] = None
    fileName: Optional[str] = None

class AgentConfigUpdate(BaseModel):
    model_config = PydanticConfigDict(extra="allow")
    title: Optional[str] = None
    depth: Optional[str] = "intermediario"
    duration: Optional[int] = 30
    modules: Optional[int] = 3
    interactivity: Optional[str] = "media"
    visualStyle: Optional[str] = "moderno e profissional"
    format: Optional[str] = "curso_completo"
    description: Optional[str] = None

class AgentChatMessage(BaseModel):
    message: str

@api_router.post("/agent/sessions")
async def create_agent_session(data: AgentSessionCreate):
    """Create a new AI agent session."""
    session_id = str(uuid.uuid4())
    session = {
        "id": session_id,
        "step": "created",  # created, analyzed, configured, structured, storyboarded, generated
        "contentText": data.contentText or "",
        "fileName": data.fileName or "",
        "analysis": None,
        "config": {},
        "structure": None,
        "storyboard": None,
        "projectId": None,
        "chatHistory": [],
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    await db.agent_sessions.insert_one({**session, "_id": session_id})
    return session

@api_router.get("/agent/sessions/{session_id}")
async def get_agent_session(session_id: str):
    """Get agent session state."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    return s

@api_router.post("/agent/sessions/{session_id}/upload")
async def agent_upload_content(session_id: str, file: UploadFile = File(None), text: str = Form(None), url: str = Form(None)):
    """Upload content to agent session (file, text, or URL)."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    content_text = text or ""
    file_name = ""

    # Handle URL scraping
    if url and url.strip():
        file_name = url.strip()
        try:
            async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
                resp = await client.get(url.strip(), headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "Accept": "text/html,application/xhtml+xml",
                    "Accept-Language": "pt-BR,pt;q=0.9,en;q=0.8",
                })
                if resp.status_code == 200:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(resp.text, "html.parser")
                    # Remove scripts, styles, nav, footer
                    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe"]):
                        tag.decompose()
                    # Try article/main content first
                    main = soup.find("article") or soup.find("main") or soup.find("body")
                    if main:
                        content_text = main.get_text(separator="\n", strip=True)
                    else:
                        content_text = soup.get_text(separator="\n", strip=True)
                    # Clean up excessive whitespace
                    import re as _re
                    content_text = _re.sub(r'\n{3,}', '\n\n', content_text).strip()
                    content_text = content_text[:50000]  # Limit
                else:
                    content_text = f"[Erro ao acessar URL: HTTP {resp.status_code}]"
        except Exception as e:
            logger.warning(f"URL scraping failed: {e}")
            content_text = f"[Erro ao acessar URL: {str(e)[:200]}]"

    # Handle file upload
    elif file:
        file_name = file.filename or ""
        file_bytes = await file.read()
        ext = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else ""

        if ext == "txt":
            content_text = file_bytes.decode("utf-8", errors="ignore")

        elif ext == "pdf":
            # Native PDF extraction
            try:
                from PyPDF2 import PdfReader
                import io as _io
                reader = PdfReader(_io.BytesIO(file_bytes))
                pages = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
                content_text = "\n\n".join(pages)
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}")
                # Fallback to ConvertAPI
                content_text = await _convert_api_extract(file_name, file_bytes, ext)

        elif ext in ("docx",):
            # Native DOCX extraction
            try:
                from docx import Document
                import io as _io
                doc = Document(_io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            paragraphs.append(" | ".join(cells))
                content_text = "\n\n".join(paragraphs)
            except Exception as e:
                logger.warning(f"python-docx extraction failed: {e}")
                content_text = await _convert_api_extract(file_name, file_bytes, ext)

        elif ext in ("pptx", "ppt", "doc"):
            # Use ConvertAPI for PPT/legacy doc
            content_text = await _convert_api_extract(file_name, file_bytes, ext)

        else:
            content_text = file_bytes.decode("utf-8", errors="ignore")

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"contentText": content_text, "fileName": file_name, "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "ok", "contentLength": len(content_text), "fileName": file_name}


async def _convert_api_extract(file_name: str, file_bytes: bytes, ext: str) -> str:
    """Fallback: use ConvertAPI to extract text from a file."""
    convert_secret = os.environ.get("CONVERTAPI_SECRET", "")
    if not convert_secret:
        return f"[ConvertAPI não configurada - arquivo {file_name} recebido]"
    try:
        async with httpx.AsyncClient(timeout=120) as client:
            files_payload = {"File": (file_name, file_bytes)}
            resp = await client.post(
                f"https://v2.convertapi.com/convert/{ext}/to/txt?Secret={convert_secret}",
                files=files_payload,
            )
            if resp.status_code == 200:
                result = resp.json()
                text = ""
                for f_item in result.get("Files", []):
                    b64 = f_item.get("FileData", "")
                    if b64:
                        text += base64.b64decode(b64).decode("utf-8", errors="ignore")
                return text
    except Exception as e:
        logger.warning(f"ConvertAPI text extraction failed: {e}")
    return f"[Erro ao extrair texto de {file_name}]"

@api_router.post("/agent/sessions/{session_id}/analyze")
async def agent_analyze(session_id: str):
    """Step 1: Analyze content with AI."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    if not s.get("contentText"):
        raise HTTPException(400, "No content uploaded")

    from services.ai_agent import analyze_content
    analysis = await analyze_content(session_id, s["contentText"], s.get("fileName", ""))

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"analysis": analysis, "step": "analyzed", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return analysis

@api_router.post("/agent/sessions/{session_id}/configure")
async def agent_configure(session_id: str, data: AgentConfigUpdate):
    """Set course configuration parameters."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    config = data.model_dump(exclude_unset=True)
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"config": config, "step": "configured", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "ok", "config": config}

@api_router.post("/agent/sessions/{session_id}/generate-structure")
async def agent_generate_structure(session_id: str, request: Request):
    """Step 2: Generate course structure (optionally from a template)."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    # Check if body contains a templateId
    template_id = None
    try:
        body = await request.json()
        template_id = body.get("templateId")
    except Exception:
        pass

    if template_id:
        from services.ai_agent import generate_structure_from_template
        structure = await generate_structure_from_template(session_id, s["contentText"], s.get("config", {}), template_id)
        if not structure:
            from services.ai_agent import generate_structure
            structure = await generate_structure(session_id, s["contentText"], s.get("config", {}))
    else:
        from services.ai_agent import generate_structure
        structure = await generate_structure(session_id, s["contentText"], s.get("config", {}))

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"structure": structure, "step": "structured", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return structure

@api_router.post("/agent/sessions/{session_id}/generate-storyboard")
async def agent_generate_storyboard(session_id: str, background_tasks: BackgroundTasks):
    """Step 3: Generate detailed storyboard (runs in thread pool to avoid blocking event loop)."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    if not s.get("structure"):
        raise HTTPException(400, "Structure not generated yet")

    # If storyboard already exists, return it
    if s.get("step") == "storyboarded" and s.get("storyboard"):
        return {"status": "already_done", "message": "Storyboard already generated"}

    # If already generating, don't start another thread
    if s.get("step") == "storyboarding":
        return {"status": "processing", "message": "Storyboard generation already in progress"}

    # Mark as processing
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"step": "storyboarding", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )

    def _sync_generate():
        """Run storyboard generation in its own event loop inside a thread."""
        import asyncio as _asyncio
        loop = _asyncio.new_event_loop()
        _asyncio.set_event_loop(loop)
        try:
            from services.ai_agent import generate_storyboard
            from motor.motor_asyncio import AsyncIOMotorClient
            _client = AsyncIOMotorClient(os.environ.get("MONGO_URL"))
            _db = _client[os.environ.get("DB_NAME")]

            async def _progress(batch_num, total, message):
                await _db.agent_sessions.update_one(
                    {"id": session_id},
                    {"$set": {
                        "storyboardProgress": {"batch": batch_num, "total": total, "message": message},
                        "updatedAt": datetime.now(timezone.utc).isoformat()
                    }}
                )

            storyboard = loop.run_until_complete(
                generate_storyboard(session_id, s["contentText"], s["structure"], s.get("config", {}), progress_callback=_progress)
            )
            loop.run_until_complete(
                _db.agent_sessions.update_one(
                    {"id": session_id},
                    {"$set": {"storyboard": storyboard, "step": "storyboarded", "storyboardProgress": None, "updatedAt": datetime.now(timezone.utc).isoformat()}}
                )
            )
            _client.close()
        except Exception as e:
            err_msg = str(e)
            logger.error(f"Storyboard generation error: {err_msg}")
            error_detail = "Erro ao gerar storyboard."
            if "Budget" in err_msg or "budget" in err_msg:
                error_detail = "Orçamento da chave LLM excedido. Acesse Profile > Universal Key > Add Balance para adicionar saldo."
            elif "502" in err_msg or "BadGateway" in err_msg:
                error_detail = "Serviço de IA temporariamente indisponível (502). Tente novamente em alguns minutos."
            try:
                from motor.motor_asyncio import AsyncIOMotorClient
                _client = AsyncIOMotorClient(os.environ.get("MONGO_URL"))
                _db = _client[os.environ.get("DB_NAME")]
                loop.run_until_complete(
                    _db.agent_sessions.update_one(
                        {"id": session_id},
                        {"$set": {"step": "structured", "error": error_detail, "storyboardProgress": None, "updatedAt": datetime.now(timezone.utc).isoformat()}}
                    )
                )
                _client.close()
            except Exception:
                pass
        finally:
            loop.close()

    # Run in thread pool - this keeps the main event loop free for other requests
    import threading
    thread = threading.Thread(target=_sync_generate, daemon=True)
    thread.start()
    return {"status": "processing", "message": "Storyboard being generated..."}


@api_router.post("/agent/generate-bg-image")
async def agent_generate_bg_image(data: dict):
    """Generate a background image using AI based on a text prompt."""
    prompt = data.get("prompt", "")
    if not prompt:
        raise HTTPException(400, "Prompt required")
    try:
        from emergentintegrations.llm.openai.image_generation import OpenAIImageGeneration
        emergent_key = os.environ.get("EMERGENT_LLM_KEY", "")
        image_gen = OpenAIImageGeneration(api_key=emergent_key)
        full_prompt = f"Abstract artistic background for presentation slide: {prompt}. No text, no watermarks, smooth gradients, professional, suitable as slide background."
        images = await image_gen.generate_images(prompt=full_prompt, model="gpt-image-1", number_of_images=1)
        if images and len(images) > 0:
            img_data = images[0]
            import hashlib
            seed = hashlib.md5(prompt.encode()).hexdigest()[:10]
            fname = f"bg_ai_{seed}.png"
            fpath = os.path.join(str(PROJECTS_DIR), "bg_temp", fname)
            os.makedirs(os.path.dirname(fpath), exist_ok=True)
            with open(fpath, "wb") as f:
                f.write(img_data)
            return {"imageUrl": f"/api/projects/bg_temp/assets/{fname}", "imageBase64": f"data:image/png;base64,{base64.b64encode(img_data).decode()}"}
        raise HTTPException(500, "No image generated")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"BG image generation error: {e}")
        raise HTTPException(500, f"Failed to generate background image: {str(e)}")


@api_router.post("/agent/sessions/{session_id}/media-config")
async def agent_set_media_config(session_id: str, data: dict):
    """Save media and background configuration for each slide."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    media_config = data.get("mediaConfig", {})
    bg_config = data.get("bgConfig", {})
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"mediaConfig": media_config, "bgConfig": bg_config, "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "ok", "configured": len(media_config), "backgrounds": len(bg_config)}


@api_router.post("/agent/sessions/{session_id}/generate-course")
async def agent_generate_course(session_id: str, background_tasks: BackgroundTasks):
    """Step 5: Generate actual Scormfy project from storyboard with media."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")
    if not s.get("storyboard"):
        raise HTTPException(400, "Storyboard not generated yet")

    config = s.get("config", {})
    media_config = s.get("mediaConfig", {})
    bg_config = s.get("bgConfig", {})
    title = config.get("title", s.get("analysis", {}).get("title", "Curso Gerado por IA"))
    desc = config.get("description", s.get("analysis", {}).get("summary", ""))

    project = Project(name=title, description=desc)
    project_dir = PROJECTS_DIR / project.id
    (project_dir / "assets").mkdir(parents=True, exist_ok=True)

    from services.ai_agent import generate_course_from_storyboard
    course_data = await generate_course_from_storyboard(
        session_id, s["storyboard"], s.get("config", {}),
        project_dir=str(PROJECTS_DIR), project_id=project.id,
        media_config=media_config, bg_config=bg_config
    )

    project.course.metadata.title = title
    project.course.metadata.description = desc
    project.course.slides = []

    project_dict = project.model_dump()
    project_dict["createdAt"] = project.createdAt.isoformat()
    project_dict["updatedAt"] = project.updatedAt.isoformat()
    project_dict["course"]["createdAt"] = project.course.createdAt.isoformat()
    project_dict["course"]["updatedAt"] = project.course.updatedAt.isoformat()
    project_dict["course"]["slides"] = course_data["slides"]

    # Mark project as created by agent
    project_dict["createdByAgent"] = True
    project_dict["agentSessionId"] = session_id

    # Store HeyGen pending info if any
    heygen_pending = course_data.get("heygenPending", [])
    if heygen_pending:
        project_dict["heygenPending"] = heygen_pending

    await db.projects.insert_one(project_dict)

    # Save quiz questions if any
    for q in course_data.get("quizQuestions", []):
        q["projectId"] = project.id
        q["createdAt"] = datetime.now(timezone.utc).isoformat()
        q["updatedAt"] = datetime.now(timezone.utc).isoformat()
        await db.questions.insert_one({**q, "_id": q["id"]})

    # Trigger HeyGen video generation in background
    if heygen_pending and HEYGEN_API_KEY:
        background_tasks.add_task(_trigger_heygen_videos, project.id, heygen_pending)

    # Trigger narration generation in background if configured
    narration_enabled = config.get("narrationEnabled", False)
    narration_voice = config.get("narrationVoiceId", "")
    narration_count = 0
    if narration_enabled and narration_voice and ELEVENLABS_API_KEY:
        # Collect slides for narration
        import re as _re
        narration_tasks = []
        for si, slide_data in enumerate(course_data["slides"]):
            texts = []
            for el in slide_data.get("elements", []):
                html_content = el.get("htmlContent", "")
                if html_content:
                    clean = _re.sub(r'<[^>]+>', '', html_content)
                    clean = _re.sub(r'\s+', ' ', clean).strip()
                    if len(clean) > 30:
                        texts.append(clean)
            if texts:
                narration_tasks.append({
                    "slideIndex": si,
                    "slideId": slide_data.get("id", ""),
                    "slideTitle": slide_data.get("title", f"Slide {si+1}"),
                    "text": " ".join(texts)[:2000],
                    "status": "pending",
                })
        if narration_tasks:
            narration_count = len(narration_tasks)
            await db.projects.update_one(
                {"id": project.id},
                {"$set": {"narrationPending": narration_tasks, "narrationVoiceId": narration_voice}}
            )
            background_tasks.add_task(_generate_narrations, project.id, narration_tasks, narration_voice)

    # Update session
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"projectId": project.id, "step": "generated", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )

    # Auto-trigger improvement suggestions analysis in background
    background_tasks.add_task(_generate_improvement_suggestions, session_id, project.id)

    return {
        "status": "ok",
        "projectId": project.id,
        "projectName": title,
        "slidesCount": len(course_data["slides"]),
        "quizCount": len(course_data.get("quizQuestions", [])),
        "heygenPending": len(heygen_pending),
        "narrationPending": narration_count,
    }



async def _generate_improvement_suggestions(session_id: str, project_id: str):
    """Background task: analyze the course creation process and generate improvement suggestions."""
    try:
        session = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
        project = await db.projects.find_one({"id": project_id}, {"_id": 0, "course": 0})
        if not session:
            return

        # Build analysis context
        config = session.get("config", {})
        analysis = session.get("analysis", {})
        storyboard = session.get("storyboard", {})
        media_config = session.get("mediaConfig", {})
        bg_config = session.get("bgConfig", {})
        slides_count = len(storyboard.get("slides", []))
        modules_count = len(storyboard.get("modules", []))

        # Collect slide types
        slide_types = {}
        for s in storyboard.get("slides", []):
            t = s.get("type", "unknown")
            slide_types[t] = slide_types.get(t, 0) + 1

        # Media usage
        media_types = {}
        for v in media_config.values():
            t = v.get("type", "none")
            media_types[t] = media_types.get(t, 0) + 1

        # Custom backgrounds
        custom_bg_count = sum(1 for v in bg_config.values() if v.get("type", "default") != "default")

        context = f"""
ANÁLISE DO PROCESSO DE CRIAÇÃO DO CURSO:
- Título: {config.get('title', 'N/A')}
- Template: {config.get('template', 'N/A')}
- Tom: {config.get('tone', 'N/A')}
- Idioma: {config.get('language', 'pt-BR')}
- Módulos: {modules_count}
- Total de slides: {slides_count}
- Tipos de slides: {slide_types}
- Mídias configuradas: {media_types}
- Fundos personalizados: {custom_bg_count}
- Narração habilitada: {config.get('narrationEnabled', False)}
- HeyGen habilitado: {any(v.get('type') == 'heygen' for v in media_config.values())}
- Conteúdo fonte: {analysis.get('contentType', 'texto')}
- Palavras no conteúdo fonte: {analysis.get('wordCount', 'N/A')}
- Tópicos identificados: {', '.join(analysis.get('topics', [])[:5])}
- Nível de complexidade: {analysis.get('complexity', 'N/A')}

ESTRUTURA DOS MÓDULOS:
{chr(10).join(f"- {m.get('title', '?')}: {m.get('slidesCount', '?')} slides" for m in storyboard.get('modules', [])[:10])}
"""

        prompt = f"""Você é um consultor especialista em design instrucional e plataformas de e-learning.
Analise o seguinte processo de criação de curso e gere sugestões de melhoria detalhadas.

{context}

Gere sugestões organizadas nas seguintes categorias. Para cada sugestão, inclua:
- "title": título curto (max 60 chars)
- "description": descrição detalhada (2-3 frases)
- "priority": "alta", "media" ou "baixa"
- "impact": breve descrição do impacto esperado

Responda APENAS em JSON válido com esta estrutura exata:
{{
  "platform_ux": [sugestões para melhorar a experiência do usuário na plataforma Scormfy],
  "platform_features": [sugestões de novas funcionalidades para a plataforma],
  "platform_performance": [sugestões para melhorar performance e confiabilidade],
  "course_content": [sugestões para melhorar a qualidade do conteúdo do curso gerado],
  "course_design": [sugestões para melhorar o design visual do curso],
  "course_pedagogy": [sugestões para melhorar a metodologia pedagógica do curso]
}}

Gere 2-3 sugestões por categoria, totalizando 12-18 sugestões. Seja específico e actionable."""

        from emergentintegrations.llm.chat import LlmChat, UserMessage
        emergent_key = os.environ.get("EMERGENT_LLM_KEY", "")

        async def _try_llm(model_name):
            chat = LlmChat(
                api_key=emergent_key,
                session_id=f"suggestions_{session_id}",
                system_message="Você é um consultor especialista em design instrucional e plataformas de e-learning. Responda sempre em JSON válido.",
            ).with_model("openai", model_name)
            return await chat.send_message(UserMessage(text=prompt))

        try:
            raw = await _try_llm("gpt-5.2")
        except Exception as llm_err:
            logger.warning(f"Suggestions: gpt-5.2 failed ({str(llm_err)[:60]}), falling back to gpt-4o")
            raw = await _try_llm("gpt-4o")

        import json
        import re
        # Extract JSON from possible markdown blocks
        json_match = re.search(r"```json\s*([\s\S]*?)```", raw)
        if json_match:
            suggestions = json.loads(json_match.group(1).strip())
        else:
            suggestions = json.loads(raw.strip())

        # Store suggestions
        await db.agent_sessions.update_one(
            {"id": session_id},
            {"$set": {
                "suggestions": suggestions,
                "suggestionsGeneratedAt": datetime.now(timezone.utc).isoformat(),
                "updatedAt": datetime.now(timezone.utc).isoformat(),
            }}
        )
        logger.info(f"Improvement suggestions generated for session {session_id}")

    except Exception as e:
        logger.error(f"Failed to generate suggestions for session {session_id}: {e}")
        await db.agent_sessions.update_one(
            {"id": session_id},
            {"$set": {"suggestionsError": str(e), "updatedAt": datetime.now(timezone.utc).isoformat()}}
        )


@api_router.get("/agent/sessions/{session_id}/suggestions")
async def agent_get_suggestions(session_id: str):
    """Get improvement suggestions for a session."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0, "suggestions": 1, "suggestionsGeneratedAt": 1, "suggestionsError": 1})
    if not s:
        raise HTTPException(404, "Session not found")
    if s.get("suggestions"):
        return {"status": "ready", "suggestions": s["suggestions"], "generatedAt": s.get("suggestionsGeneratedAt")}
    if s.get("suggestionsError"):
        return {"status": "error", "error": s["suggestionsError"]}
    return {"status": "pending"}


@api_router.post("/agent/sessions/{session_id}/suggestions/regenerate")
async def agent_regenerate_suggestions(session_id: str, background_tasks: BackgroundTasks):
    """Regenerate improvement suggestions on demand."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0, "projectId": 1, "step": 1})
    if not s:
        raise HTTPException(404, "Session not found")
    if s.get("step") != "generated":
        raise HTTPException(400, "Course not generated yet")
    # Clear old suggestions
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$unset": {"suggestions": "", "suggestionsError": ""}, "$set": {"updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    background_tasks.add_task(_generate_improvement_suggestions, session_id, s.get("projectId", ""))
    return {"status": "regenerating"}


async def _trigger_heygen_videos(project_id: str, pending_list: list):
    """Background task: trigger HeyGen video generation for each pending slide."""
    for item in pending_list:
        try:
            async with httpx.AsyncClient(timeout=60.0) as http_client:
                payload = {
                    "video_inputs": [{
                        "character": {
                            "type": "avatar",
                            "avatar_id": item["avatar_id"],
                            "avatar_style": "normal"
                        },
                        "voice": {
                            "type": "text",
                            "input_text": item["script"],
                            "voice_id": item["voice_id"]
                        }
                    }],
                    "dimension": {"width": 1280, "height": 720},
                    "title": f"Agent-{project_id}-{item['slideId']}"
                }
                response = await http_client.post(
                    f"{HEYGEN_BASE_URL}/v2/video/generate",
                    headers=HEYGEN_HEADERS,
                    json=payload
                )
                if response.status_code == 200:
                    data = response.json()
                    video_id = data.get("data", {}).get("video_id")
                    if video_id:
                        await db.heygen_videos.insert_one({
                            "video_id": video_id,
                            "avatar_id": item["avatar_id"],
                            "voice_id": item["voice_id"],
                            "script": item["script"],
                            "title": item["title"],
                            "status": "processing",
                            "transparent": False,
                            "project_id": project_id,
                            "slide_id": item["slideId"],
                            "created_at": now_utc(),
                        })
                        # Update project with video_id mapping
                        await db.projects.update_one(
                            {"id": project_id, "heygenPending.slideId": item["slideId"]},
                            {"$set": {"heygenPending.$.videoId": video_id, "heygenPending.$.status": "processing"}}
                        )
                        logger.info(f"HeyGen video triggered: {video_id} for slide {item['slideId']}")
                else:
                    logger.error(f"HeyGen video generation failed: {response.status_code} - {response.text}")
                    await db.projects.update_one(
                        {"id": project_id, "heygenPending.slideId": item["slideId"]},
                        {"$set": {"heygenPending.$.status": "failed", "heygenPending.$.error": response.text[:200]}}
                    )
        except Exception as e:
            logger.error(f"HeyGen video trigger error: {e}")
            await db.projects.update_one(
                {"id": project_id, "heygenPending.slideId": item["slideId"]},
                {"$set": {"heygenPending.$.status": "failed", "heygenPending.$.error": str(e)[:200]}}
            )


@api_router.get("/agent/projects/{project_id}/heygen-status")
async def agent_heygen_status(project_id: str):
    """Check HeyGen video generation status for a project."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0, "id": 1, "heygenPending": 1})
    if not project:
        raise HTTPException(404, "Project not found")
    pending = project.get("heygenPending", [])
    if not pending:
        return {"status": "no_heygen", "videos": []}

    results = []
    all_done = True
    for item in pending:
        video_id = item.get("videoId")
        status = item.get("status", "pending")
        if video_id and status == "processing":
            # Check HeyGen API for status
            try:
                async with httpx.AsyncClient(timeout=30.0) as http_client:
                    resp = await http_client.get(
                        f"{HEYGEN_BASE_URL}/v1/video_status.get?video_id={video_id}",
                        headers=HEYGEN_HEADERS
                    )
                    if resp.status_code == 200:
                        vdata = resp.json().get("data", {})
                        status = vdata.get("status", "processing")
                        video_url = vdata.get("video_url")

                        # Update project pending status
                        update_fields = {"heygenPending.$.status": status}
                        if video_url:
                            update_fields["heygenPending.$.videoUrl"] = video_url
                        await db.projects.update_one(
                            {"id": project_id, "heygenPending.videoId": video_id},
                            {"$set": update_fields}
                        )

                        # If completed, update the slide element
                        if status == "completed" and video_url:
                            await _update_slide_with_heygen_video(project_id, item["slideId"], video_url)
                            await db.heygen_videos.update_one(
                                {"video_id": video_id},
                                {"$set": {"status": "completed", "video_url": video_url}}
                            )
            except Exception as e:
                logger.error(f"HeyGen status check error: {e}")

        if status not in ("completed", "failed"):
            all_done = False

        results.append({
            "slideId": item.get("slideId"),
            "slideIndex": item.get("slideIndex"),
            "title": item.get("title"),
            "videoId": item.get("videoId"),
            "status": status,
            "videoUrl": item.get("videoUrl"),
        })

    return {"status": "all_done" if all_done else "processing", "videos": results}


async def _update_slide_with_heygen_video(project_id: str, slide_id: str, video_url: str):
    """Replace the HeyGen processing element with actual video player."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        return

    slides = project.get("course", {}).get("slides", [])
    for si, slide in enumerate(slides):
        if slide.get("id") == slide_id:
            # Find and replace the heygen processing element
            for ei, el in enumerate(slide.get("elements", [])):
                html = el.get("htmlContent", "")
                if f'data-heygen-slide="{slide_id}"' in html:
                    from models import generate_id
                    new_html = f'''<div style="width:100%;height:100%;border-radius:12px;overflow:hidden;background:#000;position:relative;">
<video src="{video_url}" style="width:100%;height:100%;object-fit:cover;" controls autoplay muted></video>
<div style="position:absolute;bottom:0;left:0;right:0;padding:6px 12px;background:linear-gradient(transparent,rgba(0,0,0,0.7));">
<span style="color:rgba(255,255,255,0.6);font-size:11px;">Avatar HeyGen</span>
</div>
</div>'''
                    await db.projects.update_one(
                        {"id": project_id},
                        {"$set": {f"course.slides.{si}.elements.{ei}.htmlContent": new_html}}
                    )
                    logger.info(f"Updated slide {slide_id} with HeyGen video")
                    return


@api_router.post("/agent/projects/{project_id}/generate-narration")
async def agent_generate_narration(project_id: str, background_tasks: BackgroundTasks):
    """Trigger narration generation for all content slides of a project."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(404, "Project not found")

    # Get narration config from the agent session
    session_id = project.get("agentSessionId", "")
    session = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0, "config": 1}) if session_id else None
    config = session.get("config", {}) if session else {}
    voice_id = config.get("narrationVoiceId", "")

    if not voice_id:
        raise HTTPException(400, "No narration voice configured")
    if not ELEVENLABS_API_KEY:
        raise HTTPException(400, "ElevenLabs API key not configured")

    # Collect slides that need narration
    slides = project.get("course", {}).get("slides", [])
    narration_tasks = []
    for si, slide in enumerate(slides):
        # Skip slides that already have audio
        if slide.get("audio") and len(slide["audio"]) > 0:
            continue
        # Extract text from HTML elements
        texts = []
        for el in slide.get("elements", []):
            html_content = el.get("htmlContent", "")
            if html_content:
                import re as _re
                clean = _re.sub(r'<[^>]+>', '', html_content)
                clean = _re.sub(r'\s+', ' ', clean).strip()
                if len(clean) > 30:
                    texts.append(clean)
        if texts:
            narration_tasks.append({
                "slideIndex": si,
                "slideId": slide.get("id", ""),
                "slideTitle": slide.get("title", f"Slide {si+1}"),
                "text": " ".join(texts)[:2000],
                "status": "pending",
            })

    if not narration_tasks:
        return {"status": "no_slides", "message": "No slides need narration"}

    # Store narration tasks in the project
    await db.projects.update_one(
        {"id": project_id},
        {"$set": {"narrationPending": narration_tasks, "narrationVoiceId": voice_id}}
    )

    # Trigger background generation
    background_tasks.add_task(_generate_narrations, project_id, narration_tasks, voice_id)

    return {"status": "ok", "slides": len(narration_tasks), "voiceId": voice_id}


async def _generate_narrations(project_id: str, tasks: list, voice_id: str):
    """Background task: generate ElevenLabs narration for each slide."""
    try:
        from elevenlabs import ElevenLabs
        from elevenlabs.types import VoiceSettings

        el_client = ElevenLabs(api_key=ELEVENLABS_API_KEY)
        voice_settings = VoiceSettings(stability=0.5, similarity_boost=0.75, style=0.0, use_speaker_boost=True)

        for task in tasks:
            try:
                # Generate TTS
                audio_generator = el_client.text_to_speech.convert(
                    text=task["text"][:1500],
                    voice_id=voice_id,
                    model_id="eleven_multilingual_v2",
                    voice_settings=voice_settings
                )
                audio_data = b""
                for chunk in audio_generator:
                    audio_data += chunk

                # Save audio file
                audio_id = str(uuid.uuid4())
                filename = f"narration_{project_id}_{task['slideIndex']}.mp3"
                audio_path = STORAGE_DIR / "audio" / filename
                audio_path.parent.mkdir(exist_ok=True)

                async with aiofiles.open(audio_path, "wb") as f:
                    await f.write(audio_data)

                # Add audio to the slide
                slide_audio = {
                    "id": audio_id,
                    "url": f"/api/audio/{filename}",
                    "filename": filename,
                    "type": "narration",
                    "volume": 1.0,
                    "startTime": 0,
                    "duration": len(audio_data) / 16000,  # Approximate
                }
                await db.projects.update_one(
                    {"id": project_id},
                    {"$push": {f"course.slides.{task['slideIndex']}.audio": slide_audio}}
                )

                # Update status
                await db.projects.update_one(
                    {"id": project_id, "narrationPending.slideId": task["slideId"]},
                    {"$set": {"narrationPending.$.status": "completed", "narrationPending.$.audioUrl": f"/api/audio/{filename}"}}
                )
                logger.info(f"Narration generated for slide {task['slideIndex']} of project {project_id}")

            except Exception as e:
                logger.error(f"Narration error for slide {task['slideIndex']}: {e}")
                await db.projects.update_one(
                    {"id": project_id, "narrationPending.slideId": task["slideId"]},
                    {"$set": {"narrationPending.$.status": "failed", "narrationPending.$.error": str(e)[:200]}}
                )
    except Exception as e:
        logger.error(f"Narration generation failed: {e}")


@api_router.get("/agent/projects/{project_id}/narration-status")
async def agent_narration_status(project_id: str):
    """Check narration generation status for a project."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0, "id": 1, "narrationPending": 1})
    if not project:
        raise HTTPException(404, "Project not found")
    pending = project.get("narrationPending", [])
    if not pending:
        return {"status": "no_narration", "slides": []}

    completed = sum(1 for t in pending if t.get("status") == "completed")
    failed = sum(1 for t in pending if t.get("status") == "failed")
    total = len(pending)
    all_done = (completed + failed) == total

    return {
        "status": "all_done" if all_done else "processing",
        "total": total,
        "completed": completed,
        "failed": failed,
        "slides": [
            {"slideIndex": t.get("slideIndex"), "title": t.get("slideTitle"), "status": t.get("status", "pending"), "audioUrl": t.get("audioUrl")}
            for t in pending
        ],
    }

@api_router.post("/agent/sessions/{session_id}/chat")
async def agent_chat_endpoint(session_id: str, data: AgentChatMessage):
    """Chat with the agent for adjustments."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    from services.ai_agent import agent_chat
    context = {"structure": s.get("structure"), "config": s.get("config"), "analysis": s.get("analysis")}
    response = await agent_chat(session_id, data.message, context)

    # Save to chat history
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$push": {"chatHistory": {"role": "user", "text": data.message, "ts": datetime.now(timezone.utc).isoformat()}}}
    )
    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$push": {"chatHistory": {"role": "agent", "text": response, "ts": datetime.now(timezone.utc).isoformat()}}}
    )

    return {"response": response}


# --- Agent: Templates ---

@api_router.get("/agent/templates")
async def agent_list_templates():
    """List available course templates."""
    from services.ai_agent import get_templates
    return get_templates()


@api_router.post("/agent/sessions/{session_id}/generate-structure-from-template")
async def agent_generate_structure_from_template(session_id: str, data: dict):
    """Generate course structure using a template as base."""
    s = await db.agent_sessions.find_one({"id": session_id}, {"_id": 0})
    if not s:
        raise HTTPException(404, "Session not found")

    template_id = data.get("templateId")
    if not template_id:
        raise HTTPException(400, "templateId is required")

    from services.ai_agent import generate_structure_from_template
    structure = await generate_structure_from_template(session_id, s["contentText"], s.get("config", {}), template_id)
    if not structure:
        raise HTTPException(500, "Failed to generate structure from template")

    await db.agent_sessions.update_one(
        {"id": session_id},
        {"$set": {"structure": structure, "step": "structured", "updatedAt": datetime.now(timezone.utc).isoformat()}}
    )
    return structure


# --- Agent: Course Editing ---

@api_router.get("/agent/courses")
async def agent_list_courses():
    """List courses created by the AI agent."""
    projects = await db.projects.find(
        {"createdByAgent": True}, {"_id": 0, "id": 1, "name": 1, "description": 1, "createdAt": 1, "updatedAt": 1, "agentSessionId": 1}
    ).sort("createdAt", -1).to_list(100)
    # Add slide count
    for p in projects:
        full_proj = await db.projects.find_one({"id": p["id"]}, {"_id": 0, "course.slides": 1})
        p["slidesCount"] = len(full_proj.get("course", {}).get("slides", [])) if full_proj else 0
    return projects


class AgentImprovementsApply(BaseModel):
    improvements: list


@api_router.post("/agent/courses/{project_id}/analyze")
async def agent_analyze_course(project_id: str):
    """Analyze an existing agent-created course and suggest improvements."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(404, "Project not found")

    session_id = project.get("agentSessionId") or str(uuid.uuid4())
    from services.ai_agent import analyze_existing_course
    analysis = await analyze_existing_course(session_id, project)
    return analysis


@api_router.post("/agent/courses/{project_id}/apply-improvements")
async def agent_apply_improvements(project_id: str, data: AgentImprovementsApply):
    """Apply selected improvements to an existing course."""
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(404, "Project not found")

    session_id = project.get("agentSessionId") or str(uuid.uuid4())
    from services.ai_agent import apply_course_improvements
    from models import generate_id
    result = await apply_course_improvements(session_id, project, data.improvements)

    slides = project.get("course", {}).get("slides", [])

    # Apply updated slides
    for upd in result.get("updatedSlides", []):
        idx = upd.get("slideIndex")
        if idx is not None and 0 <= idx < len(slides):
            if upd.get("title"):
                slides[idx]["title"] = upd["title"]
            if upd.get("notes"):
                slides[idx]["notes"] = upd["notes"]
            if upd.get("narrationScript"):
                slides[idx]["librasScript"] = upd["narrationScript"]
            if upd.get("librasScript"):
                slides[idx]["librasScript"] = upd["librasScript"]
            # Rebuild elements
            if upd.get("elements"):
                new_elements = []
                for elem in upd["elements"]:
                    el = {
                        "id": generate_id(),
                        "type": "html",
                        "x": 560 if elem.get("position") == "center" else 100,
                        "y": 40,
                        "width": elem.get("width", 800),
                        "height": elem.get("height", 400),
                        "content": "",
                        "htmlContent": elem.get("content", ""),
                        "style": {"fontSize": 18, "fontFamily": "Inter, sans-serif", "fontColor": "#333333"},
                        "startTime": 0,
                        "animations": [],
                    }
                    new_elements.append(el)
                slides[idx]["elements"] = new_elements

    # Insert new slides
    new_slides_added = 0
    for ns in sorted(result.get("newSlides", []), key=lambda x: x.get("afterIndex", 999), reverse=True):
        after_idx = ns.get("afterIndex", len(slides) - 1)
        insert_at = min(after_idx + 1, len(slides))
        new_elements = []
        for elem in ns.get("elements", []):
            el = {
                "id": generate_id(),
                "type": "html",
                "x": 560 if elem.get("position") == "center" else 100,
                "y": 40,
                "width": elem.get("width", 800),
                "height": elem.get("height", 400),
                "content": "",
                "htmlContent": elem.get("content", ""),
                "style": {"fontSize": 18, "fontFamily": "Inter, sans-serif", "fontColor": "#333333"},
                "startTime": 0,
                "animations": [],
            }
            new_elements.append(el)
        new_slide = {
            "id": generate_id(),
            "title": ns.get("title", "Novo Slide"),
            "order": insert_at,
            "width": 1920,
            "height": 820,
            "background": ns.get("background", "#FFFFFF"),
            "elements": new_elements,
            "annotations": [],
            "transition": {"type": "fade", "duration": 0.5},
            "audio": [],
            "notes": "",
            "librasScript": ns.get("librasScript", ""),
            "duration": 5.0,
        }
        slides.insert(insert_at, new_slide)
        new_slides_added += 1

    # Re-order
    for i, s in enumerate(slides):
        s["order"] = i

    # Save
    course = project.get("course", {})
    course["slides"] = slides
    await update_project(project_id, {"course": course})

    return {
        "status": "ok",
        "updatedSlides": len(result.get("updatedSlides", [])),
        "newSlides": new_slides_added,
        "totalSlides": len(slides),
    }


# Include router
app.include_router(api_router)

# Root-level endpoints for Kubernetes health checks and startup detection
@app.get("/")
async def root():
    return {"status": "ok"}

@app.get("/health")
async def root_health():
    return {"status": "healthy"}

# CORS - Allow all origins for cross-domain production deployments
# With credentials=True, we must echo back the specific origin (not "*")
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=r"https?://.*",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    max_age=3600,
)


@app.on_event("startup")
async def startup_migrate_urls():
    """Auto-migrate absolute asset URLs to relative on startup (non-blocking background task)"""
    import asyncio
    asyncio.create_task(_run_migrate_urls())
    logger.info("Startup URL migration: started in background task")


async def _run_migrate_urls():
    """Background task: migrate absolute URLs to relative paths"""
    import re
    try:
        migrated_count = 0
        projects = await db.projects.find({}).to_list(1000)
        for project in projects:
            updated = False
            course = project.get('course', {})
            for slide in course.get('slides', []):
                for element in slide.get('elements', []):
                    if element.get('type') == 'html' and element.get('htmlContent'):
                        html = element['htmlContent']
                        new_html = re.sub(r'https?://[^/\s"\']+/api/assets/', '/api/assets/', html)
                        new_html = re.sub(r'https?://[^/\s"\']+/api/projects/', '/api/projects/', new_html)
                        if new_html != html:
                            element['htmlContent'] = new_html
                            updated = True
                            migrated_count += 1
                    src = element.get('src', '')
                    if src and src.startswith('http') and '/api/' in src:
                        new_src = re.sub(r'https?://[^/\s"\']+(/api/.*)', r'\1', src)
                        if new_src != src:
                            element['src'] = new_src
                            updated = True
                            migrated_count += 1
                bg = slide.get('backgroundImage', '')
                if bg and bg.startswith('http') and '/api/' in bg:
                    new_bg = re.sub(r'https?://[^/\s"\']+(/api/.*)', r'\1', bg)
                    if new_bg != bg:
                        slide['backgroundImage'] = new_bg
                        updated = True
                        migrated_count += 1
            if updated:
                await db.projects.update_one({'id': project['id']}, {'$set': {'course': course}})
        if migrated_count > 0:
            logger.info(f"Startup URL migration: fixed {migrated_count} absolute URLs to relative")
        else:
            logger.info("Startup URL migration: no absolute URLs found, all clean")
    except Exception as e:
        logger.warning(f"Startup URL migration failed (non-fatal): {e}")


@app.on_event("startup")
async def startup_ensure_admin():
    """Ensure super admin user exists in database (non-blocking background task)"""
    import asyncio
    asyncio.create_task(_run_ensure_admin())
    logger.info("Startup admin check: started in background task")


async def _run_ensure_admin():
    """Background task: ensure super admin user exists"""
    import bcrypt
    try:
        admin_email = "admin@scormify.com"
        existing_admin = await db.users.find_one({"email": admin_email})
        if not existing_admin:
            logger.info("Creating default super admin user...")
            password_hash = bcrypt.hashpw("admin123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            await db.users.insert_one({
                "user_id": "user_superadmin001",
                "email": admin_email,
                "name": "Super Admin",
                "picture": None,
                "companyId": None,
                "role": "super_admin",
                "passwordHash": password_hash,
                "isActive": True,
                "createdAt": datetime.now(timezone.utc),
                "updatedAt": datetime.now(timezone.utc)
            })
            logger.info("Super admin user created successfully")
        else:
            logger.info("Super admin user already exists")
    except Exception as e:
        logger.warning(f"Failed to ensure admin user (non-fatal): {e}")


@app.on_event("startup")
async def startup_persist_local_assets():
    """Persist local assets to MongoDB on startup in background (non-blocking)"""
    import threading
    
    def _persist_assets():
        try:
            from services.asset_store import store_asset_sync
            from pymongo import MongoClient
            
            _client = MongoClient(mongo_url)
            _db = _client[os.environ['DB_NAME']]
            
            total = 0
            for project_dir in PROJECTS_DIR.iterdir():
                if not project_dir.is_dir():
                    continue
                assets_dir = project_dir / "assets"
                if not assets_dir.exists():
                    continue
                project_id = project_dir.name
                for asset in assets_dir.iterdir():
                    if asset.is_file() and asset.suffix.lower() in (
                        '.png', '.jpg', '.jpeg', '.mp3', '.wav', '.ogg',
                        '.webm', '.mp4', '.gif', '.webp', '.svg'
                    ):
                        existing = _db.project_assets.find_one(
                            {"project_id": project_id, "filename": asset.name},
                            {"_id": 1}
                        )
                        if not existing:
                            store_asset_sync(mongo_url, os.environ['DB_NAME'], project_id, asset.name, str(asset))
                            total += 1
            _client.close()
            if total > 0:
                logger.info(f"Background asset persistence: saved {total} new assets to MongoDB")
            else:
                logger.info("Background asset persistence: all assets already in MongoDB")
        except Exception as e:
            logger.warning(f"Background asset persistence failed (non-fatal): {e}")
    
    thread = threading.Thread(target=_persist_assets, daemon=True)
    thread.start()
    logger.info("Startup asset persistence: started in background thread")


@app.on_event("startup")
async def startup_check_system_deps():
    """Check system dependencies in background (non-blocking)"""
    import threading
    def _check():
        try:
            from utils.system_deps import ensure_system_dependencies
            ensure_system_dependencies()
        except Exception as e:
            logger.warning(f"System dependency check failed (non-fatal): {e}")
    threading.Thread(target=_check, daemon=True).start()


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
